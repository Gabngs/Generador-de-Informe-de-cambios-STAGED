#!/usr/bin/env python3
"""
git_diff_to_docx.py
────────────────────
Lee 'informe.txt' (en la misma carpeta) generado con:
    git --no-pager diff --staged -U9999 > informe.txt

Genera un informe profesional .docx analizando la lógica
de los cambios, resumiendo archivos nuevos y detectando impacto.
"""

import sys
import re
import argparse
from datetime import datetime
from pathlib import Path
from dataclasses import dataclass, field
from typing import List, Tuple, Optional, Dict, Set

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# =============================================================================
# PALETA DE COLORES
# =============================================================================
C_TITLE     = RGBColor(0x1E, 0x3A, 0x5F)
C_SUBTITLE  = RGBColor(0x55, 0x55, 0x55)
C_BODY      = RGBColor(0x22, 0x22, 0x22)
C_MUTED     = RGBColor(0x88, 0x88, 0x88)
C_ADD_TEXT  = RGBColor(0x16, 0x65, 0x34)
C_ADD_BG    = "F0FDF4"
C_DEL_TEXT  = RGBColor(0x99, 0x1B, 0x1B)
C_DEL_BG    = "FEF2F2"
C_MOD_TEXT  = RGBColor(0x1E, 0x3A, 0x5F)
C_MOD_BG    = "EFF6FF"
C_REF_TEXT  = RGBColor(0x92, 0x40, 0x0E)
C_REF_BG    = "FFFBEB"
C_HDR_BG    = "1E3A5F"
C_ACCENT    = "2563EB"
C_BORDER    = "CCCCCC"
C_ROW_ALT   = "F8FAFC"
C_WHITE     = "FFFFFF"

# Colores para niveles de impacto
C_IMPACT_NULA     = RGBColor(0x6B, 0x72, 0x80)
C_IMPACT_LEVE     = RGBColor(0x05, 0x96, 0x69)
C_IMPACT_BAJA     = RGBColor(0x0D, 0x94, 0x88)
C_IMPACT_MEDIA    = RGBColor(0xD9, 0x77, 0x06)
C_IMPACT_ALTA     = RGBColor(0xDC, 0x25, 0x26)
C_IMPACT_CRITICA  = RGBColor(0x7F, 0x1D, 0x1D)

BG_IMPACT_NULA    = "F3F4F6"
BG_IMPACT_LEVE    = "ECFDF5"
BG_IMPACT_BAJA    = "F0FDFA"
BG_IMPACT_MEDIA   = "FFFBEB"
BG_IMPACT_ALTA    = "FFF1F2"
BG_IMPACT_CRITICA = "FEF2F2"

DEFAULT_INPUT  = "informe.txt"
DEFAULT_OUTPUT = "informe_cambios.docx"

ANSI_RE = re.compile(r'\x1B(?:[@-Z\\-_]|\[[0-?]*[ -/]*[@-~])')

MESES_ES = {
    1: "enero", 2: "febrero", 3: "marzo", 4: "abril",
    5: "mayo", 6: "junio", 7: "julio", 8: "agosto",
    9: "septiembre", 10: "octubre", 11: "noviembre", 12: "diciembre"
}

def fecha_espanol() -> str:
    hoy = datetime.now()
    return f"{hoy.day} de {MESES_ES[hoy.month]} de {hoy.year}"

def clean(text: str) -> str:
    text = text.replace('\r\n', '\n').replace('\r', '\n')
    text = ANSI_RE.sub('', text)
    return text.replace('\x00', '')

# =============================================================================
# NIVELES DE IMPACTO
# =============================================================================
def impact_colors(level: str) -> Tuple[RGBColor, str]:
    return {
        "Nula":    (C_IMPACT_NULA,    BG_IMPACT_NULA),
        "Leve":    (C_IMPACT_LEVE,    BG_IMPACT_LEVE),
        "Baja":    (C_IMPACT_BAJA,    BG_IMPACT_BAJA),
        "Media":   (C_IMPACT_MEDIA,   BG_IMPACT_MEDIA),
        "Alta":    (C_IMPACT_ALTA,    BG_IMPACT_ALTA),
        "Critica": (C_IMPACT_CRITICA, BG_IMPACT_CRITICA),
    }.get(level, (C_BODY, C_WHITE))

# =============================================================================
# MODELO: FileChange
# =============================================================================
@dataclass
class FileChange:
    filepath: str
    added:    List[str] = field(default_factory=list)
    removed:  List[str] = field(default_factory=list)
    added_with_line: List[Tuple[Optional[int], str]] = field(default_factory=list)
    removed_with_line: List[Tuple[Optional[int], str]] = field(default_factory=list)
    full_content: List[str] = field(default_factory=list)
    contexts: Set[str]  = field(default_factory=set)
    kind: str = "modified"
    is_binary: bool = False

    @property
    def filename(self) -> str:
        return Path(self.filepath).name

    @property
    def ext(self) -> str:
        name = self.filename.lower()
        for double in ('.component.ts', '.component.html', '.component.scss',
                       '.service.ts', '.spec.ts', '.module.ts', '.pipe.ts',
                       '.directive.ts', '.guard.ts', '.interceptor.ts',
                       '.reducer.ts', '.action.ts', '.effect.ts', '.selector.ts',
                       '.resolver.ts', '.model.ts', '.interface.ts', '.enum.ts',
                       '.helper.ts', '.util.ts', '.config.ts', '.constant.ts',
                       # PHP / Laravel
                       '.blade.php'):
            if name.endswith(double):
                return double
        return Path(self.filepath).suffix.lower()

    @property
    def is_lockfile(self) -> bool:
        """Detecta archivos de lock/dependencias que no aportan valor en detalle linea a linea."""
        return self.filename.lower() in (
            'package-lock.json', 'yarn.lock', 'pnpm-lock.yaml',
            'composer.lock', 'gemfile.lock', 'poetry.lock', 'cargo.lock',
            'packages.lock.json', 'shrinkwrap.json'
        )

    @property
    def is_environment_file(self) -> bool:
        """Detecta archivos de entorno/configuracion de ambiente."""
        name = self.filename.lower()
        return (name.startswith('environment') and name.endswith('.ts')) or \
               name in ('.env', '.env.local', '.env.production', '.env.staging', '.env.qa')

    @property
    def needs_structural_summary(self) -> bool:
        """
        FIX BUG 2: Umbral inteligente por tipo de archivo.
        Determina si el archivo debe mostrarse como resumen estructural
        en vez de linea por linea.
        """
        total = len(self.added) + len(self.removed)

        # Lock files: siempre resumen (nunca linea a linea)
        if self.is_lockfile:
            return True

        # Archivos grandes (mas de 60 lineas totales)
        if total > 60:
            return True

        # Archivos nuevos con mas de 15 lineas: resumen estructural
        if self.kind == 'added' and len(self.added) > 15:
            return True

        # Archivos con interfaces/tipos exportados (aunque sean pequenos)
        all_text = "\n".join(self.added)
        if re.search(r'export\s+(interface|type|enum|class)\b', all_text) and len(self.added) > 10:
            return True

        return False

    @property
    def kind_label(self) -> str:
        if self.kind == 'added':   return 'Adicion'
        if self.kind == 'deleted': return 'Eliminacion'
        if self.kind == 'renamed': return 'Renombrado'
        if self.is_binary:         return 'Binario/Media'
        n_add, n_del = len(self.added), len(self.removed)
        if n_del == 0 and n_add > 0: return 'Adicion'
        if n_add == 0 and n_del > 0: return 'Eliminacion'
        if n_del > n_add * 2 and n_del > 10: return 'Refactor'
        if n_add == 0 and n_del == 0: return 'Configuracion'
        return 'Modificacion'

    @property
    def kind_colors(self) -> Tuple[RGBColor, str]:
        lbl = self.kind_label
        if lbl == 'Adicion':     return C_ADD_TEXT, C_ADD_BG
        if lbl == 'Eliminacion': return C_DEL_TEXT, C_DEL_BG
        if lbl == 'Refactor':    return C_REF_TEXT, C_REF_BG
        return C_MOD_TEXT, C_MOD_BG

    def extract_structure(self) -> Dict[str, List[str]]:
        """Analiza logicamente el contenido para extraer estructura por tipo de archivo."""
        structure: Dict[str, List[str]] = {
            "imports": [],
            "entities": [],
            "decorators": [],
            "routes": [],
            "ui_components": [],
            "angular_bindings": [],
            "forms": [],
            "events": [],
            "css_selectors": [],
        }
        all_lines = self.added + self.removed

        import_pattern    = re.compile(r'^\s*(import|from|require\(|include|using)\b')
        entity_pattern    = re.compile(
            r'^\s*(export )?(class|def|function|interface|const \w+\s*=\s*\(|'
            r'let \w+\s*=\s*\(|async function|type\s+\w+\s*=|enum\s+\w+)'
        )
        decorator_pattern = re.compile(r'^\s*@\w+')
        route_pattern     = re.compile(
            r"(path\s*:\s*['\"]|route\s*\(|router\.(get|post|put|delete|patch)\s*\()", re.I
        )

        # HTML / Angular template
        ui_component_pattern = re.compile(r'<\s*(p-[\w-]+|app-[\w-]+)\b', re.I)
        angular_binding_pattern = re.compile(
            r'(\[\([^\)]*\)\]|\[[^\]]+\]|\([^\)]+\)|\*ng(?:If|For|Switch)|\bngModel\b|\bformControlName\b)'
        )
        event_pattern = re.compile(r'\((click|change|input|submit|keyup|keydown|blur|focus)\)\s*=')
        form_pattern = re.compile(r'<\s*(input|select|textarea|form|button|p-dropdown|p-inputtext|p-button)\b', re.I)

        # SCSS / CSS
        css_selector_pattern = re.compile(r'^\s*([.#][\w-]+|:host|::ng-deep|[a-zA-Z][\w-]*(?:\s+[a-zA-Z][\w-]*)?)\s*[{,]')

        for line in all_lines:
            line_stripped = line.strip()

            if import_pattern.search(line) and line not in structure["imports"]:
                structure["imports"].append(
                    line.strip()[:90] + ('...' if len(line) > 90 else '')
                )
            elif decorator_pattern.match(line) and line_stripped not in structure["decorators"]:
                structure["decorators"].append(line_stripped[:60])
            elif route_pattern.search(line) and line_stripped not in structure["routes"]:
                structure["routes"].append(line_stripped[:80])
            elif entity_pattern.search(line):
                clean_entity = line_stripped.split('{')[0].strip()
                if clean_entity not in structure["entities"]:
                    structure["entities"].append(clean_entity)

            if self.ext in ('.html', '.component.html'):
                comp_match = ui_component_pattern.search(line)
                if comp_match:
                    comp = comp_match.group(1).lower()
                    if comp not in structure["ui_components"]:
                        structure["ui_components"].append(comp)

                bind_matches = angular_binding_pattern.findall(line)
                for b in bind_matches:
                    b_clean = b.strip()
                    if b_clean and b_clean not in structure["angular_bindings"]:
                        structure["angular_bindings"].append(b_clean)

                ev_match = event_pattern.search(line)
                if ev_match:
                    ev = ev_match.group(1)
                    if ev not in structure["events"]:
                        structure["events"].append(ev)

                form_match = form_pattern.search(line)
                if form_match:
                    frm = form_match.group(1).lower()
                    if frm not in structure["forms"]:
                        structure["forms"].append(frm)

            if self.ext in ('.scss', '.css', '.component.scss'):
                sel_match = css_selector_pattern.search(line)
                if sel_match:
                    sel = sel_match.group(1).strip()
                    if sel and sel not in structure["css_selectors"]:
                        structure["css_selectors"].append(sel[:80])
        return structure

    def build_functional_summary(self) -> List[str]:
        """Genera una síntesis ejecutiva del cambio por archivo (sin detalle línea a línea)."""
        summary: List[str] = []
        n_add = len(self.added)
        n_del = len(self.removed)
        struct = self.extract_structure()

        if self.ext in ('.html', '.component.html'):
            if self.kind == 'added':
                summary.append("Se incorpora una nueva plantilla de interfaz para la funcionalidad del módulo.")
            elif self.kind == 'modified':
                summary.append("Se ajusta la estructura del template para mejorar interacción y visualización.")

            if struct["ui_components"]:
                comps = ", ".join(struct["ui_components"][:6])
                summary.append(f"Se integran componentes de UI: {comps}.")
            if struct["forms"]:
                controls = ", ".join(struct["forms"][:6])
                summary.append(f"Se agregan controles de captura/acción: {controls}.")
            if struct["angular_bindings"]:
                summary.append(
                    f"Se detectan bindings/directivas Angular ({len(struct['angular_bindings'])} usos) "
                    "para enlazar estado y eventos del componente."
                )
            if struct["events"]:
                events = ", ".join(struct["events"][:6])
                summary.append(f"Se registran eventos de interacción de usuario: {events}.")

        elif self.ext in ('.scss', '.css', '.component.scss'):
            if self.kind == 'added':
                summary.append("Se incorpora hoja de estilos asociada al componente para estandarizar la presentación.")
            if struct["css_selectors"]:
                sels = ", ".join(struct["css_selectors"][:6])
                summary.append(f"Se definen selectores/reglas relevantes: {sels}.")

        elif self.ext in ('.ts', '.component.ts', '.service.ts'):
            if struct["entities"]:
                ents = ", ".join(struct["entities"][:5])
                summary.append(f"Se incorporan/ajustan estructuras lógicas: {ents}.")
            if struct["imports"]:
                summary.append("Se actualizan dependencias para soportar la nueva lógica del archivo.")

        if not summary:
            if self.kind == 'added':
                summary.append("Se agrega archivo nuevo con estructura inicial funcional.")
            elif self.kind == 'deleted':
                summary.append("Se elimina archivo y su responsabilidad asociada del módulo.")
            else:
                summary.append("Se aplican ajustes internos en la implementación del archivo.")

        summary.append(f"Balance de cambio: +{n_add} líneas / -{n_del} líneas.")
        return summary[:5]

    def _extract_deleted_identifiers(self, line: str) -> List[str]:
        line = line.strip()
        m_ts_braces = re.search(r'import\s+\{([^}]+)\}', line)
        if m_ts_braces:
            return [x.strip() for x in m_ts_braces.group(1).split(',')]
        m_ts_simple = re.search(r'import\s+([a-zA-Z0-9_]+)\s+from', line)
        if m_ts_simple:
            return [m_ts_simple.group(1)]
        m_py_from = re.search(r'from\s+\S+\s+import\s+(.+)', line)
        if m_py_from:
            return [x.strip() for x in m_py_from.group(1).split(',')]
        m_var = re.search(
            r'(?:(?:private|public|protected|readonly|const|let|var|static)\s+)+([a-zA-Z0-9_]+)',
            line
        )
        if m_var:
            return [m_var.group(1)]
        m_simple_var = re.match(r'^([a-zA-Z0-9_]+)\s*[=:]', line)
        if m_simple_var:
            return [m_simple_var.group(1)]
        return []

    def verify_dead_code(self, deleted_line: str) -> bool:
        if not self.full_content:
            return False
        identifiers = [i for i in self._extract_deleted_identifiers(deleted_line) if i]
        if not identifiers:
            return False
        for ident in identifiers:
            pattern = re.compile(rf'\b{re.escape(ident)}\b')
            if not any(pattern.search(line) for line in self.full_content):
                return True
        return False

    def detect_linter_fix(self, removed_line: str) -> str:
        """
        Detecta si la eliminacion de una linea fue por correccion de linter/formatter.
        FIX BUG 3: Corregido falso positivo en regla eqeqeq cuando r_line == a_line
                   y la linea no contiene operadores de comparacion.
        """
        if not self.added:
            return ""
        r_line = removed_line.strip()
        if not r_line:
            return ""

        for a_line in self.added:
            a_line = a_line.strip()
            if not a_line:
                continue

            # (regex negativo lookahead/lookbehind para evitar === y !==)
            has_equality_op = bool(re.search(r'(?<![=!<>])={2}(?!=)|(?<!!)!={1}(?!=)', r_line))
            if has_equality_op:
                r_eq = r_line.replace('!==', '\x00').replace('===', '\x01')
                r_eq = r_eq.replace('!=', '!==').replace('==', '===')
                r_eq = r_eq.replace('\x00', '!==').replace('\x01', '===')
                # Solo reportar si realmente cambio algo
                if r_eq == a_line and r_eq != r_line:
                    return "ESLint: Igualdad estricta (=== / !==)"

            # 2. ESLint prefer-const / no-var
            if 'let' in r_line and re.sub(r'\blet\b', 'const', r_line) == a_line:
                return "ESLint: Usar const en lugar de let"
            if 'var' in r_line and re.sub(r'\bvar\b', 'let', r_line) == a_line:
                return "ESLint: Usar let en lugar de var"
            if 'var' in r_line and re.sub(r'\bvar\b', 'const', r_line) == a_line:
                return "ESLint: Usar const en lugar de var"

            # 3. Prettier quotes
            if "'" in r_line and r_line.replace("'", '"') == a_line:
                return "Prettier: Cambio a comillas dobles"
            if '"' in r_line and r_line.replace('"', "'") == a_line:
                return "Prettier: Cambio a comillas simples"

            # 4. Prettier semi
            if r_line + ';' == a_line:
                return "Prettier: Agregar punto y coma"
            if r_line == a_line + ';':
                return "Prettier: Eliminar punto y coma"

            # 5. Prettier: espaciado / indentacion (solo si las lineas son distintas)
            r_ns = re.sub(r'\s+', '', r_line)
            a_ns = re.sub(r'\s+', '', a_line)
            if r_ns == a_ns and r_line != a_line and len(r_ns) > 2:
                return "Prettier: Arreglo de espaciado/indentacion"

            # 6. ESLint trailing comma
            if (r_line.rstrip(',') == a_line.rstrip(',') and r_line != a_line
                    and r_line.endswith(',') != a_line.endswith(',') and len(r_line) > 2):
                return "ESLint: Trailing comma"

            # 7. TSLint: tipo any -> unknown
            if 'any' in r_line and re.sub(r'\bany\b', 'unknown', r_line) == a_line:
                return "TSLint: Reemplazar 'any' por 'unknown'"

            # 8. ESLint no-console
            if re.match(r'^console\.(log|warn|error|info|debug)\s*\(', r_line):
                return "ESLint: Eliminar console.log/warn/error"

            # 9. Prettier: template literals
            if '`' in r_line and r_line.replace('`', '"') == a_line:
                return "Prettier: Template literal a comillas dobles"
            if '"' in r_line and r_line.replace('"', '`') == a_line:
                return "Prettier: Comillas a template literal"

            # 10. ESLint prefer-arrow-callback
            m_func  = re.match(r'function\s*\(([^)]*)\)\s*\{(.+)\}', r_line)
            m_arrow = re.match(r'\(([^)]*)\)\s*=>\s*\{?(.+)\}?', a_line)
            if m_func and m_arrow and m_func.group(1).strip() == m_arrow.group(1).strip():
                return "ESLint: Convertir a arrow function"

            # 11. TSLint: eliminar modificador 'public' implicito
            if 'public' in r_line and re.sub(r'\bpublic\b\s*', '', r_line).strip() == a_line.strip():
                return "TSLint: Eliminar modificador 'public' implicito"

            # 12. ESLint: object shorthand { x: x } -> { x }
            m_long  = re.search(r'\{\s*(\w+)\s*:\s*\1\s*\}', r_line)
            m_short = re.search(r'\{\s*(\w+)\s*\}', a_line)
            if m_long and m_short and m_long.group(1) == m_short.group(1):
                return "ESLint: Object shorthand"

            # 13. Prettier: parentesis innecesarios en arrow de un parametro
            if '=>' in r_line and re.sub(r'\((\w+)\)\s*=>', r'\1 =>', r_line) == a_line:
                return "Prettier: Eliminar parentesis en arrow function de un parametro"

        return ""

    def classify_removed_line(self, line: str) -> str:
        """Clasifica la razon de eliminacion con mayor precision."""
        linter = self.detect_linter_fix(line)
        if linter:
            return f"[Linter] {linter}"
        if self.verify_dead_code(line):
            return "[Limpieza] Codigo sin uso detectado"

        stripped = line.strip()
        if stripped.startswith('//') or stripped.startswith('#') \
                or stripped.startswith('*') or stripped.startswith('/*'):
            return "[Doc] Comentario o documentacion eliminada"
        if re.match(r'(console\.(log|warn|error|debug|info)|print\(|logger\.|Log\.)', stripped):
            return "[Debug] Traza o log de depuracion eliminada"
        if re.search(r'\b(TODO|FIXME|HACK|XXX|TEMP)\b', stripped, re.I):
            return "[Deuda tecnica] Comentario TODO/FIXME eliminado"
        if re.search(r'\b(isDevMode|environment\.|process\.env|DEBUG|FEATURE_FLAG)\b', stripped, re.I):
            return "[Config] Flag de entorno o feature flag"
        if stripped.startswith('//') and re.search(r'[({;=]', stripped):
            return "[Refactor] Codigo comentado eliminado"
        if re.search(r'\b(mock|stub|fake|dummy|hardcoded|temp|test_data)\b', stripped, re.I):
            return "[Test] Dato de prueba o mock eliminado"
        return ""


# =============================================================================
# PARSER DE DIFF
# =============================================================================
class DiffParser:
    _FILE   = re.compile(r'^diff --git a/(.+?) b/(.+)$')
    _NEW    = re.compile(r'^new file mode')
    _DEL    = re.compile(r'^deleted file mode')
    _REN    = re.compile(r'^rename to (.+)$')
    _BIN    = re.compile(r'^Binary files.*differ$')
    _HUNK   = re.compile(r'^@@ -(\d+)(?:,(\d+))? \+(\d+)(?:,(\d+))? @@\s*(.*)$')
    _PLUS3  = re.compile(r'^\+\+\+')
    _MIN3   = re.compile(r'^---')

    def parse(self, text: str) -> List[FileChange]:
        text = clean(text)
        files: List[FileChange] = []
        cur: Optional[FileChange] = None
        old_line_no: Optional[int] = None
        new_line_no: Optional[int] = None

        for line in text.splitlines():
            m_file = self._FILE.match(line)
            if m_file:
                cur = FileChange(filepath=m_file.group(2).strip())
                files.append(cur)
                old_line_no = None
                new_line_no = None
                continue
            if cur is None:
                continue
            if self._NEW.match(line):
                cur.kind = 'added'
            elif self._DEL.match(line):
                cur.kind = 'deleted'
            elif self._BIN.match(line):
                cur.is_binary = True
            elif self._REN.match(line):
                cur.kind = 'renamed'
                cur.filepath = self._REN.match(line).group(1).strip()
            elif m_hunk := self._HUNK.match(line):
                old_line_no = int(m_hunk.group(1))
                new_line_no = int(m_hunk.group(3))
                hint = m_hunk.group(5).strip()
                if hint and len(hint) > 2:
                    cur.contexts.add(hint[:60])
            elif self._PLUS3.match(line) or self._MIN3.match(line):
                continue
            elif line.startswith('+') and not line.startswith('+++'):
                s = line[1:].strip()
                if s:
                    cur.added.append(s)
                    cur.added_with_line.append((new_line_no, s))
                    cur.full_content.append(s)
                if new_line_no is not None:
                    new_line_no += 1
            elif line.startswith('-') and not line.startswith('---'):
                s = line[1:].strip()
                if s:
                    cur.removed.append(s)
                    cur.removed_with_line.append((old_line_no, s))
                if old_line_no is not None:
                    old_line_no += 1
            elif line.startswith(' '):
                s = line[1:].strip()
                if s:
                    cur.full_content.append(s)
                if old_line_no is not None:
                    old_line_no += 1
                if new_line_no is not None:
                    new_line_no += 1

        return [f for f in files if f.filepath]


# =============================================================================
# CATEGORIAS DE ARCHIVOS
# =============================================================================
FILE_CATEGORIES: Dict[str, str] = {
    '.component.html':  'Template Angular',
    '.component.ts':    'Componente Angular',
    '.component.scss':  'Estilos Componente',
    '.service.ts':      'Servicio Angular',
    '.spec.ts':         'Test Unitario',
    '.module.ts':       'Modulo Angular',
    '.pipe.ts':         'Pipe Angular',
    '.directive.ts':    'Directiva Angular',
    '.guard.ts':        'Guard de Ruta',
    '.interceptor.ts':  'Interceptor HTTP',
    '.reducer.ts':      'Reducer NgRx',
    '.action.ts':       'Accion NgRx',
    '.effect.ts':       'Efecto NgRx',
    '.selector.ts':     'Selector NgRx',
    '.resolver.ts':     'Resolver Angular',
    '.model.ts':        'Modelo de Datos',
    '.interface.ts':    'Interfaz TypeScript',
    '.enum.ts':         'Enumeracion',
    '.helper.ts':       'Helper/Utilidad',
    '.util.ts':         'Utilidad',
    '.config.ts':       'Configuracion',
    '.constant.ts':     'Constantes',
    '.html':            'Template HTML',
    '.scss':            'Estilos SCSS',
    '.css':             'Estilos CSS',
    '.ts':              'TypeScript',
    '.js':              'JavaScript',
    '.py':              'Python',
    '.json':            'Configuracion JSON',
    '.md':              'Documentacion',
    '.sql':             'Base de Datos',
    '.yml':             'Config YAML',
    '.yaml':            'Config YAML',
    '.env':             'Variables de Entorno',
    '.sh':              'Script Shell',
    '.xml':             'XML / Config',
    '.graphql':         'Esquema GraphQL',
    '.prisma':          'Schema Prisma ORM',
    # PHP / Laravel
    '.php':             'PHP Backend',
    '.blade.php':       'Vista Blade PHP',
}

def get_category(fc: FileChange) -> str:
    if fc.is_lockfile:
        return 'Dependencias (Lock)'
    if fc.is_environment_file:
        return 'Archivo de Entorno'
    return FILE_CATEGORIES.get(fc.ext, Path(fc.filepath).suffix.upper() or 'Archivo')


# =============================================================================
# SENALES DE IMPACTO TECNICO
# =============================================================================
IMPACT_SIGNALS: List[Tuple[re.Pattern, str]] = [
    (re.compile(r'\.subscribe\s*\('),
     'Flujo asincrono RxJS'),
    (re.compile(r'catchError|throwError|try\s*{|except\s+|\.catch\s*\('),
     'Gestion de errores y excepciones'),
    (re.compile(r'\bAuthService\b|\bAuthGuard\b|inject\s*\(.*Auth|new\s+Auth\w+'),
     'Seguridad: Servicio de autenticacion'),
    (re.compile(r'\bJWT\b|jsonwebtoken|\.sign\s*\(|\.verify\s*\(|bearerToken|refreshToken'),
     'Seguridad: Manejo de JWT/tokens'),
    (re.compile(r'(?<!\w)password(?!\w)|\bhash\b|\bbcrypt\b|\bsalt\b|\bcipher\b', re.I),
     'Seguridad: Credenciales o hashes'),
    (re.compile(r'apiLoginUrl|apiBackend|apiUrl\b|API_URL|baseUrl|apiSIAW|apiSIP', re.I),
     'Variable de entorno o endpoint de API'),
    (re.compile(r'router\.navigate|HttpResponse|location\.href|\[routerLink\]'),
     'Logica de ruteo o navegacion'),
    (re.compile(r'HttpClient|http\.(get|post|put|delete|patch)\s*[(<]', re.I),
     'Llamada HTTP a API externa'),
    (re.compile(r'console\.(log|warn|error)|print\(|logger\.'),
     'Trazabilidad (logs de depuracion)'),
    (re.compile(r'SELECT\b|INSERT\b|UPDATE\b|DELETE\b|JOIN\b|\.save\s*\(|\.find\s*\(', re.I),
     'Consulta u operacion de base de datos'),
    (re.compile(r'\bexport\s+(class|interface|type|enum)\b'),
     'Contrato exportado publicamente'),
    (re.compile(r'@Input\s*\(\)|@Output\s*\(\)|EventEmitter|@Prop\s*\(\)'),
     'Contrato de componente (inputs/outputs)'),
    (re.compile(r'@NgModule\s*\(|providers\s*:\s*\[|declarations\s*:\s*\['),
     'Configuracion de modulo Angular'),
    (re.compile(r'environment\.(prod|production|cloud|staging|qa)', re.I),
     'Configuracion de ambiente especifico'),
    (re.compile(r'dispatch\s*\(|store\.select|createAction|createReducer|createEffect'),
     'Estado global NgRx'),
    (re.compile(r'migration|schema|alembic|flyway|liquibase|ALTER TABLE|DROP TABLE', re.I),
     'Migracion de base de datos'),
    (re.compile(r'cron|@Scheduled|scheduler|setInterval|setTimeout'),
     'Tarea programada/scheduler'),
    (re.compile(r'WebSocket|socket\.io|ws://|wss://'),
     'Comunicacion en tiempo real (WebSocket)'),
    (re.compile(r'\bcache\b|redis|memcache|localStorage|sessionStorage', re.I),
     'Capa de cache o almacenamiento'),
    (re.compile(r'\bi18n\b|translate\.|locale|l10n', re.I),
     'Internacionalizacion (i18n)'),
    (re.compile(r'canActivate|canLoad|canMatch|menuGuard|authGuard'),
     'Guards de ruta (control de acceso)'),
    # --- PHP / Laravel ---
    (re.compile(r'public\s+function\s+\w+|private\s+function\s+\w+|protected\s+function\s+\w+'),
     'Modificacion de metodo de clase PHP'),
    (re.compile(r'\?\?|\?\?=|\?->|\$\w+\s*\?\s*\$\w+->|\$\w+\s*\?\s*\$\w+\s*:\s*', re.I),
     'Guard de nulidad / null-coalescing en logica PHP'),
    (re.compile(r'->where\s*\(|->select\s*\(|->with\s*\(|->find\s*\(|->create\s*\(|->update\s*\(|->save\s*\(|->delete\s*\(', re.I),
     'Operacion Eloquent ORM (consulta/persistencia)'),
    (re.compile(r'Route::(get|post|put|delete|patch|resource|apiResource)\s*\(', re.I),
     'Definicion de ruta de API/web Laravel'),
    (re.compile(r'->middleware\s*\(|\bMiddleware\b', re.I),
     'Middleware de autenticacion/autorizacion PHP'),
    (re.compile(r'return\s+response\s*\(|->json\s*\(|->response\s*\(', re.I),
     'Respuesta de endpoint HTTP'),
    (re.compile(r'->belongsTo\s*\(|->hasMany\s*\(|->hasOne\s*\(|->belongsToMany\s*\(|->morphMany\s*\('),
     'Relacion Eloquent entre modelos'),
    (re.compile(r'\$request->|Request\s+\$\w+|->validated\s*\(|->rules\s*\('),
     'Procesamiento/validacion de peticion HTTP'),
    (re.compile(r'\bLog::|\$this->error\b|\$this->success\b|ApiResponse', re.I),
     'Trazabilidad / respuesta estandarizada de API PHP'),
    (re.compile(r'->paginate\s*\(|->get\s*\(|->first\s*\(|->count\s*\(', re.I),
     'Consulta de coleccion Eloquent'),
]
# =============================================================================
# MOTOR DE ANALISIS SEMANTICO CONTEXTUAL
# =============================================================================

class SemanticInsightEngine:
    """
    Motor heurístico que intenta deducir la intención del cambio.
    No depende de ejemplos específicos.
    Detecta patrones arquitectónicos, de layout, estado y diseño.
    """

    def analyze(self, fc: FileChange) -> List[str]:
        insights = []

        added = "\n".join(fc.added)
        removed = "\n".join(fc.removed)
        full = added + "\n" + removed

        # ---------------------------------------------------------
        # 1. Migración de lógica TS hacia CSS
        # ---------------------------------------------------------
        if fc.ext in ('.component.ts', '.ts'):
            if re.search(r'\b(height|width)\s*:', removed) and \
               re.search(r'calc\(.*vh', added):
                insights.append(
                    "Se migra control de dimensiones desde lógica TypeScript "
                    "hacia cálculo dinámico en CSS, favoreciendo separación de responsabilidades."
                )

        # ---------------------------------------------------------
        # 2. Simplificación de binding Angular
        # ---------------------------------------------------------
        if fc.ext in ('.html', '.component.html'):
            if re.search(r'\[.*\]=', removed) and not re.search(r'\[.*\]=', added):
                insights.append(
                    "Se reduce uso de binding dinámico, simplificando el template "
                    "y disminuyendo dependencia del estado del componente."
                )

        # ---------------------------------------------------------
        # 3. Eliminación de propiedad posiblemente obsoleta
        # ---------------------------------------------------------
        if fc.ext in ('.ts', '.component.ts'):
            if re.search(r'\b(public|private|protected)?\s*\w+\s*:', removed):
                insights.append(
                    "Se elimina propiedad del componente, posible reducción "
                    "de estado interno o eliminación de lógica innecesaria."
                )

        # ---------------------------------------------------------
        # 4. Consolidación de estilos
        # ---------------------------------------------------------
        if fc.ext in ('.scss', '.css', '.component.scss'):
            if removed.count('}') > added.count('}'):
                insights.append(
                    "Se reduce cantidad de bloques CSS, indicando consolidación "
                    "o simplificación de reglas de estilo."
                )

        # ---------------------------------------------------------
        # 5. Eliminación de duplicidad estructural
        # ---------------------------------------------------------
        if len(set(fc.removed)) < len(fc.removed):
            insights.append(
                "Se eliminan líneas repetidas, posible corrección de duplicidad estructural."
            )

        # ---------------------------------------------------------
        # 6. Refactor hacia patrón más declarativo
        # ---------------------------------------------------------
        if re.search(r'if\s*\(', removed) and not re.search(r'if\s*\(', added):
            insights.append(
                "Se reduce lógica condicional explícita, posible migración "
                "hacia patrón más declarativo o basado en configuración."
            )

        # ---------------------------------------------------------
        # 7. Mejora responsive
        # ---------------------------------------------------------
        if re.search(r'vh|vw|%', added) and not re.search(r'vh|vw|%', removed):
            insights.append(
                "Se introducen unidades relativas (vh/vw/%), indicando mejora en comportamiento responsive."
            )

        # ---------------------------------------------------------
        # 8. Reducción de acoplamiento
        # ---------------------------------------------------------
        if re.search(r'import ', removed) and not re.search(r'import ', added):
            insights.append(
                "Se eliminan dependencias importadas, posible reducción de acoplamiento."
            )

        # ---------------------------------------------------------
        # 9. Optimización de performance
        # ---------------------------------------------------------
        if re.search(r'\.subscribe\(', removed) and re.search(r'async|await|pipe', added):
            insights.append(
                "Se modifica patrón asincrónico, posible mejora en manejo de suscripciones o performance."
            )

        # ---------------------------------------------------------
        # 10. Eliminación de estado redundante
        # ---------------------------------------------------------
        if re.search(r'\bthis\.', removed) and not re.search(r'\bthis\.', added):
            insights.append(
                "Se reduce uso de propiedades del componente, indicando simplificación del estado."
            )

        return insights


class ChangeRelationAnalyzer:
    """Detecta relaciones funcionales entre archivos modificados del mismo diff."""

    _ENV_ENTRY = re.compile(r"^\s*([a-zA-Z_]\w*)\s*:\s*['\"]([^'\"]+)['\"]\s*,?\s*$")

    def __init__(self, changes: List[FileChange]):
        self.changes = changes
        self.by_path = {c.filepath: c for c in changes}

    def _component_family(self, fc: FileChange) -> List[str]:
        """Relaciona artefactos hermanos de un componente Angular."""
        name = fc.filename
        if '.component.' not in name:
            return []

        base = name.split('.component.')[0]
        folder = str(Path(fc.filepath).parent).replace('\\', '/')
        siblings: List[str] = []
        for other in self.changes:
            if other.filepath == fc.filepath:
                continue
            other_folder = str(Path(other.filepath).parent).replace('\\', '/')
            other_name = Path(other.filepath).name
            if other_folder == folder and other_name.startswith(base + '.component.'):
                siblings.append(other_name)
        return siblings

    def _extract_env_entries(self, fc: FileChange) -> List[Tuple[str, str, Optional[int]]]:
        entries: List[Tuple[str, str, Optional[int]]] = []
        seen: Set[Tuple[str, str]] = set()
        source = fc.added_with_line or [(None, l) for l in fc.added]
        for line_no, line in source:
            m = self._ENV_ENTRY.match(line.strip())
            if not m:
                continue
            key = m.group(1)
            val = m.group(2)
            if key.lower() in ('production', 'qa', 'local', 'staging'):
                continue
            pair = (key, val)
            if pair in seen:
                continue
            seen.add(pair)
            entries.append((key, val, line_no))
        return entries

    def _usage_files_for_env_key(self, key: str, value: str, current_path: str) -> Tuple[List[str], List[str]]:
        non_env_users: List[str] = []
        env_users: List[str] = []
        key_ref = re.compile(rf'\benvironment\.{re.escape(key)}\b')
        val_ref = re.compile(re.escape(value))

        for other in self.changes:
            if other.filepath == current_path:
                continue
            sample = "\n".join(other.added + other.removed + other.full_content[:200])
            if key_ref.search(sample) or val_ref.search(sample):
                if other.is_environment_file:
                    env_users.append(other.filename)
                else:
                    non_env_users.append(other.filename)
        return non_env_users[:6], env_users[:6]

    def _endpoint_purpose(self, key: str, value: str) -> str:
        text = f"{key} {value}".lower()
        if re.search(r'mantenimiento|catalogo', text):
            return "habilitar consumo de catálogos/mantenimientos"
        if re.search(r'auth|login|token', text):
            return "habilitar flujo de autenticación"
        if re.search(r'perfil|user|usuario', text):
            return "habilitar consulta/gestión de usuarios"
        if re.search(r'param|config', text):
            return "habilitar lectura de parámetros del sistema"
        return "habilitar consumo de API del módulo"

    def analyze(self, fc: FileChange) -> List[str]:
        insights: List[str] = []

        siblings = self._component_family(fc)
        if siblings:
            insights.append(
                "Interrelación de componente: se modifican artefactos relacionados "
                f"({', '.join(siblings)}), manteniendo coherencia entre lógica, vista y estilos."
            )

        if fc.is_environment_file:
            env_entries = self._extract_env_entries(fc)
            for key, value, _line_no in env_entries[:6]:
                purpose = self._endpoint_purpose(key, value)
                non_env_users, env_users = self._usage_files_for_env_key(key, value, fc.filepath)
                msg = (
                    f"Se agrega configuración de endpoint '{key}: {value}' para {purpose}."
                )
                if non_env_users:
                    msg += f" Consumo detectado en archivos funcionales: {', '.join(non_env_users)}."
                elif env_users:
                    msg += f" Homologado en ambientes: {', '.join(env_users)}."
                else:
                    msg += " No se detecta consumo directo en otros archivos modificados de este diff."
                insights.append(msg)

        # Relación por importaciones directas entre archivos cambiados
        import_lines = [l for l in fc.added if ' from ' in l and 'import ' in l]
        related: List[str] = []
        for line in import_lines:
            m = re.search(r"from\s+['\"]([^'\"]+)['\"]", line)
            if not m:
                continue
            path_hint = m.group(1).lower()
            for other in self.changes:
                if other.filepath == fc.filepath:
                    continue
                other_name = other.filename.lower().replace('.ts', '').replace('.html', '').replace('.scss', '')
                if other_name and other_name in path_hint and other.filename not in related:
                    related.append(other.filename)

        if related:
            insights.append(
                "Dependencias cruzadas detectadas con archivos también modificados: "
                f"{', '.join(related[:6])}."
            )

        return insights[:6]


# =============================================================================
# MOTOR DE ANÁLISIS ESLINT ANGULAR (Angular 16 · strict mode)
# Reglas activas según .eslintrc.json del proyecto:
#   .html → @angular-eslint/template/recommended  (reglas 16-18)
#   .ts   → @angular-eslint/recommended            (reglas 1-15)
# =============================================================================

# Eventos DOM nativos que no deben usarse como nombre de @Output()
DOM_NATIVE_EVENTS: frozenset = frozenset({
    'click', 'focus', 'blur', 'change', 'input', 'submit',
    'keyup', 'keydown', 'keypress', 'mouseenter', 'mouseleave',
    'mouseover', 'mouseout', 'mouseup', 'mousedown', 'dblclick',
    'contextmenu', 'scroll', 'resize', 'load', 'error', 'abort',
    'select', 'reset', 'drag', 'dragstart', 'dragend', 'dragenter',
    'dragleave', 'dragover', 'drop', 'touchstart', 'touchend',
    'touchmove', 'touchcancel', 'wheel', 'copy', 'cut', 'paste',
    'beforeinput',
})

# Hooks de ciclo de vida Angular y su interfaz correspondiente
LIFECYCLE_HOOKS: frozenset = frozenset({
    'ngOnInit', 'ngOnDestroy', 'ngOnChanges', 'ngDoCheck',
    'ngAfterContentInit', 'ngAfterContentChecked',
    'ngAfterViewInit', 'ngAfterViewChecked',
})

LIFECYCLE_INTERFACES: Dict[str, str] = {
    'ngOnInit':              'OnInit',
    'ngOnDestroy':           'OnDestroy',
    'ngOnChanges':           'OnChanges',
    'ngDoCheck':             'DoCheck',
    'ngAfterContentInit':    'AfterContentInit',
    'ngAfterContentChecked': 'AfterContentChecked',
    'ngAfterViewInit':       'AfterViewInit',
    'ngAfterViewChecked':    'AfterViewChecked',
}


@dataclass
class ESLintFinding:
    """Representa una corrección o violación ESLint detectada en una línea de diff."""
    rule:         str   # ej: "@angular-eslint/template/eqeqeq"
    severity:     str   # "ERROR"
    category:     str   # "CORRECCION" | "VIOLACION"
    line_added:   str   # línea nueva  (prefijo +)
    line_removed: str   # línea previa (prefijo -), puede ser ""
    description:  str   # descripción legible


@dataclass
class ESLintFileReport:
    """Resultado del análisis ESLint de un archivo."""
    corrections: List[ESLintFinding] = field(default_factory=list)
    violations:  List[ESLintFinding] = field(default_factory=list)
    functional:  List[str]           = field(default_factory=list)


class ESLintAngularAnalyzer:
    """
    Motor de análisis ESLint para proyectos Angular 16 strict mode.

    Archivos .html  → reglas 16, 17, 18  (template/recommended)
    Archivos .ts    → reglas  1-15       (@angular-eslint/recommended)

    Reglas NO activas (no se reportan): no-call-expression, no-any,
    no-conflicting-lifecycle, @typescript-eslint/**, accesibilidad.
    """

    # ── Patrones HTML ────────────────────────────────────────────────────────
    # Extrae contenido entre comillas de atributos Angular y de interpolaciones
    _RE_EXPR_ATTR   = re.compile(
        r'(?:[\[(*][\w.@$-]+\]?|\*\w+|#\w+)\s*=\s*"([^"]*)"'
    )
    _RE_EXPR_INTERP = re.compile(r'\{\{([^}]+)\}\}')

    # Regla 17: == o != que NO sean === ni !==
    # Lookbehind/lookahead para evitar capturar ===, !==, <=, >=
    _RE_EQEQ        = re.compile(r'(?<![=!<>])={2}(?!=)|(?<![!<>])!={1}(?!=)')
    _RE_STRICT      = re.compile(r'===|!==')

    # Regla 16: banana-in-box — patrón ([attr])=
    _RE_BANANA      = re.compile(r'\(\[[\w.]+\]\)\s*=')

    # Regla 18: no-negated-async — !(expr | async)
    _RE_NEG_ASYNC   = re.compile(r'!\s*\([^)]*\|\s*async\s*\)')

    # Regla 19: accessibility-table-scope — <th> debe tener scope="col"|"row"|"colgroup"|"rowgroup"
    # Incluida en template/recommended v16.3.1; el preset accessibility está comentado en .eslintrc.json
    # pero esta regla específica ya viene en recommended desde v15+.
    _RE_TH_SCOPE    = re.compile(r'<th\b[^>]*\bscope\s*=', re.I)
    _RE_TH_TAG      = re.compile(r'<th\b', re.I)

    # ── Patrones TS ──────────────────────────────────────────────────────────
    _RE_COMP_DECOR  = re.compile(r'@Component\s*\(')
    _RE_DIR_DECOR   = re.compile(r'@Directive\s*\(')
    _RE_PIPE_DECOR  = re.compile(r'@Pipe\s*\(')

    # Regla 4: método de ciclo de vida con cuerpo vacío
    _RE_EMPTY_HOOK  = re.compile(
        r'\b(' + '|'.join(LIFECYCLE_HOOKS) + r')\s*\([^)]*\)\s*:\s*\w*\s*\{\s*\}|\b('
        + '|'.join(LIFECYCLE_HOOKS) + r')\s*\([^)]*\)\s*\{\s*\}'
    )

    # Regla 5: host property en decorador
    _RE_HOST_PROP   = re.compile(r'\bhost\s*:\s*\{')

    # Regla 6: @Input con alias (cualquier argumento de cadena)
    _RE_INPUT_ALIAS  = re.compile(r"@Input\s*\(\s*['\"](\w+)['\"]\s*\)")
    # Regla 10: @Output con alias
    _RE_OUTPUT_ALIAS = re.compile(r"@Output\s*\(\s*['\"](\w+)['\"]\s*\)")

    # Reglas 7/11: inputs/outputs en metadata del decorador
    _RE_INPUTS_META  = re.compile(r'\binputs\s*:\s*\[')
    _RE_OUTPUTS_META = re.compile(r'\boutputs\s*:\s*\[')

    # Regla 8: nombre de la variable después de @Output()
    _RE_OUTPUT_DECL  = re.compile(
        r'@Output\s*\(\s*\)\s+(?:public\s+|private\s+|protected\s+|readonly\s+)?(\w+)'
    )

    # Regla 9: @Output cuyo nombre empieza por "on"
    _RE_OUTPUT_ON    = re.compile(
        r'@Output\s*\([^)]*\)\s+(?:public\s+|private\s+|protected\s+|readonly\s+)?on[A-Z]\w*'
    )

    # Regla 12: método de ciclo de vida declarado en la clase
    _RE_HOOK_METHOD  = re.compile(r'\b(' + '|'.join(LIFECYCLE_HOOKS) + r')\s*\(')
    _RE_IMPLEMENTS   = re.compile(r'\bimplements\b([^{]+)')

    # Regla 13: @Pipe sin PipeTransform
    _RE_PIPE_TRANSFORM = re.compile(r'\bimplements\b[^{]*\bPipeTransform\b')

    # Reglas 14/15: selector
    _RE_SELECTOR     = re.compile(r"\bselector\s*:\s*['\"]([^'\"]+)['\"]")

    # Sufijos obligatorios (reglas 1 y 3)
    _RE_CLASS_DEF    = re.compile(r'\bclass\s+(\w+)')

    # ── Helpers internos ─────────────────────────────────────────────────────

    def _has_loose_equality(self, text: str) -> bool:
        """Devuelve True si 'text' contiene == o != no estrictos."""
        masked = self._RE_STRICT.sub('\x00\x00\x00', text)
        return bool(self._RE_EQEQ.search(masked))

    def _extract_angular_expressions(self, line: str) -> List[str]:
        """Extrae el contenido de atributos Angular e interpolaciones."""
        exprs: List[str] = []
        for m in self._RE_EXPR_ATTR.finditer(line):
            exprs.append(m.group(1))
        for m in self._RE_EXPR_INTERP.finditer(line):
            exprs.append(m.group(1))
        return exprs

    def _line_has_eqeqeq(self, line: str) -> bool:
        exprs = self._extract_angular_expressions(line)
        return any(self._has_loose_equality(e) for e in exprs)

    def _best_removed_match(self, added_l: str, removed_lines: List[str]) -> str:
        """Encuentra la línea eliminada más similar a la añadida (ratio > 0.5)."""
        if not removed_lines:
            return ""
        a_ns = re.sub(r'\s+', '', added_l)
        best, best_score = "", 0.0
        for r_l in removed_lines:
            r_ns = re.sub(r'\s+', '', r_l)
            common = sum(1 for c in set(a_ns) if c in r_ns)
            total  = max(len(a_ns), len(r_ns), 1)
            score  = common / total
            if score > best_score:
                best_score, best = score, r_l
        return best if best_score > 0.50 else ""

    # ── Análisis HTML ────────────────────────────────────────────────────────

    def analyze_html_line(self, added: str, removed: str) -> Optional[ESLintFinding]:
        a, r = added.strip(), removed.strip()

        # Regla 17 – eqeqeq
        a_has = self._line_has_eqeqeq(a)
        r_has = self._line_has_eqeqeq(r) if r else False
        if r_has and not a_has:
            return ESLintFinding(
                rule="@angular-eslint/template/eqeqeq", severity="ERROR",
                category="CORRECCION", line_added=a, line_removed=r,
                description="Corrección de igualdad estricta (== → === o != → !==) en expresión de plantilla.",
            )
        if a_has and not r_has:
            return ESLintFinding(
                rule="@angular-eslint/template/eqeqeq", severity="ERROR",
                category="VIOLACION", line_added=a, line_removed=r,
                description="Uso de == o != en expresión de plantilla Angular. Debe usarse === o !==.",
            )

        # Regla 16 – banana-in-box
        a_ban = bool(self._RE_BANANA.search(a))
        r_ban = bool(self._RE_BANANA.search(r)) if r else False
        if a_ban and not r_ban:
            return ESLintFinding(
                rule="@angular-eslint/template/banana-in-box", severity="ERROR",
                category="VIOLACION", line_added=a, line_removed=r,
                description="Two-way binding incorrecto: se usa ([attr])= en lugar de [(attr)]=.",
            )
        if r_ban and not a_ban:
            return ESLintFinding(
                rule="@angular-eslint/template/banana-in-box", severity="ERROR",
                category="CORRECCION", line_added=a, line_removed=r,
                description="Corrección de two-way binding: sintaxis ([]) reemplazada por [()].",
            )

        # Regla 18 – no-negated-async
        a_neg = bool(self._RE_NEG_ASYNC.search(a))
        r_neg = bool(self._RE_NEG_ASYNC.search(r)) if r else False
        if a_neg and not r_neg:
            return ESLintFinding(
                rule="@angular-eslint/template/no-negated-async", severity="ERROR",
                category="VIOLACION", line_added=a, line_removed=r,
                description="Negación directa sobre pipe async ( !(obs$ | async) ). Usar variable con 'as'.",
            )
        if r_neg and not a_neg:
            return ESLintFinding(
                rule="@angular-eslint/template/no-negated-async", severity="ERROR",
                category="CORRECCION", line_added=a, line_removed=r,
                description="Corrección: se elimina el patrón !(... | async) del template.",
            )

        # Regla 19 – accessibility-table-scope
        # Activa: viene en template/recommended v16.3.1 (el preset accessibility
        # está comentado en .eslintrc.json pero esta regla ya pertenece a recommended).
        a_th = bool(self._RE_TH_TAG.search(a))
        r_th = bool(self._RE_TH_TAG.search(r)) if r else False
        a_sc = bool(self._RE_TH_SCOPE.search(a)) if a_th else False
        r_sc = bool(self._RE_TH_SCOPE.search(r)) if r_th else False

        if a_th and a_sc and r_th and not r_sc:
            return ESLintFinding(
                rule="@angular-eslint/template/accessibility-table-scope", severity="ERROR",
                category="CORRECCION", line_added=a, line_removed=r,
                description=(
                    "Corrección: se añade scope=\"col\"/\"row\" al elemento <th> "
                    "para cumplir accesibilidad de tablas HTML."
                ),
            )
        if a_th and not a_sc and not (r_th and r_sc):
            return ESLintFinding(
                rule="@angular-eslint/template/accessibility-table-scope", severity="ERROR",
                category="VIOLACION", line_added=a, line_removed=r,
                description="Elemento <th> sin atributo scope. Debe añadirse scope=\"col\" o scope=\"row\".",
            )

        return None

    # ── Análisis TypeScript ──────────────────────────────────────────────────

    def analyze_ts_line(
        self,
        added: str,
        removed: str,
        full_content: List[str],
    ) -> Optional[ESLintFinding]:
        a, r = added.strip(), removed.strip()
        full_text = "\n".join(full_content)

        # Regla 4 – no-empty-lifecycle-method
        if self._RE_EMPTY_HOOK.search(a) and not self._RE_EMPTY_HOOK.search(r):
            return ESLintFinding(
                rule="@angular-eslint/no-empty-lifecycle-method", severity="ERROR",
                category="VIOLACION", line_added=a, line_removed=r,
                description="Método de ciclo de vida declarado con cuerpo vacío {}.",
            )
        if self._RE_EMPTY_HOOK.search(r) and not self._RE_EMPTY_HOOK.search(a):
            return ESLintFinding(
                rule="@angular-eslint/no-empty-lifecycle-method", severity="ERROR",
                category="CORRECCION", line_added=a, line_removed=r,
                description="Corrección: se elimina o implementa el método de ciclo de vida vacío.",
            )

        # Regla 5 – no-host-metadata-property
        if self._RE_HOST_PROP.search(a) and not self._RE_HOST_PROP.search(r):
            return ESLintFinding(
                rule="@angular-eslint/no-host-metadata-property", severity="ERROR",
                category="VIOLACION", line_added=a, line_removed=r,
                description="Propiedad 'host:' en decorador. Usar @HostListener / @HostBinding.",
            )
        if self._RE_HOST_PROP.search(r) and not self._RE_HOST_PROP.search(a):
            return ESLintFinding(
                rule="@angular-eslint/no-host-metadata-property", severity="ERROR",
                category="CORRECCION", line_added=a, line_removed=r,
                description="Corrección: se elimina propiedad 'host:' del decorador.",
            )

        # Regla 6 – no-input-rename
        m_in_a = self._RE_INPUT_ALIAS.search(a)
        m_in_r = self._RE_INPUT_ALIAS.search(r) if r else None
        if m_in_a and not m_in_r:
            return ESLintFinding(
                rule="@angular-eslint/no-input-rename", severity="ERROR",
                category="VIOLACION", line_added=a, line_removed=r,
                description=f"@Input() con alias '{m_in_a.group(1)}'. El alias debe coincidir con el nombre de la propiedad.",
            )
        if m_in_r and not m_in_a:
            return ESLintFinding(
                rule="@angular-eslint/no-input-rename", severity="ERROR",
                category="CORRECCION", line_added=a, line_removed=r,
                description=f"Corrección: se elimina alias '{m_in_r.group(1)}' de @Input().",
            )

        # Regla 7 – no-inputs-metadata-property
        if self._RE_INPUTS_META.search(a) and not self._RE_INPUTS_META.search(r):
            return ESLintFinding(
                rule="@angular-eslint/no-inputs-metadata-property", severity="ERROR",
                category="VIOLACION", line_added=a, line_removed=r,
                description="Propiedad 'inputs:' en decorador. Usar @Input() en cada campo.",
            )
        if self._RE_INPUTS_META.search(r) and not self._RE_INPUTS_META.search(a):
            return ESLintFinding(
                rule="@angular-eslint/no-inputs-metadata-property", severity="ERROR",
                category="CORRECCION", line_added=a, line_removed=r,
                description="Corrección: se elimina propiedad 'inputs:' del decorador.",
            )

        # Regla 8 – no-output-native
        m_out_decl = self._RE_OUTPUT_DECL.search(a)
        if m_out_decl:
            out_name = m_out_decl.group(1).strip()
            if out_name.lower() in DOM_NATIVE_EVENTS:
                return ESLintFinding(
                    rule="@angular-eslint/no-output-native", severity="ERROR",
                    category="VIOLACION", line_added=a, line_removed=r,
                    description=f"@Output() '{out_name}' colisiona con evento DOM nativo.",
                )

        # Regla 9 – no-output-on-prefix
        if self._RE_OUTPUT_ON.search(a) and not self._RE_OUTPUT_ON.search(r):
            return ESLintFinding(
                rule="@angular-eslint/no-output-on-prefix", severity="ERROR",
                category="VIOLACION", line_added=a, line_removed=r,
                description="@Output() cuyo nombre empieza por 'on'. Renombrar sin el prefijo.",
            )
        if self._RE_OUTPUT_ON.search(r) and not self._RE_OUTPUT_ON.search(a):
            return ESLintFinding(
                rule="@angular-eslint/no-output-on-prefix", severity="ERROR",
                category="CORRECCION", line_added=a, line_removed=r,
                description="Corrección: se elimina prefijo 'on' del @Output().",
            )

        # Regla 10 – no-output-rename
        m_out_a = self._RE_OUTPUT_ALIAS.search(a)
        m_out_r = self._RE_OUTPUT_ALIAS.search(r) if r else None
        if m_out_a and not m_out_r:
            return ESLintFinding(
                rule="@angular-eslint/no-output-rename", severity="ERROR",
                category="VIOLACION", line_added=a, line_removed=r,
                description=f"@Output() con alias '{m_out_a.group(1)}'. El alias debe coincidir con el nombre de la propiedad.",
            )
        if m_out_r and not m_out_a:
            return ESLintFinding(
                rule="@angular-eslint/no-output-rename", severity="ERROR",
                category="CORRECCION", line_added=a, line_removed=r,
                description=f"Corrección: se elimina alias '{m_out_r.group(1)}' de @Output().",
            )

        # Regla 11 – no-outputs-metadata-property
        if self._RE_OUTPUTS_META.search(a) and not self._RE_OUTPUTS_META.search(r):
            return ESLintFinding(
                rule="@angular-eslint/no-outputs-metadata-property", severity="ERROR",
                category="VIOLACION", line_added=a, line_removed=r,
                description="Propiedad 'outputs:' en decorador. Usar @Output() en cada campo.",
            )
        if self._RE_OUTPUTS_META.search(r) and not self._RE_OUTPUTS_META.search(a):
            return ESLintFinding(
                rule="@angular-eslint/no-outputs-metadata-property", severity="ERROR",
                category="CORRECCION", line_added=a, line_removed=r,
                description="Corrección: se elimina propiedad 'outputs:' del decorador.",
            )

        # Regla 12 – use-lifecycle-interface
        m_hook = self._RE_HOOK_METHOD.search(a)
        if m_hook:
            hook_name = m_hook.group(1)
            if hook_name not in r:
                iface = LIFECYCLE_INTERFACES.get(hook_name, "")
                if iface:
                    impl_match = self._RE_IMPLEMENTS.search(full_text)
                    if not impl_match or iface not in impl_match.group(1):
                        return ESLintFinding(
                            rule="@angular-eslint/use-lifecycle-interface", severity="ERROR",
                            category="VIOLACION", line_added=a, line_removed=r,
                            description=f"Método '{hook_name}' sin declarar 'implements {iface}' en la clase.",
                        )

        # Regla 13 – use-pipe-transform-interface
        if self._RE_PIPE_DECOR.search(a):
            if not self._RE_PIPE_TRANSFORM.search(full_text):
                return ESLintFinding(
                    rule="@angular-eslint/use-pipe-transform-interface", severity="ERROR",
                    category="VIOLACION", line_added=a, line_removed=r,
                    description="Clase con @Pipe sin declarar 'implements PipeTransform'.",
                )

        # Reglas 1/3 – component-class-suffix / directive-class-suffix
        m_cls = self._RE_CLASS_DEF.search(a)
        if m_cls:
            cls_name = m_cls.group(1)
            is_comp = self._RE_COMP_DECOR.search(full_text)
            is_dir  = self._RE_DIR_DECOR.search(full_text)
            if is_comp and not cls_name.endswith('Component'):
                return ESLintFinding(
                    rule="@angular-eslint/component-class-suffix", severity="ERROR",
                    category="VIOLACION", line_added=a, line_removed=r,
                    description=f"La clase '{cls_name}' debe terminar en 'Component'.",
                )
            if is_dir and not cls_name.endswith('Directive'):
                return ESLintFinding(
                    rule="@angular-eslint/directive-class-suffix", severity="ERROR",
                    category="VIOLACION", line_added=a, line_removed=r,
                    description=f"La clase '{cls_name}' debe terminar en 'Directive'.",
                )

        # Reglas 14/15 – component-selector / directive-selector
        m_sel = self._RE_SELECTOR.search(a)
        if m_sel:
            sel = m_sel.group(1).strip()
            is_comp = bool(self._RE_COMP_DECOR.search(full_text))
            is_dir  = bool(self._RE_DIR_DECOR.search(full_text))
            if is_comp and not re.match(r'^app-[a-z][a-z0-9-]*$', sel):
                return ESLintFinding(
                    rule="@angular-eslint/component-selector", severity="ERROR",
                    category="VIOLACION", line_added=a, line_removed=r,
                    description=f"Selector de componente '{sel}' no cumple: prefijo 'app-', kebab-case, tipo elemento.",
                )
            if is_dir and not re.match(r'^app[A-Z][a-zA-Z0-9]*$', sel):
                return ESLintFinding(
                    rule="@angular-eslint/directive-selector", severity="ERROR",
                    category="VIOLACION", line_added=a, line_removed=r,
                    description=f"Selector de directiva '{sel}' no cumple: prefijo 'app', camelCase, tipo atributo.",
                )

        return None

    # ── Análisis de archivo completo ─────────────────────────────────────────

    def analyze_file(self, fc: FileChange) -> ESLintFileReport:
        """
        Analiza todas las líneas añadidas de un FileChange.
        Empareja cada línea + con la línea - más similar del mismo hunk
        para distinguir correcciones de violaciones nuevas.
        """
        report = ESLintFileReport()
        ext = fc.ext.lower()
        is_html = ext in ('.html', '.component.html')
        is_ts   = ext in (
            '.ts', '.component.ts', '.service.ts', '.pipe.ts',
            '.directive.ts', '.guard.ts', '.interceptor.ts',
            '.module.ts', '.resolver.ts',
        )

        if not (is_html or is_ts):
            report.functional = [l for l in fc.added]
            return report

        removed_pool = list(fc.removed)
        used_findings: Set[Tuple[str, str, str]] = set()

        for added_l in fc.added:
            a_stripped = added_l.strip()
            paired_r   = self._best_removed_match(added_l, removed_pool)

            finding: Optional[ESLintFinding] = None
            if is_html:
                finding = self.analyze_html_line(added_l, paired_r)
            elif is_ts:
                finding = self.analyze_ts_line(added_l, paired_r, fc.full_content)

            if finding:
                key = (finding.rule, finding.category, a_stripped[:120])
                if key not in used_findings:
                    used_findings.add(key)
                    if finding.category == 'CORRECCION':
                        report.corrections.append(finding)
                    else:
                        report.violations.append(finding)
            else:
                report.functional.append(a_stripped)

        return report


def analyze_technical_impact(fc: FileChange) -> str:
    """
    FIX BUG 5: Archivos de entorno analizan full_content para detectar
    patrones criticos que estan en lineas de contexto (sin + ni -).
    """
    if fc.is_binary:
        return "Actualizacion de archivo binario (imagen, fuente, recurso)"
    if fc.is_lockfile:
        n = len(fc.added)
        d = len(fc.removed)
        return f"Actualizacion de dependencias: +{n} entradas nuevas, -{d} eliminadas"

    search_lines = fc.added + fc.removed
    if not search_lines:
        search_lines = fc.full_content[:40]
    search_text = "\n".join(search_lines)

    found = []
    for pattern, description in IMPACT_SIGNALS:
        if pattern.search(search_text):
            found.append(description)

    if found:
        return " | ".join(found[:4])

    category = get_category(fc)
    if len(fc.added) == 0 and len(fc.removed) == 0:
        return f"Ajuste de propiedades o permisos en {category}"
    if fc.kind == 'added':
        return f"Implementacion inicial de {category}"
    if fc.kind == 'deleted':
        return f"Eliminacion completa de {category}"
    return f"Modificacion de logica interna en {category}"


def calculate_deploy_impact(fc: FileChange) -> str:
    """
    FIX BUG 4: Regex de AuthService mas preciso (sin falsos positivos en rutas).
    FIX BUG 5: Archivos de entorno analizan full_content para score correcto.
    """
    if fc.is_binary:
        return "Leve"
    if fc.is_lockfile:
        return "Leve"

    changed_lines = "\n".join(fc.added + fc.removed)
    all_lines = changed_lines if changed_lines.strip() else "\n".join(fc.full_content[:60])

    n_add = len(fc.added)
    n_del = len(fc.removed)
    score = 0

    # --- CRITICO ---
    if re.search(r'\bAuthService\b|\bAuthGuard\b|jsonwebtoken|\.verify\s*\(.*token', all_lines):
        score += 60
    if re.search(r'(?<!\w)password(?!\w)|\bbcrypt\b|\bsalt\b', all_lines, re.I):
        score += 50
    if re.search(r'migration|alembic|flyway|DROP TABLE|ALTER TABLE', all_lines, re.I):
        score += 70

    # --- ALTO: entornos con URLs reales ---
    if fc.is_environment_file and re.search(r'apiLoginUrl|apiBackend|apiUrl\b', all_lines, re.I):
        score += 45

    # --- BAJO/MEDIO: nuevos endpoints o rutas de recursos en env ---
    if fc.is_environment_file:
        env_added = "\n".join(fc.added)
        env_entry_pattern = re.compile(r"\b[a-zA-Z_]\w*\s*:\s*['\"][^'\"]+['\"]")
        endpoint_hint = re.compile(r'api|url|endpoint|catalogo|mantenimiento|parametro|service', re.I)
        env_entries = env_entry_pattern.findall(env_added)
        if env_entries:
            score += 10
            if any(endpoint_hint.search(e) for e in env_entries):
                score += 8

    # --- ALTO: contratos publicos ---
    if re.search(r'\bexport\s+(interface|type|class|enum)\b', all_lines):
        score += 40
    if re.search(r'@Input\s*\(\)|@Output\s*\(\)|EventEmitter', all_lines):
        score += 35
    if re.search(r'@NgModule\s*\(|providers\s*:\s*\[|declarations\s*:\s*\[', all_lines):
        score += 35
    if re.search(r'dispatch\s*\(|createAction|createReducer|createEffect', all_lines):
        score += 30

    # --- MEDIO ---
    if re.search(r'HttpClient|http\.(get|post|put|delete|patch)', all_lines, re.I):
        score += 30
    if re.search(r'catchError|throwError|try\s*{|except\s+', all_lines):
        score += 20
    if re.search(r'SELECT\b|INSERT\b|UPDATE\b|DELETE\b|\.save\s*\(', all_lines, re.I):
        score += 25
    if re.search(r'router\.navigate|canActivate|canLoad|authGuard|menuGuard', all_lines):
        score += 20
    if re.search(r'cron|scheduler|setTimeout|setInterval', all_lines):
        score += 20
    if re.search(r'WebSocket|socket\.io', all_lines, re.I):
        score += 25

    # --- BAJO: UI y estilos ---
    if fc.ext in ('.scss', '.css', '.component.scss'):
        score += 5
    if fc.ext in ('.html', '.component.html'):
        score += 10

    # --- LEVE: tests, docs ---
    if fc.ext == '.spec.ts':
        score = min(score + 3, 15)
    if fc.ext == '.md':
        return "Nula"

    # --- PHP / Laravel: scoring por contexto del archivo ---
    if fc.ext == '.php':
        filepath_lower = fc.filepath.lower().replace('\\', '/')
        score += 20  # base: cualquier archivo PHP en produccion tiene impacto

        # Controlador de API: modifica contratos de respuesta
        if re.search(r'controller', filepath_lower):
            score += 15
        # Middleware: intercepta todas las peticiones
        if re.search(r'middleware', filepath_lower):
            score += 35
        # Migracion: altera el esquema de base de datos
        if re.search(r'migration[s]?/', filepath_lower):
            score += 55
        # Rutas: define o modifica endpoints publicos
        if re.search(r'routes?/(api|web|channels|console)\.php', filepath_lower):
            score += 30
        # Configuracion de framework
        if re.search(r'config/', filepath_lower):
            score += 25
        # Seeder / Factory: populate de datos, riesgo bajo
        if re.search(r'seeder|factory', filepath_lower):
            score = min(score, 25)
        # Vista Blade: solo UI, impacto reducido
        if '.blade.' in filepath_lower:
            score = min(score, 20)

        # Bonus por signales en el codigo cambiado
        if re.search(r'->where\s*\(|->select\s*\(|->save\s*\(|->create\s*\(|->update\s*\(|->delete\s*\(', all_lines, re.I):
            score += 20  # operaciones ORM directas
        if re.search(r'return\s+response\s*\(|->json\s*\(|\$this->success\b|\$this->error\b', all_lines, re.I):
            score += 15  # modifica respuesta de la API
        if re.search(r'\?\?|\?->|\?\?=', all_lines):
            score += 10  # null guards (correccion de bug potencial)
        if re.search(r'\blog::|try\s*\{|catch\s*\(\$e\b|catch\s*\(Exception', all_lines, re.I):
            score += 10  # manejo de excepciones
        if re.search(r'auth\s*\(|can\s*\(|Gate::|\bpolicy\b', all_lines, re.I):
            score += 20  # logica de autorizacion
        if re.search(r'->belongsTo\s*\(|->hasMany\s*\(|->hasOne\s*\(|->belongsToMany\s*\(', all_lines):
            score += 12  # relaciones entre modelos
        if re.search(r'->paginate\s*\(|->get\s*\(|->first\s*\(|->with\s*\(', all_lines, re.I):
            score += 8   # consulta de coleccion

    # Bonus por magnitud
    total_lines = n_add + n_del
    if total_lines > 200:   score += 20
    elif total_lines > 100: score += 12
    elif total_lines > 50:  score += 6
    elif total_lines > 20:  score += 3
    elif total_lines > 5:   score += 1  # cambios quirurgicos siguen puntuando

    # Archivo nuevo: cap de impacto
    if fc.kind == 'added' and score < 40:
        score = min(score, 30)
    # Archivo eliminado: riesgo de dependencias rotas
    if fc.kind == 'deleted':
        score += 20

    # Solo linter fixes
    if fc.removed and not fc.added:
        if all(fc.detect_linter_fix(l) for l in fc.removed if l.strip()):
            return "Nula"

    # Cambios muy pequenos en environment no deben inflarse por contexto completo
    critical_change = bool(re.search(
        r'\bAuthService\b|\bAuthGuard\b|jsonwebtoken|\.verify\s*\(.*token|(?<!\w)password(?!\w)|\bbcrypt\b|\bsalt\b|migration|alembic|flyway|DROP TABLE|ALTER TABLE',
        changed_lines,
        re.I,
    ))
    total_lines = n_add + n_del
    if fc.is_environment_file and total_lines <= 5 and not critical_change:
        score = min(score, 28)

    if score == 0:        return "Nula"
    elif score <= 10:     return "Leve"
    elif score <= 28:     return "Baja"
    elif score <= 50:     return "Media"
    elif score <= 75:     return "Alta"
    else:                 return "Critica"


# =============================================================================
# RESUMEN DE LOCK FILE
# =============================================================================
def summarize_lockfile(fc: FileChange) -> Dict[str, List[str]]:
    """Extrae nombres de paquetes nuevos/eliminados de un lock file."""
    added_pkgs:   List[str] = []
    removed_pkgs: List[str] = []
    pkg_pattern = re.compile(r'"(@?[\w/@.-]+)":\s*\{|^(@?[\w/@.-]+)@[\d.^~]')

    for line in fc.added:
        m = pkg_pattern.search(line.strip())
        if m:
            name = (m.group(1) or m.group(2) or '').strip('"')
            if name and name not in added_pkgs:
                added_pkgs.append(name)

    for line in fc.removed:
        m = pkg_pattern.search(line.strip())
        if m:
            name = (m.group(1) or m.group(2) or '').strip('"')
            if name and name not in removed_pkgs:
                removed_pkgs.append(name)

    return {"added": added_pkgs[:20], "removed": removed_pkgs[:20]}


# =============================================================================
# RECOMENDACIONES
# =============================================================================
def analyze_recommendations(changes: List[FileChange]) -> List[str]:
    recs = []
    all_removed = "\n".join(l for f in changes for l in f.removed).lower()

    if 'console.' in all_removed or 'print(' in all_removed:
        recs.append(
            "Se limpiaron logs de depuracion. Verificar que no falten "
            "trazas criticas en produccion."
        )

    env_changes = [f for f in changes if f.is_environment_file]
    if env_changes:
        ambientes = [f.filename for f in env_changes]
        recs.append(
            f"Se modificaron {len(env_changes)} archivo(s) de entorno "
            f"({', '.join(ambientes)}). Confirmar que los endpoints sean "
            f"correctos en cada ambiente antes del despliegue."
        )

    if any(f.is_lockfile for f in changes):
        recs.append(
            "Se actualizo el archivo de dependencias. Ejecutar 'npm install' "
            "en el servidor de CI/CD para sincronizar node_modules."
        )

    if any(f.kind == 'deleted' for f in changes):
        recs.append(
            "Se eliminaron archivos. Ejecutar build completo y verificar "
            "que no queden imports huerfanos."
        )

    if any(f.ext in ('.html', '.scss', '.css', '.component.html', '.component.scss')
           for f in changes):
        recs.append(
            "Cambios en interfaz de usuario. Realizar QA visual en "
            "distintas resoluciones (movil y escritorio)."
        )

    all_code = "\n".join(l for f in changes for l in f.added + f.removed)
    if re.search(r'\bexport\s+(interface|type)\b', all_code):
        recs.append(
            "Se modificaron interfaces o tipos exportados. Verificar que "
            "todos los consumidores del tipo esten actualizados y que el "
            "build compile sin errores de tipo."
        )

    if any(f.ext == '.spec.ts' for f in changes):
        recs.append(
            "Se modificaron tests unitarios. Ejecutar suite completa de "
            "pruebas antes del merge."
        )

    if re.search(r'migration|ALTER TABLE|DROP', all_code, re.I):
        recs.append(
            "ALERTA: Detectadas posibles migraciones de base de datos. "
            "Revisar con DBA antes del despliegue."
        )

    if re.search(r'@NgModule|providers\s*:|declarations\s*:', all_code):
        recs.append(
            "Se modifico un modulo Angular. Verificar que las declaraciones "
            "y providers sean correctas y que no haya duplicados."
        )

    if any(calculate_deploy_impact(f) in ("Alta", "Critica") for f in changes):
        recs.append(
            "Existen cambios de impacto ALTO o CRITICO. Se recomienda "
            "revision por pares (code review) antes del merge a develop."
        )

    if not recs:
        recs.append(
            "No se detectaron riesgos criticos. Revisar cobertura de pruebas "
            "unitarias para las nuevas estructuras anadidas."
        )

    return list(dict.fromkeys(recs))


# =============================================================================
# HELPERS DOCX
# =============================================================================
def _set_bg(cell, hex_color: str):
    tcPr = cell._tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def _set_borders(cell, color: str = "CCCCCC"):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        borders.append(el)
    tcPr.append(borders)

def _set_width(cell, cm: float):
    tcPr = cell._tc.get_or_add_tcPr()
    tcW  = OxmlElement("w:tcW")
    tcW.set(qn("w:w"), str(int(cm * 567)))
    tcW.set(qn("w:type"), "dxa")
    tcPr.append(tcW)

def _run(para, text: str, bold=False, italic=False,
         color: RGBColor = None, size: float = 10, font: str = "Arial"):
    r = para.add_run(text)
    r.bold        = bold
    r.italic      = italic
    r.font.size   = Pt(size)
    r.font.name   = font
    if color:
        r.font.color.rgb = color
    return r

def _header_row(table, cols: List[Tuple[str, float]]):
    row = table.rows[0]
    for i, (text, w) in enumerate(cols):
        cell = row.cells[i]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        _set_bg(cell, C_HDR_BG)
        _set_borders(cell, C_ACCENT)
        _set_width(cell, w)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(p, text, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF), size=11)

def _data_row(table, cells: List[Tuple[str, float, RGBColor, bool, str]]):
    row = table.add_row()
    for i, (text, w, tc, bold, bg) in enumerate(cells):
        cell = row.cells[i]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        _set_bg(cell, bg)
        _set_borders(cell)
        _set_width(cell, w)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for j, line in enumerate(text.split("\n")):
            if j > 0:
                p = cell.add_paragraph()
            _run(p, line, bold=bold, color=tc, size=10)

def _bullet(doc, text: str, symbol: str, color: RGBColor, indent: float = 1.0):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent       = Cm(indent)
    p.paragraph_format.first_line_indent = Cm(-0.5)
    p.paragraph_format.space_after       = Pt(2)
    _run(p, f"{symbol}  ", bold=True, color=color, size=9.5)
    _run(p, text, color=color, size=9.5)

def _section_title(doc, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(5)
    _run(p, text, bold=True, color=C_TITLE, size=12)
    pPr    = p._p.get_or_add_pPr()
    pBdr   = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), C_ACCENT)
    pBdr.append(bottom)
    pPr.append(pBdr)

def _divider(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "4")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), C_BORDER)
    pBdr.append(bot)
    pPr.append(pBdr)

def _h1(doc, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(7)
    _run(p, text, bold=True, color=C_TITLE, size=14)


# =============================================================================
# GENERADOR DOCX
# =============================================================================
class ReportGenerator:
    def __init__(self, changes: List[FileChange], branch_from: str, branch_to: str, output: str):
        self.changes         = changes
        self.branch_from     = branch_from
        self.branch_to       = branch_to
        self.output          = output
        self.doc             = Document()
        self._eslint_reports: Dict[str, ESLintFileReport] = {}
        self._upgrade_compatibility()
        self._page_setup()

    def _upgrade_compatibility(self):
        settings = self.doc.settings.element
        compat   = OxmlElement('w:compat')
        cs       = OxmlElement('w:compatSetting')
        cs.set(qn('w:name'), 'compatibilityMode')
        cs.set(qn('w:uri'),  'http://schemas.microsoft.com/office/word')
        cs.set(qn('w:val'),  '15')
        compat.append(cs)
        settings.append(compat)

    def _page_setup(self):
        sec = self.doc.sections[0]
        sec.page_width    = Inches(8.5)
        sec.page_height   = Inches(11)
        sec.left_margin   = Cm(2.0)
        sec.right_margin  = Cm(2.0)
        sec.top_margin    = Cm(2.0)
        sec.bottom_margin = Cm(2.0)
        style = self.doc.styles["Normal"]
        style.font.name = "Arial"
        style.font.size = Pt(10)

    def _title(self):
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(3)
        _run(p, "INFORME DE CAMBIOS DE CODIGO", bold=True, color=C_TITLE, size=20)

        p2 = self.doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_after = Pt(2)
        _run(p2, f"Rama: {self.branch_from}  hacia  {self.branch_to}", color=C_SUBTITLE, size=11)

        p3 = self.doc.add_paragraph()
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p3.paragraph_format.space_after = Pt(12)
        _run(p3, f"Fecha: {fecha_espanol()}", color=C_MUTED, size=10)

        p4 = self.doc.add_paragraph()
        p4.paragraph_format.space_after = Pt(12)
        pPr  = p4._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bot  = OxmlElement("w:bottom")
        bot.set(qn("w:val"),   "single")
        bot.set(qn("w:sz"),    "12")
        bot.set(qn("w:space"), "1")
        bot.set(qn("w:color"), C_ACCENT)
        pBdr.append(bot)
        pPr.append(pBdr)

    def _summary(self):
        _h1(self.doc, "1. Tabla Resumen de Cambios")
        total_add = sum(len(f.added) for f in self.changes)
        total_del = sum(len(f.removed) for f in self.changes)

        p = self.doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        _run(
            p,
            f"Archivos afectados: {len(self.changes)}   "
            f"Lineas anadidas: +{total_add}   "
            f"Lineas eliminadas: -{total_del}",
            color=C_MUTED, size=9, italic=True
        )

        COLS = [
            ("Archivo",           3.8),
            ("Categoria",         2.6),
            ("Tipo de Cambio",    2.2),
            ("Impacto Tecnico",   4.0),
            ("Impacto en Deploy", 1.6),
        ]
        tbl = self.doc.add_table(rows=1, cols=len(COLS))
        tbl.style = "Table Grid"
        tbl.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _header_row(tbl, COLS)

        for fc in self.changes:
            tc_kind, bg_kind = fc.kind_colors
            category         = get_category(fc)
            deploy_level     = calculate_deploy_impact(fc)
            tc_dep, bg_dep   = impact_colors(deploy_level)
            tech_impact      = analyze_technical_impact(fc)

            _data_row(tbl, [
                (fc.filename,   3.8, C_MOD_TEXT, True,  "EFF6FF"),
                (category,      2.6, C_BODY,     False, C_ROW_ALT),
                (fc.kind_label, 2.2, tc_kind,    True,  bg_kind),
                (tech_impact,   4.0, C_BODY,     False, C_WHITE),
                (deploy_level,  1.6, tc_dep,     True,  bg_dep),
            ])

        self.doc.add_paragraph()

    def _render_lockfile_detail(self, fc: FileChange):
        """Renderizado compacto para archivos de dependencias (lock files)."""
        summary = summarize_lockfile(fc)
        p_info  = self.doc.add_paragraph()
        _run(
            p_info,
            f"Actualizacion de dependencias del proyecto. "
            f"Cambios: +{len(fc.added)} / -{len(fc.removed)} entradas. "
            f"Archivo gestionado automaticamente por el package manager.",
            color=C_MUTED, size=9, italic=True
        )
        if summary["added"]:
            p_a = self.doc.add_paragraph()
            _run(p_a, "Paquetes con entradas nuevas:", bold=True, color=C_ADD_TEXT, size=9)
            for pkg in summary["added"]:
                _bullet(self.doc, pkg, "+", C_ADD_TEXT)
        if summary["removed"]:
            p_r = self.doc.add_paragraph()
            _run(p_r, "Paquetes con entradas eliminadas:", bold=True, color=C_DEL_TEXT, size=9)
            for pkg in summary["removed"]:
                _bullet(self.doc, pkg, "-", C_DEL_TEXT)
        if not summary["added"] and not summary["removed"]:
            _bullet(self.doc, "Actualizacion de metadatos internos (integridad, resoluciones).", "-", C_MUTED)

    def _render_structural_summary(self, fc: FileChange):
        """Renderizado de resumen estructural para archivos extensos o con interfaces."""
        p_lbl = self.doc.add_paragraph()
        _run(p_lbl, "Resumen Estructural del archivo:", bold=True, color=C_TITLE, size=9)

        p_exec = self.doc.add_paragraph()
        _run(p_exec, "Sintesis funcional del cambio:", bold=True, color=C_SUBTITLE, size=9)
        for item in fc.build_functional_summary():
            _bullet(self.doc, item, "-", C_BODY, indent=1.5)

        struct = fc.extract_structure()

        if struct["decorators"]:
            p_d = self.doc.add_paragraph()
            _run(p_d, "Decoradores detectados:", bold=True, color=C_SUBTITLE, size=9)
            for d in struct["decorators"]:
                _bullet(self.doc, d, "-", C_REF_TEXT, indent=1.5)

        if struct["imports"]:
            p_i = self.doc.add_paragraph()
            _run(p_i, "Dependencias / Imports:", bold=True, color=C_SUBTITLE, size=9)
            for imp in struct["imports"]:
                _bullet(self.doc, imp, "-", C_BODY, indent=1.5)

        if struct["entities"]:
            p_e = self.doc.add_paragraph()
            _run(p_e, "Estructuras Logicas (Clases / Funciones / Interfaces / Tipos):",
                 bold=True, color=C_SUBTITLE, size=9)
            for ent in struct["entities"]:
                _bullet(self.doc, ent, "-", C_BODY, indent=1.5)

        if struct["routes"]:
            p_r = self.doc.add_paragraph()
            _run(p_r, "Rutas registradas:", bold=True, color=C_SUBTITLE, size=9)
            for rt in struct["routes"]:
                _bullet(self.doc, rt, "-", C_MOD_TEXT, indent=1.5)

        if struct["ui_components"]:
            p_ui = self.doc.add_paragraph()
            _run(p_ui, "Componentes de interfaz detectados:", bold=True, color=C_SUBTITLE, size=9)
            for comp in struct["ui_components"][:12]:
                _bullet(self.doc, comp, "-", C_MOD_TEXT, indent=1.5)

        if struct["forms"]:
            p_f = self.doc.add_paragraph()
            _run(p_f, "Controles de formulario/accion:", bold=True, color=C_SUBTITLE, size=9)
            for ctrl in struct["forms"][:12]:
                _bullet(self.doc, ctrl, "-", C_BODY, indent=1.5)

        if struct["angular_bindings"]:
            p_b = self.doc.add_paragraph()
            _run(p_b, "Bindings y directivas Angular:", bold=True, color=C_SUBTITLE, size=9)
            for bind in struct["angular_bindings"][:12]:
                _bullet(self.doc, bind, "-", C_BODY, indent=1.5)

        if struct["events"]:
            p_ev = self.doc.add_paragraph()
            _run(p_ev, "Eventos de usuario detectados:", bold=True, color=C_SUBTITLE, size=9)
            for ev in struct["events"][:12]:
                _bullet(self.doc, ev, "-", C_REF_TEXT, indent=1.5)

        if struct["css_selectors"]:
            p_css = self.doc.add_paragraph()
            _run(p_css, "Selectores/Reglas de estilo detectados:", bold=True, color=C_SUBTITLE, size=9)
            for sel in struct["css_selectors"][:12]:
                _bullet(self.doc, sel, "-", C_BODY, indent=1.5)

        if not any(struct.values()):
            _bullet(
                self.doc,
                "Contenido estructurado sin entidades detectables (datos, configuracion, markup).",
                "-", C_BODY
            )

    def _render_line_detail(self, fc: FileChange):
        """Renderizado linea a linea para cambios acotados con clasificacion de razon."""
        if fc.added:
            p_add = self.doc.add_paragraph()
            _run(p_add, "Lineas anadidas:", bold=True, color=C_ADD_TEXT, size=9)
            for line_no, line in (fc.added_with_line or [(None, l) for l in fc.added]):
                if line_no is not None:
                    _bullet(self.doc, f"L{line_no}: {line}", "+", C_ADD_TEXT)
                else:
                    _bullet(self.doc, line, "+", C_ADD_TEXT)

        if fc.removed:
            p_rem = self.doc.add_paragraph()
            p_rem.paragraph_format.space_before = Pt(4)
            _run(p_rem, "Lineas eliminadas:", bold=True, color=C_DEL_TEXT, size=9)

            TAG_COLORS = {
                "Linter":        C_ADD_TEXT,
                "Limpieza":      C_MUTED,
                "Debug":         C_MUTED,
                "Doc":           C_MUTED,
                "Deuda tecnica": C_REF_TEXT,
                "Config":        C_MOD_TEXT,
                "Refactor":      C_REF_TEXT,
                "Test":          C_SUBTITLE,
            }

            removed_items = fc.removed_with_line or [(None, l) for l in fc.removed]
            for line_no, line in removed_items:
                reason = fc.classify_removed_line(line)
                line_display = f"L{line_no}: {line}" if line_no is not None else line
                if reason:
                    tag_match = re.match(r'\[([^\]]+)\]\s*(.*)', reason)
                    if tag_match:
                        tag_key  = tag_match.group(1)
                        tag_desc = tag_match.group(2)
                        tc_tag   = TAG_COLORS.get(tag_key, C_SUBTITLE)
                        p_line   = self.doc.add_paragraph()
                        p_line.paragraph_format.left_indent       = Cm(1.0)
                        p_line.paragraph_format.first_line_indent = Cm(-0.5)
                        p_line.paragraph_format.space_after       = Pt(2)
                        _run(p_line, "-  ", bold=True, color=C_DEL_TEXT, size=9)
                        _run(p_line, line_display, color=C_DEL_TEXT, size=9)
                        _run(p_line, f"  [{tag_key}: {tag_desc}]",
                             bold=True, color=tc_tag, size=8, italic=True)
                    else:
                        _bullet(self.doc, f"{line_display}  ({reason})", "-", C_DEL_TEXT)
                else:
                    _bullet(self.doc, line_display, "-", C_DEL_TEXT)

    def _render_eslint_report(self, eslint_report: "ESLintFileReport", fc: "FileChange"):
        """
        Renderiza las tablas ESLint (correcciones, violaciones y cambios funcionales)
        de un archivo Angular (.ts o .html) dentro del informe docx.
        """
        ext = fc.ext.lower()
        is_html = ext in ('.html', '.component.html')
        is_ts   = ext in (
            '.ts', '.component.ts', '.service.ts', '.pipe.ts',
            '.directive.ts', '.guard.ts', '.interceptor.ts',
            '.module.ts', '.resolver.ts',
        )
        if not (is_html or is_ts):
            return
        if not eslint_report.corrections and not eslint_report.violations:
            p_ok = self.doc.add_paragraph()
            p_ok.paragraph_format.space_before = Pt(4)
            _run(p_ok, "Análisis ESLint: ", bold=True, color=C_TITLE, size=9)
            _run(p_ok, "No se detectaron correcciones ni violaciones ESLint en las líneas añadidas.",
                 color=C_MUTED, size=9, italic=True)
            return

        # ── Correcciones ESLint ───────────────────────────────────────────────
        if eslint_report.corrections:
            p_corr = self.doc.add_paragraph()
            p_corr.paragraph_format.space_before = Pt(6)
            _run(p_corr, f"✔ Correcciones ESLint aplicadas ({len(eslint_report.corrections)})",
                 bold=True, color=C_ADD_TEXT, size=10)

            CORR_COLS = [
                ("#",             0.5),
                ("Antes (−)",     4.8),
                ("Después (+)",   4.8),
                ("Regla corregida", 3.4),
                ("Severidad",     1.3),
            ]
            tbl_c = self.doc.add_table(rows=1, cols=len(CORR_COLS))
            tbl_c.style = "Table Grid"
            _header_row(tbl_c, CORR_COLS)

            for idx, f in enumerate(eslint_report.corrections, 1):
                before = (f.line_removed[:90] + "…") if len(f.line_removed) > 90 else f.line_removed
                after  = (f.line_added[:90]   + "…") if len(f.line_added)   > 90 else f.line_added
                _data_row(tbl_c, [
                    (str(idx),    0.5,  C_MUTED,    False, C_WHITE),
                    (before,      4.8,  C_DEL_TEXT, False, C_DEL_BG),
                    (after,       4.8,  C_ADD_TEXT, False, C_ADD_BG),
                    (f.rule,      3.4,  C_MOD_TEXT, False, C_ROW_ALT),
                    (f.severity,  1.3,  C_ADD_TEXT, True,  C_ADD_BG),
                ])
            self.doc.add_paragraph()

        # ── Violaciones ESLint ────────────────────────────────────────────────
        if eslint_report.violations:
            p_viol = self.doc.add_paragraph()
            p_viol.paragraph_format.space_before = Pt(6)
            _run(p_viol, f"✘ Violaciones ESLint introducidas ({len(eslint_report.violations)})",
                 bold=True, color=C_DEL_TEXT, size=10)

            VIOL_COLS = [
                ("#",            0.5),
                ("Línea añadida (+)", 6.0),
                ("Regla violada",    5.5),
                ("Severidad",        1.3),
            ]
            tbl_v = self.doc.add_table(rows=1, cols=len(VIOL_COLS))
            tbl_v.style = "Table Grid"
            _header_row(tbl_v, VIOL_COLS)

            for idx, f in enumerate(eslint_report.violations, 1):
                line_disp = (f.line_added[:110] + "…") if len(f.line_added) > 110 else f.line_added
                _data_row(tbl_v, [
                    (str(idx),    0.5,  C_MUTED,    False, C_WHITE),
                    (line_disp,   6.0,  C_DEL_TEXT, False, C_DEL_BG),
                    (f.rule,      5.5,  C_MOD_TEXT, False, C_ROW_ALT),
                    (f.severity,  1.3,  C_DEL_TEXT, True,  C_DEL_BG),
                ])
            self.doc.add_paragraph()

    def _detail(self):
        semantic_engine = SemanticInsightEngine()
        relation_engine = ChangeRelationAnalyzer(self.changes)
        eslint_analyzer  = ESLintAngularAnalyzer()
        _h1(self.doc, "2. Detalle de Cambios por Archivo")

        for i, fc in enumerate(self.changes):
            _section_title(self.doc, fc.filename)

            p_path = self.doc.add_paragraph()
            p_path.paragraph_format.space_after = Pt(3)
            _run(p_path, fc.filepath, color=C_MUTED, size=8, italic=True)

            if fc.contexts:
                p_ctx = self.doc.add_paragraph()
                p_ctx.paragraph_format.space_after = Pt(3)
                _run(p_ctx, "Bloques / Funciones afectadas: ", bold=True, color=C_SUBTITLE, size=9)
                _run(p_ctx, ", ".join(fc.contexts), color=C_BODY, size=9, italic=True)

            deploy_level = calculate_deploy_impact(fc)
             # -----------------------------------------
            # ANALISIS SEMANTICO CONTEXTUAL
            # -----------------------------------------
            semantic_insights = semantic_engine.analyze(fc)
            if semantic_insights:
                p_sem = self.doc.add_paragraph()
                _run(p_sem, "Analisis Semantico Detectado:",
                     bold=True, color=C_TITLE, size=9)

                for insight in semantic_insights:
                    _bullet(self.doc, insight, ">", C_MOD_TEXT)

            relation_insights = relation_engine.analyze(fc)
            if relation_insights:
                p_rel = self.doc.add_paragraph()
                _run(p_rel, "Interrelaciones Detectadas:",
                     bold=True, color=C_TITLE, size=9)
                for insight in relation_insights:
                    _bullet(self.doc, insight, " > ", C_REF_TEXT)

            tc_dep, _    = impact_colors(deploy_level)
            p_stats = self.doc.add_paragraph()
            p_stats.paragraph_format.space_after = Pt(4)
            _run(p_stats, f"+{len(fc.added)} anadidas   -{len(fc.removed)} eliminadas   ",
                 color=C_MUTED, size=8)
            _run(p_stats, f"Impacto en deploy: {deploy_level}", bold=True, color=tc_dep, size=8)

            # ── Análisis ESLint (antes del renderizado de líneas) ──────────────
            eslint_report = eslint_analyzer.analyze_file(fc)
            self._eslint_reports[fc.filepath] = eslint_report

            # Elegir estrategia de renderizado
            if fc.is_binary:
                _bullet(self.doc, "Archivo binario. No se puede mostrar contenido de texto.",
                        "-", C_MUTED)
            elif fc.is_lockfile:
                self._render_lockfile_detail(fc)
            elif fc.needs_structural_summary:
                self._render_structural_summary(fc)
            else:
                self._render_line_detail(fc)

            # ── Reporte ESLint (correcciones / violaciones) ─────────────────────
            self._render_eslint_report(eslint_report, fc)

            if not fc.added and not fc.removed and not fc.is_binary:
                p_empty = self.doc.add_paragraph()
                _run(
                    p_empty,
                    "Sin modificaciones en el codigo fuente. Posible cambio de "
                    "permisos, archivo vacio o renombramiento.",
                    color=C_MUTED, size=9, italic=True
                )

            if i < len(self.changes) - 1:
                _divider(self.doc)

        self.doc.add_paragraph()

    def _impact_legend(self):
        _h1(self.doc, "3. Leyenda de Niveles de Impacto en Deploy")

        legend = [
            ("Nula",
             "Sin riesgo. Documentacion, solo linter/formatter o configuracion menor."),
            ("Leve",
             "Archivo de dependencias, nuevo archivo aislado o cambio estetico."),
            ("Baja",
             "Nuevo componente/servicio sin contratos externos. Logica interna acotada."),
            ("Media",
             "Cambio en logica de negocio, llamadas HTTP, rutas o estado compartido."),
            ("Alta",
             "Interfaces/tipos exportados, modulos Angular, guards o archivos de entorno."),
            ("Critica",
             "Autenticacion, JWT, passwords, migraciones de BD o contratos de API externa."),
        ]

        tbl = self.doc.add_table(rows=1, cols=2)
        tbl.style = "Table Grid"
        tbl.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _header_row(tbl, [("Nivel", 2.5), ("Descripcion del criterio", 11.7)])

        for level, desc in legend:
            tc_l, bg_l = impact_colors(level)
            _data_row(tbl, [
                (level, 2.5,  tc_l,   True,  bg_l),
                (desc,  11.7, C_BODY, False, C_WHITE),
            ])

        self.doc.add_paragraph()

    def _recommendations(self):
        _h1(self.doc, "4. Recomendaciones antes del Merge")
        recs = analyze_recommendations(self.changes)
        p = self.doc.add_paragraph()
        p.paragraph_format.space_after = Pt(7)
        _run(
            p,
            "Recomendaciones generadas a partir del analisis del codigo modificado:",
            color=C_BODY, size=10
        )
        for rec in recs:
            _bullet(self.doc, rec, " > ", C_MOD_TEXT)
        self.doc.add_paragraph()

    def _eslint_global_summary(self):
        """Sección 5: Resumen global del análisis ESLint para todo el diff."""
        if not self._eslint_reports:
            return

        all_corrections: List[ESLintFinding] = []
        all_violations:  List[ESLintFinding] = []
        for report in self._eslint_reports.values():
            all_corrections.extend(report.corrections)
            all_violations.extend(report.violations)

        # Contar frecuencia por regla
        from collections import Counter
        corr_counter = Counter(f.rule for f in all_corrections)
        viol_counter = Counter(f.rule for f in all_violations)

        _h1(self.doc, "5. Resumen Global del Análisis ESLint")

        # Estadísticas generales
        ts_files   = sum(1 for fc in self.changes
                         if fc.ext.lower() in ('.ts', '.component.ts', '.service.ts',
                                               '.pipe.ts', '.directive.ts', '.guard.ts',
                                               '.interceptor.ts', '.module.ts', '.resolver.ts'))
        html_files = sum(1 for fc in self.changes
                         if fc.ext.lower() in ('.html', '.component.html'))

        p_stat = self.doc.add_paragraph()
        p_stat.paragraph_format.space_after = Pt(6)
        _run(p_stat,
             f"Archivos TypeScript analizados: {ts_files}   "
             f"Archivos HTML analizados: {html_files}   "
             f"Correcciones ESLint: {len(all_corrections)}   "
             f"Violaciones ESLint: {len(all_violations)}",
             color=C_MUTED, size=9, italic=True)

        # ── Tabla de correcciones por regla ───────────────────────────────────
        if corr_counter:
            p_c = self.doc.add_paragraph()
            _run(p_c, "Reglas más frecuentemente corregidas:", bold=True,
                 color=C_ADD_TEXT, size=10)
            COLS_C = [("#", 0.5), ("Regla ESLint", 8.5), ("Veces corregida", 2.2), ("Severidad", 1.5)]
            tbl_c = self.doc.add_table(rows=1, cols=len(COLS_C))
            tbl_c.style = "Table Grid"
            _header_row(tbl_c, COLS_C)
            for rank, (rule, count) in enumerate(corr_counter.most_common(10), 1):
                _data_row(tbl_c, [
                    (str(rank), 0.5,  C_MUTED,    False, C_WHITE),
                    (rule,      8.5,  C_MOD_TEXT, False, C_ROW_ALT),
                    (str(count),2.2,  C_ADD_TEXT, True,  C_ADD_BG),
                    ("ERROR",   1.5,  C_ADD_TEXT, True,  C_ADD_BG),
                ])
            self.doc.add_paragraph()

        # ── Tabla de violaciones por regla ────────────────────────────────────
        if viol_counter:
            p_v = self.doc.add_paragraph()
            _run(p_v, "Reglas con violaciones introducidas:", bold=True,
                 color=C_DEL_TEXT, size=10)
            COLS_V = [("#", 0.5), ("Regla ESLint", 8.5), ("Ocurrencias", 2.2), ("Severidad", 1.5)]
            tbl_v = self.doc.add_table(rows=1, cols=len(COLS_V))
            tbl_v.style = "Table Grid"
            _header_row(tbl_v, COLS_V)
            for rank, (rule, count) in enumerate(viol_counter.most_common(10), 1):
                _data_row(tbl_v, [
                    (str(rank), 0.5,  C_MUTED,    False, C_WHITE),
                    (rule,      8.5,  C_MOD_TEXT, False, C_ROW_ALT),
                    (str(count),2.2,  C_DEL_TEXT, True,  C_DEL_BG),
                    ("ERROR",   1.5,  C_DEL_TEXT, True,  C_DEL_BG),
                ])
            self.doc.add_paragraph()

        if not all_corrections and not all_violations:
            p_ok = self.doc.add_paragraph()
            _run(p_ok,
                 "No se detectaron correcciones ni violaciones ESLint en ningún archivo "
                 "TypeScript o HTML del diff.",
                 color=C_MUTED, size=9, italic=True)
            self.doc.add_paragraph()

    def _footer(self):
        _divider(self.doc)
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(6)
        total_add = sum(len(f.added) for f in self.changes)
        total_del = sum(len(f.removed) for f in self.changes)
        _run(
            p,
            f"Archivos: {len(self.changes)}  |  "
            f"+{total_add} lineas  |  "
            f"-{total_del} lineas  |  "
            f"Archivo generado automaticamente - Ryu Gabo -",
            color=C_MUTED, size=9.5, italic=True
        )

    def generate(self) -> str:
        self._title()
        self._summary()
        self._detail()
        self._impact_legend()
        self._recommendations()
        self._eslint_global_summary()
        self._footer()
        self.doc.save(self.output)
        return self.output


# =============================================================================
# CLI Y EJECUCION
# =============================================================================
def find_input(arg: Optional[str]) -> Path:
    if arg:
        p = Path(arg)
        if not p.exists():
            print(f"[ERROR] No se encontro: {p}", file=sys.stderr)
            sys.exit(1)
        return p
    for candidate in (Path.cwd() / DEFAULT_INPUT, Path(__file__).parent / DEFAULT_INPUT):
        if candidate.exists():
            return candidate
    print(
        f"[ERROR] No se encontro '{DEFAULT_INPUT}'.\n"
        f"Genera el archivo con:\n"
        f"  git --no-pager diff --staged -U9999 > {DEFAULT_INPUT}",
        file=sys.stderr
    )
    sys.exit(1)


def main():
    ap = argparse.ArgumentParser(
        description="Convierte 'informe.txt' (git diff) en un informe .docx profesional."
    )
    ap.add_argument("--input",       "-i",  default=None,       help=f"Archivo diff (default: {DEFAULT_INPUT})")
    ap.add_argument("--output",      "-o",  default=None,       help=f"Archivo salida (default: {DEFAULT_OUTPUT})")
    ap.add_argument("--branch-from", "-bf", default="gabotest", help="Rama origen  (default: gabotest)")
    ap.add_argument("--branch-to",   "-bt", default="develop",  help="Rama destino (default: develop)")
    args = ap.parse_args()

    input_path = find_input(args.input)
    print(f"[INFO] Leyendo: {input_path}")

    text = clean(input_path.read_text(encoding="utf-8", errors="replace"))
    if not text.strip():
        print("[ERROR] El archivo esta vacio.", file=sys.stderr)
        sys.exit(1)

    changes = DiffParser().parse(text)
    if not changes:
        print("[AVISO] No se encontraron cambios legibles.", file=sys.stderr)
        sys.exit(1)

    output = args.output or str(input_path.parent / DEFAULT_OUTPUT)
    result = ReportGenerator(
        changes=changes,
        branch_from=args.branch_from,
        branch_to=args.branch_to,
        output=output
    ).generate()

    print(f"\n[OK] Informe generado: {result}")


if __name__ == "__main__":
    main()