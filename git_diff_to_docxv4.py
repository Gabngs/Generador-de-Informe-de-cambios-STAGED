#!/usr/bin/env python3
"""
git_diff_to_docx.py  v4
────────────────────────
Lee 'informe.txt' (en la misma carpeta) generado con:
    git --no-pager diff --staged -U9999 > informe.txt

Genera un informe profesional .docx analizando la lógica
de los cambios, resumiendo archivos nuevos y detectando impacto.

MEJORAS v4 (motores de análisis):
  • Motor infer_modification_purpose: deduce el PROPÓSITO de alto nivel de
    cada cambio (corrección de bug, refactor, nueva funcionalidad, limpieza,
    cambio de contrato, optimización, seguridad, i18n, estilo/linter).
  • build_functional_summary: ahora encabeza cada resumen con el propósito
    inferido (ícono 🎯) antes de los detalles estructurales.
  • analyze_php_logic_changes: incorpora propósito al inicio del bloque PHP.
  • SemanticInsightEngine: 10 nuevos análisis (17-26):
      17. Migración callback → async/await (TS).
      18. Extracción de métodos / SRP (TS).
      19. Refuerzo de tipado any → tipo concreto (TS).
      20. Cambio de tipo de JOIN en SQL.
      21. Gestión de efectos secundarios React/Vue (useEffect/watch).
      22. Renombrado semántico de identificador.
      23. Introducción / eliminación de try/catch.
      24. Hardcoded → constante/variable configurable.
      25. Eliminación de código comentado (deuda técnica).
      26. Píxeles → unidades relativas CSS (accesibilidad).
  • classify_removed_line: añade motivo explicativo a cada categoría
    (por qué se eliminó, qué riesgo evita, qué mejora aporta).
"""

import sys
import re
import argparse
from datetime import datetime
from pathlib import Path
from dataclasses import dataclass, field
from typing import List, Tuple, Optional, Dict, Set

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# =============================================================================
# PALETA DE COLORES
# =============================================================================
C_TITLE     = RGBColor(0x1E, 0x3A, 0x5F)
C_SUBTITLE  = RGBColor(0x55, 0x55, 0x55)
C_BODY      = RGBColor(0x22, 0x22, 0x22)
C_MUTED     = RGBColor(0x88, 0x88, 0x88)
C_ADD_TEXT  = RGBColor(0x16, 0x65, 0x34)
C_ADD_BG    = "F0FDF4"
C_DEL_TEXT  = RGBColor(0x99, 0x1B, 0x1B)
C_DEL_BG    = "FEF2F2"
C_MOD_TEXT  = RGBColor(0x1E, 0x3A, 0x5F)
C_MOD_BG    = "EFF6FF"
C_REF_TEXT  = RGBColor(0x92, 0x40, 0x0E)
C_REF_BG    = "FFFBEB"
C_HDR_BG    = "1E3A5F"
C_ACCENT    = "2563EB"
C_BORDER    = "CCCCCC"
C_ROW_ALT   = "F8FAFC"
C_WHITE     = "FFFFFF"

# Colores para niveles de impacto
C_IMPACT_NULA     = RGBColor(0x6B, 0x72, 0x80)
C_IMPACT_LEVE     = RGBColor(0x05, 0x96, 0x69)
C_IMPACT_BAJA     = RGBColor(0x0D, 0x94, 0x88)
C_IMPACT_MEDIA    = RGBColor(0xD9, 0x77, 0x06)
C_IMPACT_ALTA     = RGBColor(0xDC, 0x25, 0x26)
C_IMPACT_CRITICA  = RGBColor(0x7F, 0x1D, 0x1D)

BG_IMPACT_NULA    = "F3F4F6"
BG_IMPACT_LEVE    = "ECFDF5"
BG_IMPACT_BAJA    = "F0FDFA"
BG_IMPACT_MEDIA   = "FFFBEB"
BG_IMPACT_ALTA    = "FFF1F2"
BG_IMPACT_CRITICA = "FEF2F2"

DEFAULT_INPUT  = "informe.txt"
DEFAULT_OUTPUT = "informe_cambios.docx"

ANSI_RE = re.compile(r'\x1B(?:[@-Z\\-_]|\[[0-?]*[ -/]*[@-~])')

MESES_ES = {
    1: "enero", 2: "febrero", 3: "marzo", 4: "abril",
    5: "mayo", 6: "junio", 7: "julio", 8: "agosto",
    9: "septiembre", 10: "octubre", 11: "noviembre", 12: "diciembre"
}

def fecha_espanol() -> str:
    hoy = datetime.now()
    return f"{hoy.day} de {MESES_ES[hoy.month]} de {hoy.year}"

def clean(text: str) -> str:
    text = text.replace('\r\n', '\n').replace('\r', '\n')
    text = ANSI_RE.sub('', text)
    return text.replace('\x00', '')

# =============================================================================
# NIVELES DE IMPACTO
# =============================================================================
def impact_colors(level: str) -> Tuple[RGBColor, str]:
    return {
        "Nula":    (C_IMPACT_NULA,    BG_IMPACT_NULA),
        "Leve":    (C_IMPACT_LEVE,    BG_IMPACT_LEVE),
        "Baja":    (C_IMPACT_BAJA,    BG_IMPACT_BAJA),
        "Media":   (C_IMPACT_MEDIA,   BG_IMPACT_MEDIA),
        "Alta":    (C_IMPACT_ALTA,    BG_IMPACT_ALTA),
        "Critica": (C_IMPACT_CRITICA, BG_IMPACT_CRITICA),
    }.get(level, (C_BODY, C_WHITE))

# =============================================================================
# MODELO: FileChange
# =============================================================================
@dataclass
class FileChange:
    filepath: str
    added:    List[str] = field(default_factory=list)
    removed:  List[str] = field(default_factory=list)
    added_with_line: List[Tuple[Optional[int], str]] = field(default_factory=list)
    removed_with_line: List[Tuple[Optional[int], str]] = field(default_factory=list)
    full_content: List[str] = field(default_factory=list)
    contexts: Set[str]  = field(default_factory=set)
    kind: str = "modified"
    is_binary: bool = False

    @property
    def filename(self) -> str:
        return Path(self.filepath).name

    @property
    def ext(self) -> str:
        name = self.filename.lower()
        for double in ('.component.ts', '.component.html', '.component.scss',
                       '.service.ts', '.spec.ts', '.module.ts', '.pipe.ts',
                       '.directive.ts', '.guard.ts', '.interceptor.ts',
                       '.reducer.ts', '.action.ts', '.effect.ts', '.selector.ts',
                       '.resolver.ts', '.model.ts', '.interface.ts', '.enum.ts',
                       '.helper.ts', '.util.ts', '.config.ts', '.constant.ts',
                       # PHP / Laravel
                       '.blade.php'):
            if name.endswith(double):
                return double
        return Path(self.filepath).suffix.lower()

    @property
    def is_lockfile(self) -> bool:
        """Detecta archivos de lock/dependencias que no aportan valor en detalle linea a linea."""
        name = self.filename.lower()
        # Sufijo .bak: siempre es copia de respaldo, sin valor linea a linea
        if name.endswith('.bak'):
            return True
        return name in (
            'package-lock.json', 'yarn.lock', 'pnpm-lock.yaml',
            'composer.lock', 'gemfile.lock', 'poetry.lock', 'cargo.lock',
            'packages.lock.json', 'shrinkwrap.json',
            # .NET / NuGet
            'packages.config',
            # Go
            'go.sum',
            # Python
            'pipfile.lock',
            # Elixir
            'mix.lock',
            # Flutter / Dart
            'pubspec.lock',
            # Nix
            'flake.lock',
            # Swift
            'package.resolved',
        )

    @property
    def is_environment_file(self) -> bool:
        """Detecta archivos de entorno/configuracion de ambiente."""
        name = self.filename.lower()
        # Angular environments
        if name.startswith('environment') and name.endswith('.ts'):
            return True
        # ASP.NET Core appsettings
        if name.startswith('appsettings') and name.endswith('.json'):
            return True
        # Spring Boot
        if name.startswith('application') and name.endswith(('.properties', '.yml', '.yaml')):
            return True
        return name in (
            '.env', '.env.local', '.env.production', '.env.staging', '.env.qa',
            '.env.example',
            # .NET Framework
            'web.config', 'app.config',
            # Django / Flask
            'settings.py',
            # Rails
            'database.yml',
        )

    @property
    def needs_structural_summary(self) -> bool:
        """
        FIX BUG 2: Umbral inteligente por tipo de archivo.
        Determina si el archivo debe mostrarse como resumen estructural
        en vez de linea por linea.
        """
        total = len(self.added) + len(self.removed)

        # Lock files: siempre resumen (nunca linea a linea)
        if self.is_lockfile:
            return True

        # Archivos grandes (mas de 60 lineas totales)
        if total > 60:
            return True

        # Archivos nuevos con mas de 15 lineas: resumen estructural
        if self.kind == 'added' and len(self.added) > 15:
            return True

        # Archivos con interfaces/tipos exportados (aunque sean pequenos)
        all_text = "\n".join(self.added)
        if re.search(r'export\s+(interface|type|enum|class)\b', all_text) and len(self.added) > 10:
            return True

        # Archivos de respaldo: siempre resumen (son copias completas)
        if self.ext == '.bak':
            return True

        # Soluciones y proyectos de Visual Studio: siempre resumen
        if self.ext in ('.sln', '.vbproj', '.csproj', '.vcxproj', '.fsproj'):
            return True

        # Archivos de recursos / XML / config grandes
        if self.ext in ('.resx', '.xml', '.config', '.manifest') and total > 20:
            return True

        return False

    @property
    def kind_label(self) -> str:
        if self.kind == 'added':   return 'Adicion'
        if self.kind == 'deleted': return 'Eliminacion'
        if self.kind == 'renamed': return 'Renombrado'
        if self.is_binary:         return 'Binario/Media'
        n_add, n_del = len(self.added), len(self.removed)
        if n_del == 0 and n_add > 0: return 'Adicion'
        if n_add == 0 and n_del > 0: return 'Eliminacion'
        if n_del > n_add * 2 and n_del > 10: return 'Refactor'
        if n_add == 0 and n_del == 0: return 'Configuracion'
        return 'Modificacion'

    @property
    def kind_colors(self) -> Tuple[RGBColor, str]:
        lbl = self.kind_label
        if lbl == 'Adicion':     return C_ADD_TEXT, C_ADD_BG
        if lbl == 'Eliminacion': return C_DEL_TEXT, C_DEL_BG
        if lbl == 'Refactor':    return C_REF_TEXT, C_REF_BG
        return C_MOD_TEXT, C_MOD_BG

    def extract_structure(self) -> Dict[str, List[str]]:
        """Analiza logicamente el contenido para extraer estructura por tipo de archivo."""
        structure: Dict[str, List[str]] = {
            "imports": [],
            "entities": [],
            "decorators": [],
            "routes": [],
            "ui_components": [],
            "angular_bindings": [],
            "forms": [],
            "events": [],
            "css_selectors": [],
        }
        all_lines = self.added + self.removed

        import_pattern    = re.compile(r'^\s*(import|from|require\(|include|using)\b')
        entity_pattern    = re.compile(
            r'^\s*(export )?(class|def|function|interface|const \w+\s*=\s*\(|'
            r'let \w+\s*=\s*\(|async function|type\s+\w+\s*=|enum\s+\w+|'
            # VB.NET
            r'(?:Public|Private|Protected|Friend)\s+(?:Class|Module|Sub|Function|Property|Enum|Interface|Structure)\s+|'
            r'Public\s+(?:Shared\s+)?(?:ReadOnly\s+)?(?:Sub|Function)\s+|'
            # C# / .NET
            r'(?:public|private|protected|internal|static|override|virtual|abstract)\s+(?:class|interface|record|struct|enum|void|async\s+Task)\s+|'
            r'(?:public|private|protected)\s+\w+\s+\w+\s*(?:\(|{)|'
            # Go
            r'func\s+(?:\([^)]*\)\s*)?\w+|'
            # Rust
            r'(?:pub\s+)?(?:fn|struct|enum|trait|impl)\s+\w+|'
            # Java / Kotlin
            r'(?:public|private|protected)\s+(?:static\s+)?(?:class|interface|enum|\w+)\s+\w+)'
        )
        decorator_pattern = re.compile(r'^\s*@\w+')
        route_pattern     = re.compile(
            r"(path\s*:\s*['\"]|route\s*\(|router\.(get|post|put|delete|patch)\s*\()", re.I
        )

        # HTML / Angular template
        ui_component_pattern = re.compile(r'<\s*(p-[\w-]+|app-[\w-]+)\b', re.I)
        angular_binding_pattern = re.compile(
            r'(\[\([^\)]*\)\]|\[[^\]]+\]|\([^\)]+\)|\*ng(?:If|For|Switch)|\bngModel\b|\bformControlName\b)'
        )
        event_pattern = re.compile(r'\((click|change|input|submit|keyup|keydown|blur|focus)\)\s*=')
        form_pattern = re.compile(r'<\s*(input|select|textarea|form|button|p-dropdown|p-inputtext|p-button)\b', re.I)

        # SCSS / CSS
        css_selector_pattern = re.compile(r'^\s*([.#][\w-]+|:host|::ng-deep|[a-zA-Z][\w-]*(?:\s+[a-zA-Z][\w-]*)?)\s*[{,]')

        for line in all_lines:
            line_stripped = line.strip()

            if import_pattern.search(line) and line not in structure["imports"]:
                structure["imports"].append(
                    line.strip()[:90] + ('...' if len(line) > 90 else '')
                )
            elif decorator_pattern.match(line) and line_stripped not in structure["decorators"]:
                structure["decorators"].append(line_stripped[:60])
            elif route_pattern.search(line) and line_stripped not in structure["routes"]:
                structure["routes"].append(line_stripped[:80])
            elif entity_pattern.search(line):
                clean_entity = line_stripped.split('{')[0].strip()
                if clean_entity not in structure["entities"]:
                    structure["entities"].append(clean_entity)

            if self.ext in ('.html', '.component.html'):
                comp_match = ui_component_pattern.search(line)
                if comp_match:
                    comp = comp_match.group(1).lower()
                    if comp not in structure["ui_components"]:
                        structure["ui_components"].append(comp)

                bind_matches = angular_binding_pattern.findall(line)
                for b in bind_matches:
                    b_clean = b.strip()
                    if b_clean and b_clean not in structure["angular_bindings"]:
                        structure["angular_bindings"].append(b_clean)

                ev_match = event_pattern.search(line)
                if ev_match:
                    ev = ev_match.group(1)
                    if ev not in structure["events"]:
                        structure["events"].append(ev)

                form_match = form_pattern.search(line)
                if form_match:
                    frm = form_match.group(1).lower()
                    if frm not in structure["forms"]:
                        structure["forms"].append(frm)

            if self.ext in ('.scss', '.css', '.component.scss'):
                sel_match = css_selector_pattern.search(line)
                if sel_match:
                    sel = sel_match.group(1).strip()
                    if sel and sel not in structure["css_selectors"]:
                        structure["css_selectors"].append(sel[:80])
        return structure

    def build_functional_summary(self) -> List[str]:
        """Genera una síntesis ejecutiva del cambio por archivo (sin detalle línea a línea).
        Incluye inferencia del propósito del cambio al inicio."""
        summary: List[str] = []
        n_add = len(self.added)
        n_del = len(self.removed)
        struct = self.extract_structure()

        # Propósito general inferido al inicio del resumen
        purpose = infer_modification_purpose(self.added, self.removed, self.ext)
        if purpose:
            summary.append(f"🎯 {purpose}")

        if self.ext in ('.html', '.component.html'):
            if self.kind == 'added':
                summary.append("Se incorpora una nueva plantilla de interfaz para la funcionalidad del módulo.")
            elif self.kind == 'modified':
                summary.append("Se ajusta la estructura del template para mejorar interacción y visualización.")

            if struct["ui_components"]:
                comps = ", ".join(struct["ui_components"][:6])
                summary.append(f"Se integran componentes de UI: {comps}.")
            if struct["forms"]:
                controls = ", ".join(struct["forms"][:6])
                summary.append(f"Se agregan controles de captura/acción: {controls}.")
            if struct["angular_bindings"]:
                summary.append(
                    f"Se detectan bindings/directivas Angular ({len(struct['angular_bindings'])} usos) "
                    "para enlazar estado y eventos del componente."
                )
            if struct["events"]:
                events = ", ".join(struct["events"][:6])
                summary.append(f"Se registran eventos de interacción de usuario: {events}.")

        elif self.ext in ('.scss', '.css', '.component.scss'):
            if self.kind == 'added':
                summary.append("Se incorpora hoja de estilos asociada al componente para estandarizar la presentación.")
            if struct["css_selectors"]:
                sels = ", ".join(struct["css_selectors"][:6])
                summary.append(f"Se definen selectores/reglas relevantes: {sels}.")

        elif self.ext in ('.ts', '.component.ts', '.service.ts'):
            if struct["entities"]:
                ents = ", ".join(struct["entities"][:5])
                summary.append(f"Se incorporan/ajustan estructuras lógicas: {ents}.")
            if struct["imports"]:
                summary.append("Se actualizan dependencias para soportar la nueva lógica del archivo.")

        elif self.ext in ('.vb',):
            all_text = "\n".join(self.added + self.removed)
            if struct["entities"]:
                ents = ", ".join(struct["entities"][:5])
                summary.append(f"Se incorporan/modifican rutinas VB.NET: {ents}.")
            if re.search(r'\bSub\s+New\b|\bSub\s+Form_Load\b', all_text, re.I):
                summary.append("Se ajusta la inicialización del formulario.")
            if re.search(r"'\s*(Private|Public)\s+(Sub|Function)", all_text):
                summary.append("Se comenta o elimina código de eventos/funciones previamente activo.")
            if re.search(r'EsCajaApp|Obtener_Caja_App|EstamosEnCajaPrincipal', all_text):
                summary.append("Se introduce lógica de tipo CajaApp para bypass condicional de reglas de caja.")
            if re.search(r'\bRealizarConsulta\b|\bDataTable\b|\bDataRow\b', all_text):
                summary.append("Se aplican cambios en acceso a datos (consultas o manejo de DataTable/DataRow).")

        elif self.ext in ('.cs',):
            all_text = "\n".join(self.added + self.removed)
            if struct["entities"]:
                ents = ", ".join(struct["entities"][:5])
                summary.append(f"Se incorporan/modifican clases o métodos C#: {ents}.")
            if re.search(r'\[HttpGet\]|\[HttpPost\]|\[HttpPut\]|\[HttpDelete\]|\[Route\(', all_text):
                summary.append("Se ajustan endpoints de API REST (atributos de enrutamiento HTTP).")
            if re.search(r'DbContext|SaveChangesAsync|AddDbContext', all_text):
                summary.append("Se modifica la interacción con base de datos mediante Entity Framework.")

        elif self.ext in ('.sln',):
            summary.append("Se modifica la solución Visual Studio: cambio de versión, proyectos referenciados o configuración de build.")

        elif self.ext in ('.bak',):
            summary.append("Archivo de respaldo generado automáticamente por el IDE. Refleja el estado previo antes de la modificación del archivo original.")
            summary.append(f"Contiene {len(self.added)} líneas del snapshot anterior.")

        elif self.ext in ('.java', '.kt'):
            all_text = "\n".join(self.added + self.removed)
            if struct["entities"]:
                ents = ", ".join(struct["entities"][:5])
                summary.append(f"Se incorporan/modifican clases o métodos: {ents}.")
            if re.search(r'@RestController|@Controller|@Service|@Repository', all_text):
                summary.append("Se ajustan componentes de la capa Spring (Controller/Service/Repository).")
            if re.search(r'@GetMapping|@PostMapping|@PutMapping|@DeleteMapping', all_text):
                summary.append("Se modifican endpoints REST de Spring MVC.")

        elif self.ext in ('.go',):
            all_text = "\n".join(self.added + self.removed)
            if struct["entities"]:
                ents = ", ".join(struct["entities"][:5])
                summary.append(f"Se incorporan/modifican funciones Go: {ents}.")
            if re.search(r'\bhttp\.HandleFunc\b|\bgin\.\w+\b|\becho\.\w+\b', all_text):
                summary.append("Se ajustan handlers de rutas HTTP en Go.")

        elif self.ext in ('.rs',):
            if struct["entities"]:
                ents = ", ".join(struct["entities"][:5])
                summary.append(f"Se incorporan/modifican estructuras Rust: {ents}.")

        elif self.ext in ('.php', '.blade.php'):
            php_insights = analyze_php_logic_changes(self.added, self.removed,
                                                     self.added_with_line, self.removed_with_line)
            summary.extend(php_insights[:4])
            if struct["entities"]:
                ents = ", ".join(struct["entities"][:4])
                summary.append(f"Funciones/métodos afectados: {ents}.")

        if not summary:
            if self.kind == 'added':
                summary.append("Se agrega archivo nuevo con estructura inicial funcional.")
            elif self.kind == 'deleted':
                summary.append("Se elimina archivo y su responsabilidad asociada del módulo.")
            else:
                summary.append("Se aplican ajustes internos en la implementación del archivo.")

        summary.append(f"Balance de cambio: +{n_add} líneas / -{n_del} líneas.")
        return summary[:5]

    def _extract_deleted_identifiers(self, line: str) -> List[str]:
        line = line.strip()
        m_ts_braces = re.search(r'import\s+\{([^}]+)\}', line)
        if m_ts_braces:
            return [x.strip() for x in m_ts_braces.group(1).split(',')]
        m_ts_simple = re.search(r'import\s+([a-zA-Z0-9_]+)\s+from', line)
        if m_ts_simple:
            return [m_ts_simple.group(1)]
        m_py_from = re.search(r'from\s+\S+\s+import\s+(.+)', line)
        if m_py_from:
            return [x.strip() for x in m_py_from.group(1).split(',')]
        m_var = re.search(
            r'(?:(?:private|public|protected|readonly|const|let|var|static)\s+)+([a-zA-Z0-9_]+)',
            line
        )
        if m_var:
            return [m_var.group(1)]
        m_simple_var = re.match(r'^([a-zA-Z0-9_]+)\s*[=:]', line)
        if m_simple_var:
            return [m_simple_var.group(1)]
        return []

    def verify_dead_code(self, deleted_line: str) -> bool:
        if not self.full_content:
            return False
        identifiers = [i for i in self._extract_deleted_identifiers(deleted_line) if i]
        if not identifiers:
            return False
        for ident in identifiers:
            pattern = re.compile(rf'\b{re.escape(ident)}\b')
            if not any(pattern.search(line) for line in self.full_content):
                return True
        return False

    def detect_linter_fix(self, removed_line: str) -> str:
        """
        Detecta si la eliminacion de una linea fue por correccion de linter/formatter.
        FIX BUG 3: Corregido falso positivo en regla eqeqeq cuando r_line == a_line
                   y la linea no contiene operadores de comparacion.
        """
        if not self.added:
            return ""
        r_line = removed_line.strip()
        if not r_line:
            return ""

        for a_line in self.added:
            a_line = a_line.strip()
            if not a_line:
                continue

            # (regex negativo lookahead/lookbehind para evitar === y !==)
            has_equality_op = bool(re.search(r'(?<![=!<>])={2}(?!=)|(?<!!)!={1}(?!=)', r_line))
            if has_equality_op:
                r_eq = r_line.replace('!==', '\x00').replace('===', '\x01')
                r_eq = r_eq.replace('!=', '!==').replace('==', '===')
                r_eq = r_eq.replace('\x00', '!==').replace('\x01', '===')
                # Solo reportar si realmente cambio algo
                if r_eq == a_line and r_eq != r_line:
                    return "ESLint: Igualdad estricta (=== / !==)"

            # 2. ESLint prefer-const / no-var
            if 'let' in r_line and re.sub(r'\blet\b', 'const', r_line) == a_line:
                return "ESLint: Usar const en lugar de let"
            if 'var' in r_line and re.sub(r'\bvar\b', 'let', r_line) == a_line:
                return "ESLint: Usar let en lugar de var"
            if 'var' in r_line and re.sub(r'\bvar\b', 'const', r_line) == a_line:
                return "ESLint: Usar const en lugar de var"

            # 3. Prettier quotes
            if "'" in r_line and r_line.replace("'", '"') == a_line:
                return "Prettier: Cambio a comillas dobles"
            if '"' in r_line and r_line.replace('"', "'") == a_line:
                return "Prettier: Cambio a comillas simples"

            # 4. Prettier semi
            if r_line + ';' == a_line:
                return "Prettier: Agregar punto y coma"
            if r_line == a_line + ';':
                return "Prettier: Eliminar punto y coma"

            # 5. Prettier: espaciado / indentacion (solo si las lineas son distintas)
            r_ns = re.sub(r'\s+', '', r_line)
            a_ns = re.sub(r'\s+', '', a_line)
            if r_ns == a_ns and r_line != a_line and len(r_ns) > 2:
                return "Prettier: Arreglo de espaciado/indentacion"

            # 6. ESLint trailing comma
            if (r_line.rstrip(',') == a_line.rstrip(',') and r_line != a_line
                    and r_line.endswith(',') != a_line.endswith(',') and len(r_line) > 2):
                return "ESLint: Trailing comma"

            # 7. TSLint: tipo any -> unknown
            if 'any' in r_line and re.sub(r'\bany\b', 'unknown', r_line) == a_line:
                return "TSLint: Reemplazar 'any' por 'unknown'"

            # 8. ESLint no-console
            if re.match(r'^console\.(log|warn|error|info|debug)\s*\(', r_line):
                return "ESLint: Eliminar console.log/warn/error"

            # 9. Prettier: template literals
            if '`' in r_line and r_line.replace('`', '"') == a_line:
                return "Prettier: Template literal a comillas dobles"
            if '"' in r_line and r_line.replace('"', '`') == a_line:
                return "Prettier: Comillas a template literal"

            # 10. ESLint prefer-arrow-callback
            m_func  = re.match(r'function\s*\(([^)]*)\)\s*\{(.+)\}', r_line)
            m_arrow = re.match(r'\(([^)]*)\)\s*=>\s*\{?(.+)\}?', a_line)
            if m_func and m_arrow and m_func.group(1).strip() == m_arrow.group(1).strip():
                return "ESLint: Convertir a arrow function"

            # 11. TSLint: eliminar modificador 'public' implicito
            if 'public' in r_line and re.sub(r'\bpublic\b\s*', '', r_line).strip() == a_line.strip():
                return "TSLint: Eliminar modificador 'public' implicito"

            # 12. ESLint: object shorthand { x: x } -> { x }
            m_long  = re.search(r'\{\s*(\w+)\s*:\s*\1\s*\}', r_line)
            m_short = re.search(r'\{\s*(\w+)\s*\}', a_line)
            if m_long and m_short and m_long.group(1) == m_short.group(1):
                return "ESLint: Object shorthand"

            # 13. Prettier: parentesis innecesarios en arrow de un parametro
            if '=>' in r_line and re.sub(r'\((\w+)\)\s*=>', r'\1 =>', r_line) == a_line:
                return "Prettier: Eliminar parentesis en arrow function de un parametro"

            # 14. VB.NET: correccion de casing de identificador (gd_Fproceso -> gd_fproceso)
            # SOLO en .vb: en otros lenguajes los cambios de casing tienen otra semántica
            if self.ext == '.vb' and r_line.lower() == a_line.lower() and r_line != a_line and not r_line.startswith("'"):
                return "Linter VB.NET: Correccion de casing de identificador"

            # 15. VB.NET / .NET: correccion de tipo estatico (file. -> File.)
            # SOLO en .vb/.cs: en Python/Go file. es un objeto valido que no debe reescribirse
            if self.ext in ('.vb', '.cs') and (
                re.sub(r'\bfile\.', 'File.', r_line) == a_line
                or re.sub(r'\bFile\.', 'file.', r_line) == a_line
            ):
                return "Linter .NET: Correccion de casing de tipo estatico (File/file)"

            # 16. VB.NET: redundancia booleana (= True / = False)
            # SOLO en .vb: Python tambien usa = True pero con semántica diferente (no es redundancia)
            if self.ext == '.vb':
                r_notrue = re.sub(r'\s*=\s*True\b', '', r_line)
                a_notrue = re.sub(r'\s*=\s*True\b', '', a_line)
                if r_notrue == a_notrue and r_notrue != r_line:
                    return "Linter VB.NET: Comparacion booleana redundante (= True eliminado)"

        return ""

    def classify_removed_line(self, line: str) -> str:
        """Clasifica la razon de eliminacion con mayor precision y explica el propósito."""
        linter = self.detect_linter_fix(line)
        if linter:
            # Añadir contexto del propósito del fix de linter
            rule_purpose = {
                "ESLint: Igualdad estricta": "para evitar comparaciones implícitas que pueden dar falsos positivos con tipos distintos",
                "ESLint: Usar const":        "para prevenir reasignaciones accidentales de variables que no deben cambiar",
                "ESLint: Usar let":          "para reemplazar 'var' con alcance de bloque más predecible",
                "Prettier: Cambio a comillas": "para unificar el estilo de cadenas en el proyecto",
                "Prettier: Arreglo de espaciado": "para mantener consistencia de formato legible",
                "ESLint: Trailing comma":    "para simplificar los diffs de git en cambios futuros",
                "TSLint: Reemplazar 'any'":  "para mejorar la seguridad de tipos y detectar errores en compilación",
                "ESLint: Eliminar console":  "para que no queden trazas de depuración en producción",
                "ESLint: Convertir a arrow": "para usar la sintaxis moderna y capturar correctamente el contexto 'this'",
                "ESLint: Object shorthand":  "para reducir redundancia en literales de objeto",
            }
            extra = ""
            for key, reason in rule_purpose.items():
                if key in linter:
                    extra = f" — {reason}"
                    break
            return f"[Linter] {linter}{extra}"
        if self.verify_dead_code(line):
            return "[Limpieza] Codigo sin uso detectado — se elimina para reducir deuda técnica y evitar confusión"

        stripped = line.strip()
        if stripped.startswith('//') or stripped.startswith('#') \
                or stripped.startswith('*') or stripped.startswith('/*'):
            return "[Doc] Comentario o documentacion eliminada"
        # VB.NET: comentario con apostrofe — SOLO para archivos .vb
        # (en PHP/Ruby/SQL la comilla simple tiene otros usos: strings, claves de array, etc.)
        if self.ext == '.vb' and stripped.startswith("'"):
            if re.search(r"'\s*(Private|Public|Protected|Friend)\s+(Sub|Function|Class)", stripped):
                return "[Refactor] Codigo VB.NET comentado eliminado (evento/rutina)"
            return "[Doc] Comentario VB.NET eliminado"
        # SQL: comentario de doble guion
        # EXCLUIR CSS/SCSS: las custom properties empiezan con -- (--primary-color: #fff)
        if stripped.startswith('--') and self.ext not in ('.css', '.scss', '.component.scss', '.sass', '.less'):
            return "[Doc] Comentario SQL eliminado"
        # VB.NET legacy: REM — SOLO para archivos .vb
        if self.ext == '.vb' and stripped.upper().startswith('REM '):
            return "[Doc] Comentario legado VB (REM) eliminado"
        if re.match(r'(console\.(log|warn|error|debug|info)|print\(|logger\.|Log\.)', stripped):
            return "[Debug] Traza o log de depuracion eliminada — evita exposición de información en producción"
        # VB.NET debug: MsgBox / Debug.Print
        if re.match(r'(MsgBox\(|Debug\.Print\b|MessageBox\.Show\()', stripped):
            return "[Debug] Cuadro de dialogo o traza de depuracion VB.NET eliminada"
        if re.search(r'\b(TODO|FIXME|HACK|XXX|TEMP)\b', stripped, re.I):
            return "[Deuda tecnica] Comentario TODO/FIXME eliminado — saldando deuda técnica pendiente"
        if re.search(r'\b(isDevMode|environment\.|process\.env|DEBUG|FEATURE_FLAG)\b', stripped, re.I):
            return "[Config] Flag de entorno o feature flag — posible limpieza de lógica de desarrollo"
        if stripped.startswith('//') and re.search(r'[({;=]', stripped):
            return "[Refactor] Codigo comentado eliminado — limpieza de código inactivo"
        if re.search(r'\b(mock|stub|fake|dummy|hardcoded|temp|test_data)\b', stripped, re.I):
            return "[Test] Dato de prueba o mock eliminado — no debe llegar a producción"
        return ""


# =============================================================================
# MOTOR DE INFERENCIA DE PROPÓSITO  (compartido por todos los motores)
# =============================================================================

def infer_modification_purpose(
    added: List[str],
    removed: List[str],
    ext: str = "",
) -> Optional[str]:
    """
    Deduce el MOTIVO / PROPÓSITO de alto nivel de un cambio comparando
    líneas añadidas y eliminadas, independientemente del lenguaje.

    Devuelve una frase en español que explica para qué se hizo la modificación,
    o None si no se puede inferir con suficiente confianza.

    Estrategia:
      - Compara la INTENCIÓN semántica del bloque eliminado vs el añadido.
      - Usa señales léxicas y estructurales para identificar el patrón.
      - Prioriza explicaciones de negocio sobre explicaciones técnicas.
    """
    added_text   = "\n".join(added)
    removed_text = "\n".join(removed)

    if not added_text.strip() and not removed_text.strip():
        return None

    # ── Patrón: corrección de bug / error ──────────────────────────────────
    bug_fix_signals_rem = [
        r'\bnull\b', r'undefined', r'TypeError', r'NullPointerException',
        r'NullReferenceException', r'Uncaught', r'Exception', r'error\b',
    ]
    bug_fix_signals_add = [
        r'\?\?', r'\?->', r'try\s*{', r'catch\s*\(', r'guard\b',
        r'if\s*\(\s*\$?\w+\s*(?:!==?|===?)\s*null',
        r'Optional\.', r'orElse\(',
    ]
    rem_bug = sum(1 for p in bug_fix_signals_rem if re.search(p, removed_text, re.I))
    add_fix = sum(1 for p in bug_fix_signals_add if re.search(p, added_text, re.I))
    if rem_bug >= 1 and add_fix >= 2:
        return ("Corrección de bug: se añaden guardas de nulidad / manejo de excepciones "
                "para evitar errores en tiempo de ejecución cuando un valor puede ser nulo o indefinido.")

    # ── Patrón: refactor de legibilidad ────────────────────────────────────
    if (removed_text.count('?') > 2 and
            re.search(r'if\s*\(', added_text) and
            not re.search(r'if\s*\(', removed_text)):
        return ("Refactor de legibilidad: expresiones ternarias anidadas se reescriben "
                "como bloques if/else explícitos para facilitar la comprensión y el mantenimiento.")

    # ── Patrón: optimización de rendimiento ────────────────────────────────
    perf_add = [r'\.cache\b', r'memoize', r'lazy', r'async\s+', r'await\s+',
                r'Promise', r'->with\s*\(', r'eager', r'index\b', r'LIMIT\b']
    perf_rem = [r'n\s*\+\s*1', r'foreach.*foreach', r'while.*query', r'SELECT \*']
    if (sum(1 for p in perf_add if re.search(p, added_text, re.I)) >= 2 or
            sum(1 for p in perf_rem if re.search(p, removed_text, re.I)) >= 1):
        return ("Optimización de rendimiento: se introduce carga diferida, caché o "
                "eager-loading para reducir la cantidad de consultas o procesamiento innecesario.")

    # ── Patrón: nueva funcionalidad ────────────────────────────────────────
    if not removed_text.strip() and added_text.strip():
        fn_match = re.search(
            r'(?:function|def|func|sub|procedure|public\s+\w+\s+\w+\s*\()\s+(\w+)',
            added_text, re.I
        )
        if fn_match:
            return f"Nueva funcionalidad: se implementa '{fn_match.group(1)}' como punto de entrada o lógica nueva."
        return "Nueva funcionalidad: se agrega código sin precedente en el archivo (no hay líneas eliminadas equivalentes)."

    # ── Patrón: eliminación de código muerto / limpieza ───────────────────
    if not added_text.strip() and removed_text.strip():
        if re.search(r'console\.|print\(|logger\.|MsgBox|Debug\.Print', removed_text, re.I):
            return "Limpieza: se eliminan trazas de depuración que no deben llegar a producción."
        return "Limpieza: se elimina código sin reemplazo (posiblemente obsoleto o sin uso detectado)."

    # ── Patrón: cambio de contrato de API / interfaz ──────────────────────
    if (re.search(r'return\s+', removed_text) and re.search(r'return\s+', added_text)):
        ret_rem = re.findall(r'return\s+(.+?)(?:;|\n)', removed_text)
        ret_add = re.findall(r'return\s+(.+?)(?:;|\n)', added_text)
        if ret_rem and ret_add and ret_rem[0].strip() != ret_add[0].strip():
            return ("Cambio de contrato: el valor de retorno de la función fue modificado, "
                    "lo que puede afectar a los consumidores de esta API o método.")

    # ── Patrón: cambio de lógica de negocio (condicional principal) ────────
    cond_rem = len(re.findall(r'\bif\s*\(', removed_text))
    cond_add = len(re.findall(r'\bif\s*\(', added_text))
    if abs(cond_add - cond_rem) >= 2:
        if cond_add > cond_rem:
            return ("Cambio de lógica de negocio: se agregan condiciones que bifurcan el flujo "
                    "para cubrir nuevos casos o escenarios de uso.")
        else:
            return ("Simplificación de lógica de negocio: se eliminan condiciones redundantes "
                    "o se consolida el flujo en una rama única más directa.")

    # ── Patrón: actualización de dependencia / importación ────────────────
    import_rem = [l for l in removed if re.match(r'\s*(import|use|require|from|using)\b', l)]
    import_add = [l for l in added   if re.match(r'\s*(import|use|require|from|using)\b', l)]
    if import_rem and import_add and not (
        set(l.strip() for l in import_rem) == set(l.strip() for l in import_add)
    ):
        return ("Actualización de dependencias: se cambian los módulos o clases importadas, "
                "posiblemente para usar una versión actualizada o un reemplazo de la dependencia.")

    # ── Patrón: seguridad / autenticación ─────────────────────────────────
    sec_add = [r'auth\b', r'token\b', r'jwt\b', r'middleware', r'guard\b',
               r'authorize', r'permission', r'role\b', r'sanitize', r'escape\b']
    if sum(1 for p in sec_add if re.search(p, added_text, re.I)) >= 2:
        return ("Mejora de seguridad: se refuerzan controles de autenticación, "
                "autorización o sanitización de datos de entrada.")

    # ── Patrón: internacionalización / localización ────────────────────────
    if re.search(r'translate\.|i18n\.|locale\b|__\(', added_text, re.I):
        return ("Internacionalización: se incorporan traducciones o literales internacionalizados "
                "para soportar múltiples idiomas o regiones.")

    # ── Patrón: corrección de estilo / linter ─────────────────────────────
    r_ns = re.sub(r'\s+', '', removed_text)
    a_ns = re.sub(r'\s+', '', added_text)
    similarity = len(r_ns) > 0 and len(a_ns) > 0 and (
        abs(len(r_ns) - len(a_ns)) / max(len(r_ns), len(a_ns)) < 0.15
    )
    if similarity and removed_text != added_text:
        return ("Corrección de estilo/linter: los cambios son mínimos en contenido lógico "
                "pero ajustan formato, espaciado, comillas o convenciones del equipo (ESLint/Prettier).")

    return None


# =============================================================================
# MOTOR DE ANÁLISIS LÓGICO PHP
# =============================================================================
def analyze_php_logic_changes(
    added: List[str],
    removed: List[str],
    added_wl: List[Tuple[Optional[int], str]],
    removed_wl: List[Tuple[Optional[int], str]],
) -> List[str]:
    """
    Motor heurístico de análisis lógico para archivos PHP / Laravel.

    Detecta múltiples escenarios e INFIERE EL PROPÓSITO de cada cambio:
      1.  Descomposición de ternarios en bloques if/else explícitos.
      2.  Refactor de acceso directo a relación → acceso con null-safe.
      3.  Introducción / eliminación de variables intermedias.
      4.  Cambio de lógica de determinación de campos (tienda, lugar, motivo…).
      5.  Adición / eliminación de guard de nulidad (?? / ?->).
      6.  Cambio en estructura de array de retorno / respuesta.
      7.  Extracción de lógica inline a bloque condicional separado.
      8.  Adición de comentarios descriptivos de sección.
      9.  Refactor de acceso a propiedad directa → método relacional.
     10.  Cambio de consulta Eloquent (scope, método, relaciones eager).
     11.  Adición / eliminación de validación de petición.
     12.  Cambio en respuesta HTTP (campos, estructura JSON).
     13.  Cambio en definición de función / firma de método.
     14.  Eliminación de acceso directo a propiedad sin guard.
     15.  Descomposición de expresión compleja en pasos intermedios.
     16+. OAuth2/Keycloak, JWT, Guzzle, duplicidades, claims.
     NEW: Propósito general inferido al inicio de cada análisis.
    """
    insights: List[str] = []

    added_text   = "\n".join(added)
    removed_text = "\n".join(removed)

    # ── Propósito general del cambio (inferencia de alto nivel) ─────────────
    purpose = infer_modification_purpose(added, removed, ext='.php')
    if purpose:
        insights.append(f"🎯 Propósito del cambio: {purpose}")

    # ── Helpers ─────────────────────────────────────────────────────────────

    def _extract_ternary_field(line: str) -> Optional[str]:
        """Extrae el nombre del campo de un ternario tipo 'campo' => expr ? a : b."""
        m = re.search(r"['\"](\w+)['\"]\s*=>\s*\$\w+.*\?", line)
        if m:
            return m.group(1)
        m2 = re.search(r'\$(\w+)\s*=\s*\$\w+.*\?', line)
        if m2:
            return m2.group(1)
        return None

    def _extract_var_assignments(lines: List[str]) -> Dict[str, str]:
        """Devuelve {$var: valor} para asignaciones simples."""
        result: Dict[str, str] = {}
        for ln in lines:
            m = re.match(r"\s*\$(\w+)\s*=\s*(.+?)\s*;?\s*$", ln)
            if m:
                result[f"${m.group(1)}"] = m.group(2).strip()
        return result

    def _count_ternaries(lines: List[str]) -> int:
        return sum(1 for ln in lines if '?' in ln and ':' in ln
                   and not ln.strip().startswith('//') and '?->' not in ln)

    def _count_if_blocks(lines: List[str]) -> int:
        return sum(1 for ln in lines if re.search(r'\bif\s*\(', ln))

    def _null_safe_count(lines: List[str]) -> int:
        return sum(1 for ln in lines if '??' in ln or '?->' in ln)

    def _find_changed_array_keys(added_l: List[str], removed_l: List[str]) -> Tuple[List[str], List[str]]:
        """Detecta claves de array PHP que aparecen solo en añadidos o solo en eliminados."""
        key_pat = re.compile(r"['\"](\w+)['\"]\s*=>")
        keys_add = {m.group(1) for ln in added_l for m in key_pat.finditer(ln)}
        keys_rem = {m.group(1) for ln in removed_l for m in key_pat.finditer(ln)}
        new_keys  = sorted(keys_add - keys_rem)
        lost_keys = sorted(keys_rem - keys_add)
        return new_keys, lost_keys

    def _find_new_variables(added_l: List[str], removed_l: List[str]) -> List[str]:
        var_pat = re.compile(r"\$(\w+)\s*=")
        vars_add = {m.group(1) for ln in added_l for m in var_pat.finditer(ln)}
        vars_rem = {m.group(1) for ln in removed_l for m in var_pat.finditer(ln)}
        return sorted(vars_add - vars_rem)

    def _find_removed_variables(added_l: List[str], removed_l: List[str]) -> List[str]:
        var_pat = re.compile(r"\$(\w+)\s*=")
        vars_add = {m.group(1) for ln in added_l for m in var_pat.finditer(ln)}
        vars_rem = {m.group(1) for ln in removed_l for m in var_pat.finditer(ln)}
        return sorted(vars_rem - vars_add)

    def _extract_section_comments(lines: List[str]) -> List[str]:
        """Extrae comentarios descriptivos de sección (// Determinar …)."""
        return [
            ln.strip().lstrip('/').strip()
            for ln in lines
            if re.match(r'\s*//\s*\w', ln) and len(ln.strip()) > 4
        ]

    # ── Análisis 1: Descomposición de ternarios en if/else ──────────────────
    tern_rem = _count_ternaries(removed)
    tern_add = _count_ternaries(added)
    if_add    = _count_if_blocks(added)
    if_rem    = _count_if_blocks(removed)

    if tern_rem > 0 and tern_add < tern_rem and if_add > if_rem:
        # Hay ternarios que desaparecen y aparecen bloques if nuevos
        # Intentar identificar qué campos se descompusieron
        fields: List[str] = []
        for ln in removed:
            f = _extract_ternary_field(ln)
            if f and f not in fields:
                fields.append(f)
        if fields:
            flds = ", ".join(f"'{f}'" for f in fields[:4])
            insights.append(
                f"Se descomponen ternarios en bloques if/else explícitos para los campos {flds}: "
                "permite manejar casos adicionales, mejorar legibilidad y agregar lógica intermedia."
            )
        else:
            insights.append(
                f"Se reemplazan {tern_rem} expresión(es) ternaria(s) por bloques if/else explícitos, "
                "incrementando la expresividad de la lógica condicional."
            )

    # ── Análisis 2: Variables intermedias nuevas ─────────────────────────────
    new_vars = _find_new_variables(added, removed)
    rem_vars  = _find_removed_variables(added, removed)

    if new_vars and len(new_vars) <= 6:
        vnames = ", ".join(f"${v}" for v in new_vars[:5])
        insights.append(
            f"Se introducen variables intermedias ({vnames}) para encapsular "
            "sub-lógicas antes calculadas inline, mejorando legibilidad y mantenibilidad."
        )
    if rem_vars and len(rem_vars) <= 6:
        vnames = ", ".join(f"${v}" for v in rem_vars[:5])
        insights.append(
            f"Se eliminan variables intermedias ({vnames}); "
            "la lógica asociada fue simplificada o consolidada en otra expresión."
        )

    # ── Análisis 3: Cambio en claves del array de respuesta/datos ────────────
    new_keys, lost_keys = _find_changed_array_keys(added, removed)
    if new_keys:
        insights.append(
            f"Se añaden campos al array de datos/respuesta: {', '.join(repr(k) for k in new_keys[:6])}."
        )
    if lost_keys:
        insights.append(
            f"Se eliminan campos del array de datos/respuesta: {', '.join(repr(k) for k in lost_keys[:6])}."
        )

    # ── Análisis 4: Guard de nulidad / null-safe ────────────────────────────
    ns_add = _null_safe_count(added)
    ns_rem = _null_safe_count(removed)
    if ns_add > ns_rem + 1:
        insights.append(
            f"Se agregan {ns_add - ns_rem} guard(s) de nulidad (?? / ?->) "
            "para prevenir errores al acceder a relaciones o propiedades opcionales."
        )
    if ns_rem > ns_add + 1:
        insights.append(
            f"Se eliminan {ns_rem - ns_add} guard(s) de nulidad previos; "
            "verificar que las relaciones referenciadas siempre existan en este contexto."
        )

    # ── Análisis 5: Acceso directo → guard de nulidad en campos de relación ──
    # Patrón: ->'campo'->propiedad reemplazado por ?->propiedad o ?? ''
    direct_access = re.findall(r'\$\w+->(\w+)->(\w+)', removed_text)
    null_safe      = re.findall(r'\$\w+\?->(\w+)', added_text)
    if direct_access and null_safe:
        rels = list({f[0] for f in direct_access})[:3]
        insights.append(
            f"Se corrige acceso directo a relación sin guard en "
            f"{', '.join(rels)}: se agrega ?-> o ?? para evitar errores si la relación es null."
        )

    # ── Análisis 6: Lógica de determinación de campo basada en condición ─────
    # Detecta el patrón: variable se asigna en bloque if/else según un flag (is_app, etc.)
    condition_flags = re.findall(r'\bif\s*\(\s*\$\w+->(\w+)\s*(?:\&\&|\|\|)?\s*\$?\w*\s*\)', added_text)
    if condition_flags:
        flags = list(dict.fromkeys(condition_flags))[:3]
        insights.append(
            f"La determinación de campos ahora depende de condición(es) sobre: "
            f"{', '.join(flags)}; se bifurca la lógica según estado del objeto."
        )

    # ── Análisis 7: Comentarios de sección añadidos ──────────────────────────
    new_comments = _extract_section_comments(added)
    if new_comments and len(new_comments) >= 2:
        topics = "; ".join(new_comments[:3])
        insights.append(
            f"Se documentan secciones lógicas con comentarios descriptivos: «{topics}»."
        )

    # ── Análisis 8: Cambio en consulta Eloquent / relaciones eager ────────────
    with_added   = re.findall(r'->with\s*\(\s*[\'"](\w+)[\'"]', added_text)
    with_removed = re.findall(r'->with\s*\(\s*[\'"](\w+)[\'"]', removed_text)
    new_eager = set(with_added) - set(with_removed)
    rem_eager = set(with_removed) - set(with_added)
    if new_eager:
        insights.append(
            f"Se agregan relaciones eager-load ({', '.join(sorted(new_eager))}) "
            "a la consulta Eloquent, optimizando la carga de datos relacionados."
        )
    if rem_eager:
        insights.append(
            f"Se eliminan relaciones eager-load ({', '.join(sorted(rem_eager))}) "
            "de la consulta; revisar si puede generarse problema N+1."
        )

    # ── Análisis 9: Métodos de consulta Eloquent modificados ─────────────────
    query_methods_add = set(re.findall(r'->(\w+)\s*\(', added_text)) & {
        'where', 'whereIn', 'whereBetween', 'orWhere', 'having',
        'orderBy', 'groupBy', 'limit', 'skip', 'select', 'join',
        'leftJoin', 'rightJoin', 'whereHas', 'doesntHave',
    }
    query_methods_rem = set(re.findall(r'->(\w+)\s*\(', removed_text)) & {
        'where', 'whereIn', 'whereBetween', 'orWhere', 'having',
        'orderBy', 'groupBy', 'limit', 'skip', 'select', 'join',
        'leftJoin', 'rightJoin', 'whereHas', 'doesntHave',
    }
    new_qm = sorted(query_methods_add - query_methods_rem)
    rem_qm = sorted(query_methods_rem - query_methods_add)
    if new_qm:
        insights.append(
            f"Se añaden métodos de filtrado/consulta Eloquent: {', '.join(new_qm)}."
        )
    if rem_qm:
        insights.append(
            f"Se eliminan métodos de filtrado/consulta Eloquent: {', '.join(rem_qm)}."
        )

    # ── Análisis 10: Cambio en firma de función PHP ───────────────────────────
    fn_add = re.findall(r'(?:public|private|protected|static)?\s*function\s+(\w+)\s*\(([^)]*)\)', added_text)
    fn_rem = re.findall(r'(?:public|private|protected|static)?\s*function\s+(\w+)\s*\(([^)]*)\)', removed_text)
    fn_add_dict = {name: params for name, params in fn_add}
    fn_rem_dict = {name: params for name, params in fn_rem}
    for fname in set(fn_add_dict) & set(fn_rem_dict):
        if fn_add_dict[fname].strip() != fn_rem_dict[fname].strip():
            insights.append(
                f"Se modifica la firma de la función '{fname}': "
                f"parámetros cambiaron de ({fn_rem_dict[fname].strip()}) "
                f"a ({fn_add_dict[fname].strip()})."
            )

    # ── Análisis 11: Respuesta HTTP / JSON ────────────────────────────────────
    resp_add = re.search(r'response\s*\(|->json\s*\(|ApiResponse|JsonResponse', added_text, re.I)
    resp_rem = re.search(r'response\s*\(|->json\s*\(|ApiResponse|JsonResponse', removed_text, re.I)
    if resp_add and not resp_rem:
        insights.append(
            "Se introduce respuesta HTTP estructurada (response/json), "
            "posiblemente convirtiendo un retorno directo en respuesta de API."
        )
    elif resp_rem and not resp_add:
        insights.append(
            "Se elimina respuesta HTTP estructurada previa; "
            "la función puede haber cambiado su contrato de retorno."
        )

    # ── Análisis 12: Validación de request ────────────────────────────────────
    val_add = re.search(r'\$request->validate\s*\(|->validated\s*\(|->rules\s*\(', added_text)
    val_rem = re.search(r'\$request->validate\s*\(|->validated\s*\(|->rules\s*\(', removed_text)
    if val_add and not val_rem:
        insights.append(
            "Se agrega validación de datos de la petición HTTP, "
            "reforzando la integridad de entrada antes de procesar."
        )
    elif val_rem and not val_add:
        insights.append(
            "Se elimina validación explícita de la petición; "
            "verificar que la integridad de datos quede cubierta por otro mecanismo."
        )

    # ── Análisis 13: Cambio en Log / trazabilidad ─────────────────────────────
    log_add = re.findall(r'Log::\w+\s*\(|logger\(\)', added_text)
    log_rem = re.findall(r'Log::\w+\s*\(|logger\(\)', removed_text)
    if len(log_add) > len(log_rem):
        insights.append(
            f"Se agregan {len(log_add) - len(log_rem)} traza(s) de log, "
            "mejorando la observabilidad del flujo."
        )
    elif len(log_rem) > len(log_add):
        insights.append(
            f"Se eliminan {len(log_rem) - len(log_add)} traza(s) de log previas."
        )

    # ── Análisis 14: Middleware / autorización ────────────────────────────────
    mw_add = re.search(r'->middleware\s*\(|Gate::|@can|can\(', added_text, re.I)
    mw_rem = re.search(r'->middleware\s*\(|Gate::|@can|can\(', removed_text, re.I)
    if mw_add and not mw_rem:
        insights.append(
            "Se introduce restricción de middleware/autorización, "
            "agregando control de acceso al flujo."
        )
    elif mw_rem and not mw_add:
        insights.append(
            "Se elimina middleware/autorización previa; "
            "verificar que el acceso a este endpoint siga protegido."
        )

    # ── Análisis 15: Flujo OAuth2 / Keycloak eliminado o añadido ─────────────
    keycloak_signals = [
        r"config\s*\(\s*['\"]services\.keycloak",
        r"grant_type.*password|grant_type.*client_credentials",
        r"preferred_username|realm.*protocol.*openid-connect",
        r"client_secret.*keycloak|keycloak.*client_id",
    ]
    kc_rem = sum(1 for pat in keycloak_signals if re.search(pat, removed_text, re.I))
    kc_add = sum(1 for pat in keycloak_signals if re.search(pat, added_text, re.I))
    if kc_rem >= 2 and kc_add == 0:
        # Detectar nombre del método donde ocurrió, si es posible
        fn_ctx = re.search(
            r'(?:public|private|protected)\s+function\s+(\w+)', removed_text
        )
        fn_name = f" en '{fn_ctx.group(1)}'" if fn_ctx else ""
        insights.append(
            f"Se elimina flujo completo de autenticación OAuth2/Keycloak (grant_type=password){fn_name}: "
            "se removió la obtención de token de acceso delegado, la extracción del claim "
            "'preferred_username' y la validación de error del token. "
            "Verificar que un mecanismo alternativo (middleware, helper de Token, etc.) "
            "cubra la autenticación del usuario."
        )
    elif kc_add >= 2 and kc_rem == 0:
        insights.append(
            "Se introduce flujo de autenticación OAuth2/Keycloak (grant_type=password): "
            "el método ahora delega la verificación de credenciales al servidor de identidad Keycloak."
        )

    # ── Análisis 16: Procesamiento manual de JWT (decode de payload) ──────────
    jwt_pipeline_rem = (
        re.search(r"explode\s*\(\s*'\.'", removed_text) and
        re.search(r'base64_decode\s*\(\s*\$\w+\s*\[', removed_text) and
        re.search(r'json_decode\s*\(\s*\$\w+\s*,\s*true\s*\)', removed_text)
    )
    jwt_pipeline_add = (
        re.search(r"explode\s*\(\s*'\.'", added_text) and
        re.search(r'base64_decode\s*\(\s*\$\w+\s*\[', added_text) and
        re.search(r'json_decode\s*\(\s*\$\w+\s*,\s*true\s*\)', added_text)
    )
    if jwt_pipeline_rem and not jwt_pipeline_add:
        # Contar cuántas veces se repite el patrón (duplicidad)
        jwt_count = len(re.findall(r'base64_decode\s*\(\s*\$\w+\s*\[', removed_text))
        dup_note = (
            f" El pipeline apareció {jwt_count} veces en el código eliminado, "
            "indicando duplicidad previa del bloque de decodificación."
            if jwt_count > 1 else ""
        )
        insights.append(
            "Se elimina procesamiento manual del token JWT (explode('.', $token), "
            "base64_decode($token[1]), json_decode): el sistema dejaba de delegar "
            "el parse del payload al middleware y lo hacía directamente en el método." + dup_note
        )
    elif jwt_pipeline_add and not jwt_pipeline_rem:
        insights.append(
            "Se introduce procesamiento manual del token JWT (explode/base64_decode/json_decode) "
            "directamente en el método: considerar centralizar en un helper de Token."
        )

    # ── Análisis 17: Cliente HTTP Guzzle eliminado o añadido ──────────────────
    guzzle_rem = re.search(r'new\s+Client\s*\(\s*\)|\$client\s*->\s*(?:post|get|put|delete)\s*\(', removed_text)
    guzzle_add = re.search(r'new\s+Client\s*\(\s*\)|\$client\s*->\s*(?:post|get|put|delete)\s*\(', added_text)
    if guzzle_rem and not guzzle_add:
        # Intentar extraer el host/servicio al que se llamaba
        url_hint = re.search(r"config\s*\(\s*['\"]([^'\"]+)['\"]", removed_text)
        service_hint = f" (endpoint: {url_hint.group(1)})" if url_hint else ""
        insights.append(
            f"Se elimina llamada HTTP directa mediante Guzzle (new Client()){service_hint}: "
            "el método ya no realiza peticiones salientes síncronas; "
            "verificar si la integración fue reemplazada por un servicio/helper centralizado."
        )
    elif guzzle_add and not guzzle_rem:
        url_hint = re.search(r"config\s*\(\s*['\"]([^'\"]+)['\"]", added_text)
        service_hint = f" (endpoint: {url_hint.group(1)})" if url_hint else ""
        insights.append(
            f"Se introduce llamada HTTP directa mediante Guzzle (new Client()){service_hint}: "
            "el método ahora realiza una petición saliente síncrona a un servicio externo."
        )

    # ── Análisis 18: Bloque de código duplicado eliminado ─────────────────────
    if removed:
        from collections import Counter
        line_counts = Counter(ln.strip() for ln in removed if len(ln.strip()) > 15)
        duplicated = [(ln, cnt) for ln, cnt in line_counts.items() if cnt >= 2]
        if duplicated:
            # Identificar tema del duplicado (token, query, assignación...)
            dup_sample = duplicated[0][0]
            if re.search(r'base64_decode|json_decode|explode.*token', dup_sample, re.I):
                dup_topic = "decodificación de token JWT"
            elif re.search(r'\$\w+->\w+\s*\(', dup_sample):
                dup_topic = "llamada a método/relación"
            elif re.search(r'\$\w+\s*=\s*.+;', dup_sample):
                dup_topic = "asignación de variable"
            else:
                dup_topic = "bloque de lógica"
            insights.append(
                f"Se detecta y elimina duplicidad de código: el bloque de {dup_topic} "
                f"aparecía {duplicated[0][1]} veces ({len(duplicated)} línea(s) repetida(s)). "
                "La eliminación unifica el flujo evitando inconsistencias por mantenimiento paralelo."
            )

    # ── Análisis 19: Validación de error en token OAuth2 / claim assertion ────
    token_err_rem = re.search(
        r"isset\s*\(\s*\$\w+\s*\[\s*['\"]error['\"]\s*\]\s*\)"
        r"|array_key_exists\s*\(\s*['\"]error['\"]"
        r"|\$\w+\s*\[\s*['\"]error_description['\"]\s*\]",
        removed_text
    )
    token_err_add = re.search(
        r"isset\s*\(\s*\$\w+\s*\[\s*['\"]error['\"]\s*\]\s*\)"
        r"|array_key_exists\s*\(\s*['\"]error['\"]",
        added_text
    )
    if token_err_rem and not token_err_add:
        insights.append(
            "Se elimina validación explícita de error en respuesta OAuth2 "
            "(isset($token['error'])): el flujo ya no verifica manualmente si el servidor "
            "de identidad retornó un error; asegurar que el middleware o try/catch "
            "cubra este caso de fallo de autenticación."
        )

    # ── Análisis 20: Extracción de claim de usuario del payload JWT ───────────
    claim_rem = re.findall(
        r"\$\w+\s*\[\s*['\"]"
        r"(preferred_username|sub|email|given_name|family_name|realm_access|resource_access)"
        r"['\"]\s*\]",
        removed_text
    )
    claim_add = re.findall(
        r"\$\w+\s*\[\s*['\"]"
        r"(preferred_username|sub|email|given_name|family_name|realm_access|resource_access)"
        r"['\"]\s*\]",
        added_text
    )
    lost_claims = sorted(set(claim_rem) - set(claim_add))
    new_claims  = sorted(set(claim_add) - set(claim_rem))
    if lost_claims:
        insights.append(
            f"Se elimina uso de claim(s) JWT: {', '.join(repr(c) for c in lost_claims)}. "
            "El método ya no extrae estos atributos del payload del token; "
            "verificar que la identidad del usuario sea obtenida por otro medio."
        )
    if new_claims:
        insights.append(
            f"Se introduce lectura de claim(s) JWT: {', '.join(repr(c) for c in new_claims)}. "
            "El método ahora extrae estos atributos del payload del token de acceso."
        )

    return insights


# =============================================================================
# PARSER DE DIFF
# =============================================================================
class DiffParser:
    _FILE   = re.compile(r'^diff --git a/(.+?) b/(.+)$')
    _NEW    = re.compile(r'^new file mode')
    _DEL    = re.compile(r'^deleted file mode')
    _REN    = re.compile(r'^rename to (.+)$')
    _BIN    = re.compile(r'^Binary files.*differ$')
    _HUNK   = re.compile(r'^@@ -(\d+)(?:,(\d+))? \+(\d+)(?:,(\d+))? @@\s*(.*)$')
    _PLUS3  = re.compile(r'^\+\+\+')
    _MIN3   = re.compile(r'^---')

    def parse(self, text: str) -> List[FileChange]:
        text = clean(text)
        files: List[FileChange] = []
        cur: Optional[FileChange] = None
        old_line_no: Optional[int] = None
        new_line_no: Optional[int] = None

        for line in text.splitlines():
            m_file = self._FILE.match(line)
            if m_file:
                cur = FileChange(filepath=m_file.group(2).strip())
                files.append(cur)
                old_line_no = None
                new_line_no = None
                continue
            if cur is None:
                continue
            if self._NEW.match(line):
                cur.kind = 'added'
            elif self._DEL.match(line):
                cur.kind = 'deleted'
            elif self._BIN.match(line):
                cur.is_binary = True
            elif self._REN.match(line):
                cur.kind = 'renamed'
                cur.filepath = self._REN.match(line).group(1).strip()
            elif m_hunk := self._HUNK.match(line):
                old_line_no = int(m_hunk.group(1))
                new_line_no = int(m_hunk.group(3))
                hint = m_hunk.group(5).strip()
                if hint and len(hint) > 2:
                    cur.contexts.add(hint[:60])
            elif self._PLUS3.match(line) or self._MIN3.match(line):
                continue
            elif line.startswith('+') and not line.startswith('+++'):
                s = line[1:].strip()
                if s:
                    cur.added.append(s)
                    cur.added_with_line.append((new_line_no, s))
                    cur.full_content.append(s)
                if new_line_no is not None:
                    new_line_no += 1
            elif line.startswith('-') and not line.startswith('---'):
                s = line[1:].strip()
                if s:
                    cur.removed.append(s)
                    cur.removed_with_line.append((old_line_no, s))
                if old_line_no is not None:
                    old_line_no += 1
            elif line.startswith(' '):
                s = line[1:].strip()
                if s:
                    cur.full_content.append(s)
                if old_line_no is not None:
                    old_line_no += 1
                if new_line_no is not None:
                    new_line_no += 1

        return [f for f in files if f.filepath]


# =============================================================================
# CATEGORIAS DE ARCHIVOS
# =============================================================================
FILE_CATEGORIES: Dict[str, str] = {
    '.component.html':  'Template Angular',
    '.component.ts':    'Componente Angular',
    '.component.scss':  'Estilos Componente',
    '.service.ts':      'Servicio Angular',
    '.spec.ts':         'Test Unitario',
    '.module.ts':       'Modulo Angular',
    '.pipe.ts':         'Pipe Angular',
    '.directive.ts':    'Directiva Angular',
    '.guard.ts':        'Guard de Ruta',
    '.interceptor.ts':  'Interceptor HTTP',
    '.reducer.ts':      'Reducer NgRx',
    '.action.ts':       'Accion NgRx',
    '.effect.ts':       'Efecto NgRx',
    '.selector.ts':     'Selector NgRx',
    '.resolver.ts':     'Resolver Angular',
    '.model.ts':        'Modelo de Datos',
    '.interface.ts':    'Interfaz TypeScript',
    '.enum.ts':         'Enumeracion',
    '.helper.ts':       'Helper/Utilidad',
    '.util.ts':         'Utilidad',
    '.config.ts':       'Configuracion',
    '.constant.ts':     'Constantes',
    '.html':            'Template HTML',
    '.scss':            'Estilos SCSS',
    '.css':             'Estilos CSS',
    '.ts':              'TypeScript',
    '.js':              'JavaScript',
    '.py':              'Python',
    '.json':            'Configuracion JSON',
    '.md':              'Documentacion',
    '.sql':             'Base de Datos',
    '.yml':             'Config YAML',
    '.yaml':            'Config YAML',
    '.env':             'Variables de Entorno',
    '.sh':              'Script Shell',
    '.xml':             'XML / Config',
    '.graphql':         'Esquema GraphQL',
    '.prisma':          'Schema Prisma ORM',
    # PHP / Laravel
    '.php':             'PHP Backend',
    '.blade.php':       'Vista Blade PHP',
    # VB.NET / Visual Studio
    '.vb':              'Formulario / Clase VB.NET',
    '.vbproj':          'Proyecto VB.NET',
    '.sln':             'Solucion Visual Studio',
    '.bak':             'Archivo de Respaldo',
    # C# / .NET
    '.cs':              'Clase C#',
    '.csproj':          'Proyecto C#',
    '.razor':           'Componente Blazor',
    '.cshtml':          'Vista Razor MVC',
    '.resx':            'Recursos .NET',
    '.xaml':            'Interfaz XAML',
    '.config':          'Configuracion App (.NET)',
    # Java / Kotlin
    '.java':            'Clase Java',
    '.kt':              'Clase Kotlin',
    # Go / Rust / Ruby
    '.go':              'Archivo Go',
    '.rs':              'Archivo Rust',
    '.rb':              'Clase Ruby',
    # Frontend
    '.vue':             'Componente Vue',
    '.jsx':             'Componente React JSX',
    '.tsx':             'Componente React TSX',
    '.svelte':          'Componente Svelte',
    # Mobile
    '.dart':            'Clase Dart/Flutter',
    '.swift':           'Codigo Swift',
    # C / C++
    '.cpp':             'Clase C++',
    '.cc':              'Clase C++',
    '.h':               'Header C/C++',
    '.hpp':             'Header C++',
    # Infraestructura
    '.tf':              'Infraestructura Terraform',
    '.dockerfile':      'Dockerfile',
    # Otros
    '.toml':            'Configuracion TOML',
    '.ini':             'Configuracion INI',
    '.proto':           'Definicion gRPC',
    '.r':               'Script R',
    '.ipynb':           'Notebook Jupyter',
    '.editorconfig':    'Configuracion Editor',
    '.gitignore':       'Git Ignore',
}

def get_category(fc: FileChange) -> str:
    if fc.is_lockfile:
        return 'Dependencias (Lock)'
    if fc.is_environment_file:
        return 'Archivo de Entorno'
    return FILE_CATEGORIES.get(fc.ext, Path(fc.filepath).suffix.upper() or 'Archivo')


# =============================================================================
# SENALES DE IMPACTO TECNICO
# =============================================================================
IMPACT_SIGNALS: List[Tuple[re.Pattern, str]] = [
    (re.compile(r'\.subscribe\s*\('),
     'Flujo asincrono RxJS'),
    (re.compile(r'catchError|throwError|try\s*{|except\s+|\.catch\s*\('),
     'Gestion de errores y excepciones'),
    (re.compile(r'\bAuthService\b|\bAuthGuard\b|inject\s*\(.*Auth|new\s+Auth\w+'),
     'Seguridad: Servicio de autenticacion'),
    (re.compile(r'\bJWT\b|jsonwebtoken|\.sign\s*\(|\.verify\s*\(|bearerToken|refreshToken'),
     'Seguridad: Manejo de JWT/tokens'),
    (re.compile(r'(?<!\w)password(?!\w)|\bhash\b|\bbcrypt\b|\bsalt\b|\bcipher\b', re.I),
     'Seguridad: Credenciales o hashes'),
    (re.compile(r'apiLoginUrl|apiBackend|apiUrl\b|API_URL|baseUrl|apiSIAW|apiSIP', re.I),
     'Variable de entorno o endpoint de API'),
    (re.compile(r'router\.navigate|HttpResponse|location\.href|\[routerLink\]'),
     'Logica de ruteo o navegacion'),
    (re.compile(r'HttpClient|http\.(get|post|put|delete|patch)\s*[(<]', re.I),
     'Llamada HTTP a API externa'),
    (re.compile(r'console\.(log|warn|error)|print\(|logger\.'),
     'Trazabilidad (logs de depuracion)'),
    (re.compile(r'SELECT\b|INSERT\b|UPDATE\b|DELETE\b|JOIN\b|\.save\s*\(|\.find\s*\(', re.I),
     'Consulta u operacion de base de datos'),
    (re.compile(r'\bexport\s+(class|interface|type|enum)\b'),
     'Contrato exportado publicamente'),
    (re.compile(r'@Input\s*\(\)|@Output\s*\(\)|EventEmitter|@Prop\s*\(\)'),
     'Contrato de componente (inputs/outputs)'),
    (re.compile(r'@NgModule\s*\(|providers\s*:\s*\[|declarations\s*:\s*\['),
     'Configuracion de modulo Angular'),
    (re.compile(r'environment\.(prod|production|cloud|staging|qa)', re.I),
     'Configuracion de ambiente especifico'),
    (re.compile(r'dispatch\s*\(|store\.select|createAction|createReducer|createEffect'),
     'Estado global NgRx'),
    (re.compile(r'migration|schema|alembic|flyway|liquibase|ALTER TABLE|DROP TABLE', re.I),
     'Migracion de base de datos'),
    (re.compile(r'cron|@Scheduled|scheduler|setInterval|setTimeout'),
     'Tarea programada/scheduler'),
    (re.compile(r'WebSocket|socket\.io|ws://|wss://'),
     'Comunicacion en tiempo real (WebSocket)'),
    (re.compile(r'\bcache\b|redis|memcache|localStorage|sessionStorage', re.I),
     'Capa de cache o almacenamiento'),
    (re.compile(r'\bi18n\b|translate\.|locale|l10n', re.I),
     'Internacionalizacion (i18n)'),
    (re.compile(r'canActivate|canLoad|canMatch|menuGuard|authGuard'),
     'Guards de ruta (control de acceso)'),
    # --- PHP / Laravel ---
    (re.compile(r'public\s+function\s+\w+|private\s+function\s+\w+|protected\s+function\s+\w+'),
     'Modificacion de metodo de clase PHP'),
    (re.compile(r'\?\?|\?\?=|\?->|\$\w+\s*\?\s*\$\w+->|\$\w+\s*\?\s*\$\w+\s*:\s*', re.I),
     'Guard de nulidad / null-coalescing en logica PHP'),
    (re.compile(r'->where\s*\(|->select\s*\(|->with\s*\(|->find\s*\(|->create\s*\(|->update\s*\(|->save\s*\(|->delete\s*\(', re.I),
     'Operacion Eloquent ORM (consulta/persistencia)'),
    (re.compile(r'Route::(get|post|put|delete|patch|resource|apiResource)\s*\(', re.I),
     'Definicion de ruta de API/web Laravel'),
    (re.compile(r'->middleware\s*\(|\bMiddleware\b', re.I),
     'Middleware de autenticacion/autorizacion PHP'),
    (re.compile(r'return\s+response\s*\(|->json\s*\(|->response\s*\(', re.I),
     'Respuesta de endpoint HTTP'),
    (re.compile(r'->belongsTo\s*\(|->hasMany\s*\(|->hasOne\s*\(|->belongsToMany\s*\(|->morphMany\s*\('),
     'Relacion Eloquent entre modelos'),
    (re.compile(r'\$request->|Request\s+\$\w+|->validated\s*\(|->rules\s*\('),
     'Procesamiento/validacion de peticion HTTP'),
    (re.compile(r'\bLog::|\$this->error\b|\$this->success\b|ApiResponse', re.I),
     'Trazabilidad / respuesta estandarizada de API PHP'),
    (re.compile(r'->paginate\s*\(|->get\s*\(|->first\s*\(|->count\s*\(', re.I),
     'Consulta de coleccion Eloquent'),
    # --- PHP: Autenticación OAuth2 / Keycloak ---
    (re.compile(r"config\s*\(\s*['\"]services\.keycloak", re.I),
     'Flujo de autenticacion Keycloak (OAuth2)'),
    (re.compile(r"grant_type.*password|grant_type.*client_credentials", re.I),
     'OAuth2: credenciales de usuario (password grant)'),
    (re.compile(r"preferred_username|realm_access|resource_access", re.I),
     'Claim JWT de identidad de usuario (Keycloak)'),
    (re.compile(r'base64_decode\s*\(\s*\$\w+\s*\[.*\]\s*\)', re.I),
     'Decodificacion manual de payload JWT (base64_decode)'),
    (re.compile(r"explode\s*\(\s*'\.'\s*,\s*\$\w+", re.I),
     'Separacion de secciones de token JWT (explode)'),
    (re.compile(r'new\s+Client\s*\(\s*\)|\bGuzzleHttp\\Client\b', re.I),
     'Cliente HTTP Guzzle (peticion saliente sincrona)'),
    (re.compile(r"isset\s*\(\s*\$\w+\s*\[\s*['\"]error['\"]\s*\]\s*\)", re.I),
     'Verificacion de error en respuesta OAuth2/token'),
    # --- VB.NET / Windows Forms ---
    (re.compile(r'\bRealizarConsulta\s*\(|\bEjecutarSP\s*\(|\bCALL\s+usp_', re.I),
     'Llamada a procedimiento almacenado VB.NET'),
    (re.compile(r'\bDataTable\b|\bDataRow\b|\bDataSet\b|\bDataAdapter\b'),
     'Manejo de datos (DataTable/DataRow) VB.NET'),
    (re.compile(r'\bMsgBox\s*\(|\bMessageBox\.Show\s*\(', re.I),
     'Cuadro de dialogo (MsgBox) VB.NET'),
    (re.compile(r'\bTry\b[\s\S]*?\bCatch\s+ex\s+As\s+Exception\b', re.I),
     'Gestion de errores Try/Catch VB.NET'),
    (re.compile(r'\.ShowDialog\s*\(|\.Show\s*\(|Me\.Close\s*\(|Me\.Dispose\s*\(', re.I),
     'Control de ciclo de vida de formulario VB.NET'),
    (re.compile(r'\bEsCajaApp\b|\bObtener_Caja_App\b|\bEstamosEnCajaPrincipal\b', re.I),
     'Logica de tipo CajaApp (bypass de reglas de caja)'),
    (re.compile(r'\bfValidacionPrincipal\s*\(|\bfValidacionTurnos\s*\(', re.I),
     'Validacion de caja principal o turnos VB.NET'),
    (re.compile(r'\bbalancin\b', re.I),
     'Verificacion de balancin (caja)'),
    # --- C# / ASP.NET ---
    (re.compile(r'\[HttpGet\]|\[HttpPost\]|\[HttpPut\]|\[HttpDelete\]|\[Route\('),
     'Endpoint de API REST ASP.NET'),
    (re.compile(r'\[Authorize\]|\[AllowAnonymous\]'),
     'Control de autorizacion ASP.NET'),
    (re.compile(r'\bDbContext\b|SaveChangesAsync|AddDbContext\b'),
     'Acceso a base de datos Entity Framework'),
    (re.compile(r'\bILogger\b|_logger\.Log|_logger\.Error|_logger\.Info', re.I),
     'Trazabilidad con ILogger .NET'),
    # --- Solucion Visual Studio (.sln) ---
    (re.compile(r'VisualStudioVersion\s*='),
     'Version de Visual Studio en solucion'),
    (re.compile(r'Project\s*\(', re.I),
     'Proyecto referenciado en solucion .sln'),
    # --- Java / Spring ---
    (re.compile(r'@RestController|@Controller|@Service|@Repository|@Component'),
     'Componente de capa Spring (Controller/Service/Repository)'),
    (re.compile(r'@GetMapping|@PostMapping|@PutMapping|@DeleteMapping|@RequestMapping'),
     'Endpoint REST Spring MVC'),
    (re.compile(r'@Transactional|@Rollback|EntityManager\b', re.I),
     'Transaccion JPA/Hibernate'),
    (re.compile(r'@Autowired|@Inject\b|@Bean\b'),
     'Inyeccion de dependencias Spring'),
    # --- Python / Django / Flask ---
    (re.compile(r'urlpatterns\s*=|path\s*\(|include\s*\(', re.I),
     'Definicion de rutas Django/Flask'),
    (re.compile(r'models\.Model|models\.CharField|models\.IntegerField', re.I),
     'Modelo Django'),
    (re.compile(r'@app\.route|@blueprint\.route', re.I),
     'Endpoint Flask'),
    (re.compile(r'\.delay\s*\(|celery\.task|@shared_task', re.I),
     'Tarea asincrona Celery'),
    # --- Go ---
    (re.compile(r'\bhttp\.HandleFunc\b|\bgin\.Default\b|\becho\.New\b'),
     'Handler HTTP Go (net/http / Gin / Echo)'),
    (re.compile(r'\bsql\.Open\b|\bdb\.Exec\b|\bdb\.Query\b'),
     'Acceso a base de datos Go'),
    # --- Rust ---
    (re.compile(r'\bactix_web\b|\bwarp::|\.route\s*\(|\baxum::', re.I),
     'Handler HTTP Rust (Actix/Warp/Axum)'),
]
# =============================================================================
# MOTOR DE ANALISIS SEMANTICO CONTEXTUAL
# =============================================================================

class SemanticInsightEngine:
    """
    Motor heurístico que intenta deducir la intención del cambio.
    No depende de ejemplos específicos.
    Detecta patrones arquitectónicos, de layout, estado y diseño.
    Versión mejorada: incluye inferencia de PROPÓSITO del cambio y análisis
    de patrones adicionales (TypeScript avanzado, SQL, Vue, React, etc.).
    """

    def analyze(self, fc: FileChange) -> List[str]:
        insights = []

        added = "\n".join(fc.added)
        removed = "\n".join(fc.removed)
        full = added + "\n" + removed

        # ── Propósito general del cambio ──────────────────────────────────
        purpose = infer_modification_purpose(fc.added, fc.removed, fc.ext)
        if purpose:
            insights.append(f"🎯 Propósito del cambio: {purpose}")

        # ---------------------------------------------------------
        # 1. Migración de lógica TS hacia CSS
        # ---------------------------------------------------------
        if fc.ext in ('.component.ts', '.ts'):
            if re.search(r'\b(height|width)\s*:', removed) and \
               re.search(r'calc\(.*vh', added):
                insights.append(
                    "Se migra control de dimensiones desde lógica TypeScript "
                    "hacia cálculo dinámico en CSS, favoreciendo separación de responsabilidades."
                )

        # ---------------------------------------------------------
        # 2. Simplificación de binding Angular
        # ---------------------------------------------------------
        if fc.ext in ('.html', '.component.html'):
            if re.search(r'\[.*\]=', removed) and not re.search(r'\[.*\]=', added):
                insights.append(
                    "Se reduce uso de binding dinámico, simplificando el template "
                    "y disminuyendo dependencia del estado del componente."
                )

        # ---------------------------------------------------------
        # 3. Eliminación de propiedad posiblemente obsoleta
        # ---------------------------------------------------------
        if fc.ext in ('.ts', '.component.ts'):
            if re.search(r'\b(public|private|protected)?\s*\w+\s*:', removed):
                insights.append(
                    "Se elimina propiedad del componente, posible reducción "
                    "de estado interno o eliminación de lógica innecesaria."
                )

        # ---------------------------------------------------------
        # 4. Consolidación de estilos
        # ---------------------------------------------------------
        if fc.ext in ('.scss', '.css', '.component.scss'):
            if removed.count('}') > added.count('}'):
                insights.append(
                    "Se reduce cantidad de bloques CSS, indicando consolidación "
                    "o simplificación de reglas de estilo."
                )

        # ---------------------------------------------------------
        # 5. Eliminación de duplicidad estructural (con contexto)
        # ---------------------------------------------------------
        if len(set(fc.removed)) < len(fc.removed):
            from collections import Counter as _Counter
            dup_counts = _Counter(l.strip() for l in fc.removed if len(l.strip()) > 12)
            dup_lines  = [(ln, cnt) for ln, cnt in dup_counts.items() if cnt >= 2]
            if dup_lines and fc.ext in ('.php', '.blade.php'):
                # Intentar dar un nombre al tema duplicado
                sample = dup_lines[0][0]
                if re.search(r'base64_decode|json_decode|explode.*token', sample, re.I):
                    topic = "decodificación de token JWT (base64_decode/json_decode)"
                elif re.search(r'new\s+Client|guzzle|->post\(|->get\(', sample, re.I):
                    topic = "instanciación de cliente HTTP (Guzzle)"
                elif re.search(r'\$request->|->validate', sample, re.I):
                    topic = "acceso a datos de la petición HTTP"
                elif re.search(r"config\s*\(", sample, re.I):
                    topic = "lectura de configuración"
                else:
                    topic = "lógica interna"
                insights.append(
                    f"Se corrige duplicidad estructural: el bloque de {topic} "
                    f"aparecía {dup_lines[0][1]} veces en el mismo método. "
                    "La eliminación de las copias evita inconsistencias por mantenimiento paralelo."
                )
            else:
                insights.append(
                    "Se eliminan líneas repetidas, posible corrección de duplicidad estructural."
                )

        # ---------------------------------------------------------
        # 6. Refactor hacia patrón más declarativo (con contexto)
        # ---------------------------------------------------------
        if re.search(r'if\s*\(', removed) and not re.search(r'if\s*\(', added):
            # Intentar caracterizar qué tipo de condicionales desaparecieron
            if fc.ext in ('.php', '.blade.php'):
                # Condición sobre error de token / claim OAuth2
                if re.search(r"isset\s*\(\s*\$\w+\s*\[\s*['\"]error", removed):
                    insights.append(
                        "Se elimina condicional de validación de error OAuth2 "
                        "(isset($token['error'])): la comprobación de fallo de autenticación "
                        "fue centralizada o delegada a otro componente del flujo."
                    )
                # Condición sobre autorización de usuario
                elif re.search(r"isset\s*\(\s*\$\w+\s*\[\s*['\"](?:sub|preferred_username|realm_access)", removed):
                    insights.append(
                        "Se elimina condicional de verificación de identidad sobre payload JWT; "
                        "la lógica de autorización fue refactorizada o centralizada."
                    )
                # Condición de nulidad / guard
                elif re.search(r'if\s*\(\s*\$\w+\s*(?:===?\s*null|!==?\s*null|\?\?)', removed):
                    insights.append(
                        "Se elimina guard de nulidad explícito (if ($x === null)); "
                        "posible migración al operador null-coalescing (??) o al optional-chaining (?->)."
                    )
                else:
                    insights.append(
                        "Se reduce lógica condicional explícita en archivo PHP, "
                        "posible consolidación de flujo o delegación a middleware/helper."
                    )
            else:
                insights.append(
                    "Se reduce lógica condicional explícita, posible migración "
                    "hacia patrón más declarativo o basado en configuración."
                )

        # ---------------------------------------------------------
        # 7. Mejora responsive
        # ---------------------------------------------------------
        if re.search(r'vh|vw|%', added) and not re.search(r'vh|vw|%', removed):
            insights.append(
                "Se introducen unidades relativas (vh/vw/%), indicando mejora en comportamiento responsive."
            )

        # ---------------------------------------------------------
        # 8. Reducción de acoplamiento
        # ---------------------------------------------------------
        if re.search(r'import ', removed) and not re.search(r'import ', added):
            insights.append(
                "Se eliminan dependencias importadas, posible reducción de acoplamiento."
            )

        # ---------------------------------------------------------
        # 9. Optimización de performance
        # ---------------------------------------------------------
        if re.search(r'\.subscribe\(', removed) and re.search(r'async|await|pipe', added):
            insights.append(
                "Se modifica patrón asincrónico, posible mejora en manejo de suscripciones o performance."
            )

        # ---------------------------------------------------------
        # 10. Eliminación de estado redundante
        # ---------------------------------------------------------
        if re.search(r'\bthis\.', removed) and not re.search(r'\bthis\.', added):
            insights.append(
                "Se reduce uso de propiedades del componente, indicando simplificación del estado."
            )

        # ---------------------------------------------------------
        # 11. VB.NET: eventos comentados (posible desactivación temporal)
        # ---------------------------------------------------------
        if fc.ext == '.vb':
            commented_events = re.findall(
                r"'\s*(?:Private|Public)\s+Sub\s+(\w+)\s*\(", removed
            )
            if commented_events:
                evs = ", ".join(commented_events[:4])
                insights.append(
                    f"Se comentan manejadores de eventos VB.NET ({evs}), "
                    "posible desactivación temporal de funcionalidades interactivas."
                )

        # ---------------------------------------------------------
        # 12. VB.NET: parametrización de función (eliminación de global)
        # ---------------------------------------------------------
        if fc.ext == '.vb':
            if re.search(r'\b(idCaja|idUsuario|idTerminal)\b', added) and \
               not re.search(r'\b(idCaja|idUsuario|idTerminal)\b', removed):
                insights.append(
                    "Se parametrizan funciones VB.NET que antes usaban variables globales, "
                    "mejorando la reutilización y testabilidad del código."
                )

        # ---------------------------------------------------------
        # 13. VB.NET / .NET: corrección de casing de tipo estático
        # ---------------------------------------------------------
        if fc.ext in ('.vb', '.cs'):
            if re.search(r'\bfile\.', removed) and re.search(r'\bFile\.', added):
                insights.append(
                    "Se corrige casing de tipo estático .NET (file. → File.), "
                    "alineando con convenciones BCL de .NET Framework."
                )

        # ---------------------------------------------------------
        # 14-PHP. Análisis lógico profundo para archivos PHP
        # ---------------------------------------------------------
        if fc.ext in ('.php', '.blade.php'):
            php_insights = analyze_php_logic_changes(
                fc.added, fc.removed,
                fc.added_with_line, fc.removed_with_line
            )
            # Solo añadir los que no estén ya en insights
            existing_kw = set(" ".join(insights).lower().split())
            for pi in php_insights:
                pi_words = set(pi.lower().split())
                # heurística: si comparte menos de 3 palabras clave con existentes, agregar
                overlap = pi_words & existing_kw
                if len(overlap) < 5:
                    insights.append(pi)
                    existing_kw.update(pi_words)

        # ---------------------------------------------------------
        # 14. Archivo de respaldo (.bak): sin valor funcional
        # ---------------------------------------------------------
        if fc.ext == '.bak':
            insights.append(
                "Archivo de respaldo automático (IDE). No representa lógica nueva; "
                "contiene snapshot del estado anterior del archivo original."
            )

        # ---------------------------------------------------------
        # 15. Solución Visual Studio (.sln): cambio de scope del proyecto
        # ---------------------------------------------------------
        if fc.ext == '.sln':
            if re.search(r'VisualStudioVersion', full):
                insights.append(
                    "Se actualiza la versión del entorno de desarrollo en el archivo de solución, "
                    "lo que puede requerir que todos los colaboradores actualicen su Visual Studio."
                )
            removed_projects = len(re.findall(r'^-.*Project\s*\(', removed, re.M))
            added_projects   = len(re.findall(r'^\+?.*Project\s*\(', added,   re.M))
            if removed_projects > added_projects:
                insights.append(
                    f"Se eliminan {removed_projects - added_projects} proyecto(s) de la solución, "
                    "reduciendo el scope del build y la solución."
                )
            elif added_projects > removed_projects:
                insights.append(
                    f"Se incorporan {added_projects - removed_projects} proyecto(s) nuevo(s) a la solución."
                )

        # ---------------------------------------------------------
        # 16. Lógica de bypass condicional (CajaApp / feature flag)
        # ---------------------------------------------------------
        if re.search(r'\bEsCajaApp\b|\bObtener_Caja_App\b|\bEstamosEnCajaPrincipal\b', full, re.I):
            insights.append(
                "Se introduce bifurcación condicional basada en tipo de caja (CajaApp), "
                "permitiendo que ciertos registros omitan restricciones de IP y balance. "
                "Patrón equivalente a feature-flag a nivel de entidad."
            )

        # ---------------------------------------------------------
        # 17. TypeScript: migración de callback a async/await
        # ---------------------------------------------------------
        if fc.ext in ('.ts', '.component.ts', '.service.ts'):
            callbacks_rem = len(re.findall(r'\.then\s*\(|\.catch\s*\(|new\s+Promise\s*\(', removed))
            async_add     = len(re.findall(r'\bawait\b|\basync\b', added))
            if callbacks_rem >= 2 and async_add >= 1:
                insights.append(
                    "Modernización asíncrona: se migran cadenas de Promises/callbacks (.then/.catch) "
                    "al patrón async/await, mejorando legibilidad y manejo de errores."
                )

        # ---------------------------------------------------------
        # 18. TypeScript: extracción de método (método largo → varios cortos)
        # ---------------------------------------------------------
        if fc.ext in ('.ts', '.component.ts', '.service.ts'):
            fn_add = len(re.findall(r'(?:private|public|protected)?\s*\w+\s*\([^)]*\)\s*[:{]', added))
            fn_rem = len(re.findall(r'(?:private|public|protected)?\s*\w+\s*\([^)]*\)\s*[:{]', removed))
            if fn_add > fn_rem + 1:
                insights.append(
                    f"Extracción de métodos: se dividen responsabilidades en {fn_add - fn_rem} función(es) "
                    "adicional(es), favoreciendo el principio de responsabilidad única (SRP)."
                )

        # ---------------------------------------------------------
        # 19. Introducción de tipado fuerte (any → tipo concreto)
        # ---------------------------------------------------------
        if fc.ext in ('.ts', '.component.ts', '.service.ts', '.interface.ts'):
            any_rem = len(re.findall(r':\s*any\b', removed))
            any_add = len(re.findall(r':\s*any\b', added))
            typed_add = len(re.findall(
                r':\s*(?:string|number|boolean|Date|Observable|Promise|Array|Record|Map)\b', added
            ))
            if any_rem > any_add and typed_add > 0:
                insights.append(
                    f"Refuerzo de tipado: se reemplazan {any_rem - any_add} uso(s) de 'any' "
                    "por tipos concretos, mejorando la detección de errores en tiempo de compilación."
                )

        # ---------------------------------------------------------
        # 20. SQL: cambio de tipo de JOIN (seguridad / datos faltantes)
        # ---------------------------------------------------------
        if fc.ext in ('.sql', '.py', '.php', '.ts', '.js', '.cs', '.java'):
            join_rem = re.findall(r'\b(INNER\s+JOIN|LEFT\s+JOIN|RIGHT\s+JOIN|FULL\s+JOIN|CROSS\s+JOIN)\b',
                                  removed, re.I)
            join_add = re.findall(r'\b(INNER\s+JOIN|LEFT\s+JOIN|RIGHT\s+JOIN|FULL\s+JOIN|CROSS\s+JOIN)\b',
                                  added, re.I)
            if join_rem and join_add:
                j_rem_set = {j.upper().strip() for j in join_rem}
                j_add_set = {j.upper().strip() for j in join_add}
                if j_rem_set != j_add_set:
                    from_j = ", ".join(sorted(j_rem_set))
                    to_j   = ", ".join(sorted(j_add_set))
                    insights.append(
                        f"Cambio de JOIN en consulta SQL: de '{from_j}' a '{to_j}'. "
                        "Esto puede afectar qué registros se incluyen cuando no hay datos relacionados."
                    )

        # ---------------------------------------------------------
        # 21. React/Vue: gestión de efectos secundarios
        # ---------------------------------------------------------
        if fc.ext in ('.tsx', '.jsx', '.vue'):
            use_effect_rem = len(re.findall(r'useEffect\s*\(', removed))
            use_effect_add = len(re.findall(r'useEffect\s*\(', added))
            watch_rem = len(re.findall(r'\bwatch\s*\(|\bwatchEffect\s*\(', removed))
            watch_add = len(re.findall(r'\bwatch\s*\(|\bwatchEffect\s*\(', added))

            if use_effect_add > use_effect_rem:
                insights.append(
                    f"Se agregan {use_effect_add - use_effect_rem} efecto(s) secundario(s) (useEffect) "
                    "para sincronizar estado con dependencias externas o del DOM."
                )
            elif use_effect_rem > use_effect_add:
                insights.append(
                    f"Se eliminan {use_effect_rem - use_effect_add} efecto(s) secundario(s) (useEffect); "
                    "posible simplificación del ciclo de vida del componente."
                )
            if watch_add > watch_rem:
                insights.append(
                    f"Se agregan {watch_add - watch_rem} watcher(s) Vue para reaccionar "
                    "a cambios de estado reactivo."
                )

        # ---------------------------------------------------------
        # 22. Detección de renombrado semántico de variable/método
        # ---------------------------------------------------------
        identifiers_rem = set(re.findall(r'\b([a-z][a-zA-Z0-9]{3,})\b', removed))
        identifiers_add = set(re.findall(r'\b([a-z][a-zA-Z0-9]{3,})\b', added))
        only_rem = identifiers_rem - identifiers_add - {'true','false','null','undefined','this','self'}
        only_add = identifiers_add - identifiers_rem - {'true','false','null','undefined','this','self'}
        # Filtrar: solo palabras que parecen identificadores de dominio (no palabras comunes)
        stop = {'return','const','let','var','function','class','import','export',
                'from','public','private','protected','static','async','await',
                'interface','type','enum','string','number','boolean','void',
                'that','with','have','this','your','from','they','their'}
        only_rem -= stop
        only_add -= stop
        if 1 <= len(only_rem) <= 4 and 1 <= len(only_add) <= 4:
            rem_list = sorted(only_rem)[:3]
            add_list = sorted(only_add)[:3]
            if rem_list and add_list:
                insights.append(
                    f"Posible renombrado semántico: identificador(es) '{', '.join(rem_list)}' "
                    f"reemplazado(s) por '{', '.join(add_list)}', sugiriendo "
                    "aclaración del nombre o cambio en el modelo de dominio."
                )

        # ---------------------------------------------------------
        # 23. Introducción de manejo de errores donde no había
        # ---------------------------------------------------------
        has_try_add = bool(re.search(r'\btry\s*[{\(]|\bcatch\s*\(|\bexcept\b|\brescue\b', added, re.I))
        has_try_rem = bool(re.search(r'\btry\s*[{\(]|\bcatch\s*\(|\bexcept\b|\brescue\b', removed, re.I))
        if has_try_add and not has_try_rem:
            insights.append(
                "Se introduce manejo de excepciones (try/catch) donde antes el código era "
                "susceptible a fallar silenciosamente; mejora la robustez ante errores inesperados."
            )
        elif has_try_rem and not has_try_add:
            insights.append(
                "Se elimina bloque try/catch previo; verificar que los errores posibles "
                "queden manejados en una capa superior o que el flujo garantice no lanzar excepciones."
            )

        # ---------------------------------------------------------
        # 24. Cambio de valor literal hardcodeado a variable/constante
        # ---------------------------------------------------------
        hardcoded_rem = re.findall(r'[\'"][A-Z0-9_]{4,}[\'"]|=\s*\d{2,}', removed)
        hardcoded_add = re.findall(r'[\'"][A-Z0-9_]{4,}[\'"]|=\s*\d{2,}', added)
        const_add     = bool(re.search(r'\b(const|final|readonly|CONSTANT|CONFIG)\b', added, re.I))
        if len(hardcoded_rem) > len(hardcoded_add) and const_add:
            insights.append(
                "Se reemplazan valores hardcodeados por constantes o variables configurables, "
                "centralizando la configuración y facilitando futuros cambios sin tocar la lógica."
            )

        # ---------------------------------------------------------
        # 25. Eliminación de código comentado (deuda técnica)
        # ---------------------------------------------------------
        comment_rem_lines = [
            l for l in fc.removed
            if re.match(r'\s*(?://|#|\'|\*|<!--|/\*)', l.strip()) and len(l.strip()) > 5
        ]
        if len(comment_rem_lines) >= 3:
            insights.append(
                f"Se eliminan {len(comment_rem_lines)} línea(s) de código comentado (deuda técnica), "
                "limpiando el historial del archivo y reduciendo confusión para futuros revisores."
            )

        # ---------------------------------------------------------
        # 26. CSS: cambio de unidades absolutas a relativas (accesibilidad)
        # ---------------------------------------------------------
        if fc.ext in ('.css', '.scss', '.component.scss'):
            px_rem = len(re.findall(r'\d+px', removed))
            px_add = len(re.findall(r'\d+px', added))
            rem_add_units = len(re.findall(r'\d+(?:rem|em|%|vh|vw)', added))
            if px_rem > px_add and rem_add_units > 0:
                insights.append(
                    f"Mejora de accesibilidad: se reducen {px_rem - px_add} valor(es) en píxeles (px) "
                    "reemplazados por unidades relativas (rem/em/%), respetando la configuración "
                    "de fuente del navegador del usuario."
                )

        return insights


class ChangeRelationAnalyzer:
    """Detecta relaciones funcionales entre archivos modificados del mismo diff."""

    _ENV_ENTRY = re.compile(r"^\s*([a-zA-Z_]\w*)\s*:\s*['\"]([^'\"]+)['\"]\s*,?\s*$")

    def __init__(self, changes: List[FileChange]):
        self.changes = changes
        self.by_path = {c.filepath: c for c in changes}

    def _component_family(self, fc: FileChange) -> List[str]:
        """Relaciona artefactos hermanos de un componente Angular."""
        name = fc.filename
        if '.component.' not in name:
            return []

        base = name.split('.component.')[0]
        folder = str(Path(fc.filepath).parent).replace('\\', '/')
        siblings: List[str] = []
        for other in self.changes:
            if other.filepath == fc.filepath:
                continue
            other_folder = str(Path(other.filepath).parent).replace('\\', '/')
            other_name = Path(other.filepath).name
            if other_folder == folder and other_name.startswith(base + '.component.'):
                siblings.append(other_name)
        return siblings

    def _extract_env_entries(self, fc: FileChange) -> List[Tuple[str, str, Optional[int]]]:
        entries: List[Tuple[str, str, Optional[int]]] = []
        seen: Set[Tuple[str, str]] = set()
        source = fc.added_with_line or [(None, l) for l in fc.added]
        for line_no, line in source:
            m = self._ENV_ENTRY.match(line.strip())
            if not m:
                continue
            key = m.group(1)
            val = m.group(2)
            if key.lower() in ('production', 'qa', 'local', 'staging'):
                continue
            pair = (key, val)
            if pair in seen:
                continue
            seen.add(pair)
            entries.append((key, val, line_no))
        return entries

    def _usage_files_for_env_key(self, key: str, value: str, current_path: str) -> Tuple[List[str], List[str]]:
        non_env_users: List[str] = []
        env_users: List[str] = []
        key_ref = re.compile(rf'\benvironment\.{re.escape(key)}\b')
        val_ref = re.compile(re.escape(value))

        for other in self.changes:
            if other.filepath == current_path:
                continue
            sample = "\n".join(other.added + other.removed + other.full_content[:200])
            if key_ref.search(sample) or val_ref.search(sample):
                if other.is_environment_file:
                    env_users.append(other.filename)
                else:
                    non_env_users.append(other.filename)
        return non_env_users[:6], env_users[:6]

    def _endpoint_purpose(self, key: str, value: str) -> str:
        text = f"{key} {value}".lower()
        if re.search(r'mantenimiento|catalogo', text):
            return "habilitar consumo de catálogos/mantenimientos"
        if re.search(r'auth|login|token', text):
            return "habilitar flujo de autenticación"
        if re.search(r'perfil|user|usuario', text):
            return "habilitar consulta/gestión de usuarios"
        if re.search(r'param|config', text):
            return "habilitar lectura de parámetros del sistema"
        return "habilitar consumo de API del módulo"

    def analyze(self, fc: FileChange) -> List[str]:
        insights: List[str] = []

        siblings = self._component_family(fc)
        if siblings:
            insights.append(
                "Interrelación de componente: se modifican artefactos relacionados "
                f"({', '.join(siblings)}), manteniendo coherencia entre lógica, vista y estilos."
            )

        if fc.is_environment_file:
            env_entries = self._extract_env_entries(fc)
            for key, value, _line_no in env_entries[:6]:
                purpose = self._endpoint_purpose(key, value)
                non_env_users, env_users = self._usage_files_for_env_key(key, value, fc.filepath)
                msg = (
                    f"Se agrega configuración de endpoint '{key}: {value}' para {purpose}."
                )
                if non_env_users:
                    msg += f" Consumo detectado en archivos funcionales: {', '.join(non_env_users)}."
                elif env_users:
                    msg += f" Homologado en ambientes: {', '.join(env_users)}."
                else:
                    msg += " No se detecta consumo directo en otros archivos modificados de este diff."
                insights.append(msg)

        # Relación por importaciones directas entre archivos cambiados
        import_lines = [l for l in fc.added if ' from ' in l and 'import ' in l]
        related: List[str] = []
        for line in import_lines:
            m = re.search(r"from\s+['\"]([^'\"]+)['\"]", line)
            if not m:
                continue
            path_hint = m.group(1).lower()
            for other in self.changes:
                if other.filepath == fc.filepath:
                    continue
                other_name = other.filename.lower().replace('.ts', '').replace('.html', '').replace('.scss', '')
                if other_name and other_name in path_hint and other.filename not in related:
                    related.append(other.filename)

        if related:
            insights.append(
                "Dependencias cruzadas detectadas con archivos también modificados: "
                f"{', '.join(related[:6])}."
            )

        return insights[:6]


# =============================================================================
# MOTOR DE ANÁLISIS ESLINT ANGULAR (Angular 16 · strict mode)
# Reglas activas según .eslintrc.json del proyecto:
#   .html → @angular-eslint/template/recommended  (reglas 16-18)
#   .ts   → @angular-eslint/recommended            (reglas 1-15)
# =============================================================================

# Eventos DOM nativos que no deben usarse como nombre de @Output()
DOM_NATIVE_EVENTS: frozenset = frozenset({
    'click', 'focus', 'blur', 'change', 'input', 'submit',
    'keyup', 'keydown', 'keypress', 'mouseenter', 'mouseleave',
    'mouseover', 'mouseout', 'mouseup', 'mousedown', 'dblclick',
    'contextmenu', 'scroll', 'resize', 'load', 'error', 'abort',
    'select', 'reset', 'drag', 'dragstart', 'dragend', 'dragenter',
    'dragleave', 'dragover', 'drop', 'touchstart', 'touchend',
    'touchmove', 'touchcancel', 'wheel', 'copy', 'cut', 'paste',
    'beforeinput',
})

# Hooks de ciclo de vida Angular y su interfaz correspondiente
LIFECYCLE_HOOKS: frozenset = frozenset({
    'ngOnInit', 'ngOnDestroy', 'ngOnChanges', 'ngDoCheck',
    'ngAfterContentInit', 'ngAfterContentChecked',
    'ngAfterViewInit', 'ngAfterViewChecked',
})

LIFECYCLE_INTERFACES: Dict[str, str] = {
    'ngOnInit':              'OnInit',
    'ngOnDestroy':           'OnDestroy',
    'ngOnChanges':           'OnChanges',
    'ngDoCheck':             'DoCheck',
    'ngAfterContentInit':    'AfterContentInit',
    'ngAfterContentChecked': 'AfterContentChecked',
    'ngAfterViewInit':       'AfterViewInit',
    'ngAfterViewChecked':    'AfterViewChecked',
}


@dataclass
class ESLintFinding:
    """Representa una corrección o violación ESLint detectada en una línea de diff."""
    rule:         str   # ej: "@angular-eslint/template/eqeqeq"
    severity:     str   # "ERROR"
    category:     str   # "CORRECCION" | "VIOLACION"
    line_added:   str   # línea nueva  (prefijo +)
    line_removed: str   # línea previa (prefijo -), puede ser ""
    description:  str   # descripción legible


@dataclass
class FrontBackRiskFinding:
    """
    Riesgo de incompatibilidad de tipos entre el front (Angular/JS)
    y el back-end según el operador de comparación utilizado.

    Escenarios típicos:
      • Backend envía 1 (número) y el front compara con true (boolean):  1 === true → FALSE
      • Backend envía null para campo opcional y el front verifica === undefined → FALSE
      • Backend serializa booleano como string "true" / "1"; el front falla con ===
    """
    scenario:       str   # Nombre corto del escenario (ej: '1 vs true')
    risk_level:     str   # 'ALTO' | 'MEDIO' | 'BAJO'
    expression:     str   # Expresión concreta detectada
    explanation:    str   # Qué falla y por qué
    recommendation: str   # Cómo corregirlo
    context_line:   str   # Línea de código donde se detectó


@dataclass
class JSStructuralFinding:
    """
    Hallazgo de cambio estructural en código TypeScript/JavaScript.

    Detecta patrones que alteran la semántica de guardas de nulidad,
    evaluación de falsy/nullish, coerción implícita o aserciones de tipo.

    Ejemplos cubiertos:
      • 'x && x.prop'            →  'x?.prop'        (guarda && → optional chaining)
      • 'x || default'           →  'x ?? default'   (logical OR → nullish coalescing)
      • 'x.prop'                 →  'x!.prop'        (non-null assertion añadida)
      • 'typeof x !== "undefined"' → 'x !== undefined' (typeof guard eliminada)
      • 'x !== null ? x : d'     →  'x ?? d'         (ternario nulo → ??)
    """
    pattern:        str   # Nombre corto del patrón detectado
    severity:       str   # 'ATENCION' | 'INFORMATIVO'
    before:         str   # Fragmento representativo antes del cambio
    after:          str   # Fragmento representativo después del cambio
    explanation:    str   # Diferencia semántica exacta que introduce el cambio
    recommendation: str   # Cómo validar que el cambio es correcto


@dataclass
class ESLintFileReport:
    """Resultado del análisis ESLint de un archivo."""
    corrections: List[ESLintFinding]      = field(default_factory=list)
    violations:  List[ESLintFinding]      = field(default_factory=list)
    functional:  List[str]                = field(default_factory=list)
    removed_had_violations: bool          = False   # True si alguna línea eliminada tenía patrón ESLint
    front_back_risks: "List[FrontBackRiskFinding]" = field(default_factory=list)
    structural_findings: "List[JSStructuralFinding]" = field(default_factory=list)


# =============================================================================
# MOTOR DE RIESGOS FRONT-BACK: COERCIÓN DE TIPOS Y CONTRATO DE DATOS
# =============================================================================

class FrontBackTypeRiskEngine:
    """
    Motor heurístico que analiza comparaciones JS/TS para detectar riesgos de
    incompatibilidad de tipos entre el front-end Angular y el back-end.

    El problema central es que JavaScript tiene dos tipos de igualdad:
      • == (abstracta): convierte tipos antes de comparar  → 1 == true  es TRUE
      • === (estricta): compara tipo Y valor sin conversión → 1 === true es FALSE

    Cuando ESLint corrige == → ===, puede alterar silenciosamente la lógica si
    el back-end envía tipos distintos a los que el front-end espera.

    Escenarios críticos detectados:
      1. null/undefined mismatch  – backend retorna null, front verifica ===undefined → FALLA
      2. 1 vs true               – backend usa 0/1 para booleanos, front compara ===true/false
      3. "1" vs 1                – backend PHP/MySQL serializa numérico como string
      4. "" vs false/null        – string vacío coercionado a falsy
      5. Comparación null        – x==null captura null Y undefined (patrón intencional JS)
    """

    # Literales de riesgo y su metadata de escenario
    # (valor_literal): (nivel_riesgo, escenario, explicacion, recomendacion)
    #
    # _NEGATED_RISKY_LITERALS: se usa cuando el operador ORIGINAL era != (negado).
    # El efecto de != → !== es el contrario al de == → ===:
    # en vez de "dejar de capturar" algo, EMPIEZA a capturar algo que antes quedaba excluido.
    _NEGATED_RISKY_LITERALS: Dict[str, Tuple[str, str, str, str]] = {
        'undefined': (
            'ALTO',
            'null visible cuando debería estar oculto (null/undefined mismatch negado)',
            "Con != el operador evalúa null como equivalente a undefined, por lo que "
            "la condición es FALSE (oculta/no ejecuta) cuando el valor es null O undefined. "
            "Con !== en cambio, null !== undefined es TRUE: si el back-end devuelve null "
            "para un campo opcional (patrón habitual en PHP/Laravel/Java), la condición "
            "PASA aunque el dato no exista, mostrando o ejecutando lógica que debería "
            "estar desactivada.",
            "Verificar si el back-end puede retornar null para este campo. Si es así, "
            "añadir una guarda explícita: 'value !== undefined && value !== null', "
            "o volver al patrón idiomático '!= null' que excluye ambos valores de referencia."
        ),
        'null': (
            'ALTO',
            'undefined visible cuando debería estar oculto (null/undefined mismatch negado)',
            "Con != el operador trata undefined como equivalente a null, por lo que "
            "la condición es FALSE (oculta/no ejecuta) cuando el valor es null O undefined. "
            "Con !== en cambio, undefined !== null es TRUE: si el campo nunca fue "
            "inicializado en el front (undefined), la condición PASA y la lógica se activa "
            "aunque no haya dato real.",
            "Añadir guarda explícita: 'value !== undefined && value !== null', "
            "o usar el patrón idiomático '!= null' que excluye ambos valores de referencia."
        ),
        'true': (
            'ALTO',
            '1 oculto cuando debería estar visible (boolean/number negado)',
            "Con != el número 1 se trata como equivalente a true, por lo que "
            "la condición es FALSE (oculta/no ejecuta) para tanto true como 1. "
            "Con !== en cambio, 1 !== true es TRUE: si el back-end retorna 1 (entero) "
            "como representación del activo/habilitado (patrón común en APIs PHP/MySQL/Laravel), "
            "la condición PASA y la lógica se activa cuando debería estar desactivada.",
            "Asegurarse de que el back-end serialice el campo como boolean en JSON, "
            "o en el front normalizar: 'value !== true && value !== 1'."
        ),
        'false': (
            'ALTO',
            '0/string-vacío visible cuando debería estar oculto (boolean/number negado)',
            "Con != el valor 0 (entero) y '' (string vacío) se tratan como equivalentes a false, "
            "por lo que la condición es FALSE (oculta/no ejecuta) para false, 0 y ''. "
            "Con !== en cambio, 0 !== false y '' !== false son TRUE: si el back-end retorna "
            "0 o '' para indicar inactivo/vacío, la condición PASA y la lógica se activa "
            "cuando debería estar desactivada, rompiendo silenciosamente la guarda.",
            "Verificar que el back-end serialice como boolean. En el front usar "
            "'value !== false && value !== 0 && value !== ''', o Boolean(value) === false."
        ),
        '0': (
            'ALTO',
            'false/string-vacío visible cuando debería estar oculto (boolean/number negado)',
            "Con != el boolean false y el string '' se tratan como equivalentes a 0, "
            "por lo que la condición es FALSE (oculta/no ejecuta) para 0, false y ''. "
            "Con !== en cambio, false !== 0 y '' !== 0 son TRUE: si el back-end retorna "
            "false o '' para indicar cero/inactivo, la condición PASA incorrectamente.",
            "Tipificar el campo y normalizar con Number(value) !== 0, "
            "o usar Boolean(value) para flags."
        ),
        '1': (
            'ALTO',
            'true/string-uno visible cuando debería estar oculto (boolean/number negado)',
            "Con != el boolean true y el string '1' se tratan como equivalentes a 1, "
            "por lo que la condición es FALSE (oculta/no ejecuta) para 1, true y '1'. "
            "Con !== en cambio, true !== 1 y '1' !== 1 son TRUE: si el back-end retorna "
            "true o '1' para indicar activo, la condición PASA cuando no debería.",
            "Tipificar el campo y normalizar con Number(value) !== 1, "
            "o usar Boolean(value) para flags."
        ),
        '"1"': (
            'MEDIO',
            'número-uno visible cuando debería estar oculto (string numérico negado)',
            "Con != el número 1 se trata como equivalente al string '1'. "
            "Con !== en cambio, 1 !== '1' es TRUE: si el back-end retorna el entero 1 "
            "(API JSON tipada), la condición PASA cuando antes quedaba bloqueada.",
            "Usar Number(value) !== 1 o value?.toString() !== '1' para comparación robusta."
        ),
        "'1'": (
            'MEDIO',
            'número-uno visible cuando debería estar oculto (string numérico negado)',
            "Con != el número 1 se trata como equivalente al string '1'. "
            "Con !== en cambio, 1 !== '1' es TRUE: si el back-end retorna el entero 1, "
            "la condición PASA cuando antes quedaba bloqueada.",
            "Usar Number(value) !== 1 o value?.toString() !== '1' para comparación robusta."
        ),
        '"0"': (
            'MEDIO',
            'número-cero visible cuando debería estar oculto (string numérico negado)',
            "Con != el número 0 se trata como equivalente al string '0'. "
            "Con !== en cambio, 0 !== '0' es TRUE: si el back-end retorna el entero 0, "
            "la condición PASA cuando antes quedaba bloqueada.",
            "Usar Number(value) !== 0 o value?.toString() !== '0'."
        ),
        "'0'": (
            'MEDIO',
            'número-cero visible cuando debería estar oculto (string numérico negado)',
            "Con != el número 0 se trata como equivalente al string '0'. "
            "Con !== en cambio, 0 !== '0' es TRUE: si el back-end retorna el entero 0, "
            "la condición PASA cuando antes quedaba bloqueada.",
            "Usar Number(value) !== 0 o value?.toString() !== '0'."
        ),
    }

    _RISKY_LITERALS: Dict[str, Tuple[str, str, str, str]] = {
        'undefined':  (
            'ALTO',
            'null/undefined mismatch',
            "Con == el operador captura tanto null como undefined. "
            "Con === solo captura undefined. Si el back-end devuelve null para un "
            "campo opcional (comportamiento habitual en PHP/Laravel/Java), "
            "la condición dejará de activarse silenciosamente.",
            "Verificar si el back-end puede retornar null. Si es así, usar "
            "'=== null || === undefined', o mantener '== null' que captura ambos de forma intencional."
        ),
        'null': (
            'ALTO',
            'null/undefined mismatch (reversed)',
            "El patrón 'x == null' es idiomático en JS: captura null Y undefined con == . "
            "Cambiarlo a '=== null' rompe el guard de undefined, exponiendo errores si "
            "el campo nunca fue inicializado en el front.",
            "Mantener '== null' cuando se quiere capturar ambos valores falsy de referencia, "
            "o verificar explícitamente '=== null || === undefined'."
        ),
        'true': (
            'ALTO',
            '1 vs true (boolean/number)',
            "Con == el valor numérico 1 equivale a true dando TRUE. Con === NO. "
            "Es común que APIs PHP, MySQL, C# o Laravel retornen 1/0 como representación "
            "de booleanos (ej: activo=1, habilitado=0). El front que compare ===true fallará.",
            "Asegurarse que el back-end serialice el campo como boolean (true/false) en JSON, "
            "o en el front castear explícitamente: 'value === true || value === 1'."
        ),
        'false': (
            'ALTO',
            '0 vs false (boolean/number)',
            "Con == el valor 0 y '' (string vacío) equivalen a false. Con === NO. "
            "Si el back-end retorna 0 (entero) o '0' (string) para indicar inactivo, "
            "la comparación ===false dejará de funcionar.",
            "Verificar que el back-end serialice como boolean. En el front usar Boolean(value) "
            "o comparar explícitamente contra 0 y false."
        ),
        '1': (
            'ALTO',
            'boolean/string como número 1',
            "Con == el valor true (boolean) y '1' (string) equivalen a 1 (number). Con === NO. "
            "Si el campo del back-end puede llegar como boolean true o string '1', "
            "la comparación ===1 fallará en ambos casos.",
            "Tipificar el campo en la interfaz TypeScript y asegurarse que el back-end siempre "
            "retorne number. O usar Number(value) === 1 para normalizar."
        ),
        '0': (
            'ALTO',
            'boolean/string como número 0',
            "Con == el valor false (boolean) y '0' (string) equivalen a 0 (number). Con === NO. "
            "APIs que usan 0/1 como flags retornarían false o '0' y la comparación ===0 fallaría.",
            "Tipificar el campo y normalizar: Number(value) === 0 o Boolean(value) === false."
        ),
        '"1"': (
            'MEDIO',
            'string numérico vs number',
            "Con == el número 1 equivale al string '1'. Con === NO. "
            "Comun en APIs que serializan IDs o flags numéricos como strings (PHP echo, MySQL).",
            "Usar Number(value) === 1 o value?.toString() === '1' para comparación robusta."
        ),
        "'1'": (
            'MEDIO',
            'string numérico vs number',
            "Con == el número 1 equivale al string '1'. Con === NO. "
            "Comun en APIs que serializan IDs o flags numéricos como strings (PHP echo, MySQL).",
            "Usar Number(value) === 1 o value?.toString() === '1' para comparación robusta."
        ),
        '"0"': (
            'MEDIO',
            'string numérico vs number',
            "Con == el número 0 equivale al string '0'. Con === NO. "
            "APIs PHP/MySQL pueden retornar '0' (string) para valores nulos o falsos.",
            "Usar Number(value) === 0 o value?.toString() === '0'."
        ),
        "'0'": (
            'MEDIO',
            'string numérico vs number',
            "Con == el número 0 equivale al string '0'. Con === NO. "
            "APIs PHP/MySQL pueden retornar '0' (string) para valores nulos o falsos.",
            "Usar Number(value) === 0 o value?.toString() === '0'."
        ),
        '"true"': (
            'MEDIO',
            'boolean serializado como string',
            "Algunos backends (especialmente query params y form data) envían 'true' como string. "
            "Con === 'true' solo coincide el string, no el boolean true.",
            "Normalizar en el servicio: value === true || value === 'true'."
        ),
        '"false"': (
            'MEDIO',
            'boolean serializado como string',
            "Backends pueden enviar 'false' como string en query params o form data. "
            "Con === 'false' solo coincide el string, no el boolean false.",
            "Normalizar en el servicio: value === false || value === 'false'."
        ),
    }

    # Regex para extraer comparaciones: ident == literal  o  literal == ident
    _RE_CMP = re.compile(
        r'(?P<lhs>[\w.$?\[\]]+)\s*(?P<op>===|!==|==|!=)\s*(?P<rhs>["\']?[\w.]+["\']?)'
        r'|'
        r'(?P<rhs2>["\']?[\w.]+["\']?)\s*(?P<op2>===|!==|==|!=)\s*(?P<lhs2>[\w.$?\[\]]+)',
        re.IGNORECASE,
    )

    # Regex para detectar == / != sin triple
    _RE_LOOSE = re.compile(r'(?<![=!<>])={2}(?!=)|(?<![!<>])!={1}(?!=)')

    def _extract_comparisons(self, line: str) -> List[Tuple[str, str, str]]:
        """
        Extrae tuplas (lhs, operador, rhs) de las comparaciones en la línea.
        _RE_CMP tiene === y !== antes de == y != en la alternación → operadores
        capturados correctamente.

        También limpia comillas de delimitador HTML que pueden filtrarse en los
        tokens capturados (ej. la apertura de *ngIf="expr" hace que el regex
        capture '"selectedTurno' en lugar de 'selectedTurno').
        Una comilla SOLA en un lado indica delimitador HTML; si hay comillas
        coincidentes en ambos lados es un literal de cadena JS ('0', "1") y se
        preserva intacto.
        """
        def _clean_token(tok: str) -> str:
            """Elimina comillas de delimitador HTML filtradas (solo un lado)."""
            if len(tok) < 2:
                return tok
            lead  = tok[0]  in ('"', "'")
            trail = tok[-1] in ('"', "'")
            if lead and trail and tok[0] == tok[-1]:
                return tok           # literal de cadena JS ("1", '0') → conservar
            if lead and not trail:
                return tok.lstrip("\"'")   # comilla de apertura filtrada
            if trail and not lead:
                return tok.rstrip("\"'")   # comilla de cierre filtrada
            return tok

        results: List[Tuple[str, str, str]] = []
        for m in self._RE_CMP.finditer(line):
            lhs = _clean_token((m.group('lhs') or m.group('lhs2') or '').strip())
            op  =              (m.group('op')  or m.group('op2')  or '').strip()
            rhs = _clean_token((m.group('rhs') or m.group('rhs2') or '').strip())
            if lhs and op and rhs:
                results.append((lhs, op, rhs))
        return results

    def analyze_line_pair(
        self,
        removed_line: str,
        added_line:   str,
    ) -> List[FrontBackRiskFinding]:
        """
        Analiza un par (línea eliminada → línea añadida) donde ESLint corrigió
        == a ===. Detecta riesgos de incompatibilidad de tipos Front-Back SÓLO
        para las comparaciones que EFECTIVAMENTE cambiaron de == a === en esa
        línea concreta, evitando falsos positivos por expresiones que ya eran ===.

        Algoritmo:
          1. Extraer pares (lhs, rhs) con == / != de la línea ELIMINADA.
          2. Para cada comparación === / !== de la línea AÑADIDA, verificar que
             el mismo par concreto existía como == en la línea eliminada.
          3. Solo reportar si hay coincidencia exacta del par → cambio real.
        """
        risks: List[FrontBackRiskFinding] = []
        r = removed_line.strip()
        a = added_line.strip()

        # ── Paso 1: índice de comparaciones LAXAS en la línea ELIMINADA ─────
        # Clave: (lhs_norm, rhs_norm) → operador original ('==' o '!=')
        # Se indexan ambos órdenes (a == b y b == a) por si hay espejo.
        loose_pairs: Dict[Tuple[str, str], str] = {}
        for lhs_r, op_r, rhs_r in self._extract_comparisons(r):
            if op_r in ('==', '!='):
                k_fwd = (lhs_r.strip().lower(), rhs_r.strip().lower())
                k_rev = (rhs_r.strip().lower(), lhs_r.strip().lower())
                loose_pairs[k_fwd] = op_r
                loose_pairs[k_rev] = op_r

        # Si la línea eliminada no tenía ningún == / !=, no hubo conversión real
        if not loose_pairs and r:
            return risks

        # ── Paso 2: iterar comparaciones ESTRICTAS en la línea AÑADIDA ──────
        for lhs, op, rhs in self._extract_comparisons(a):
            if op not in ('===', '!=='):
                continue

            # ── Validación cruzada: ¿este par concreto era == en la eliminada? ──
            pair_key = (lhs.strip().lower(), rhs.strip().lower())
            if pair_key not in loose_pairs:
                # Esta comparación ya era === antes del cambio → no reportar
                continue
            original_op = loose_pairs[pair_key]  # '==' o '!='

            # Revisar si el rhs (o lhs) es un literal de riesgo
            literal_check = rhs.strip()
            meta = self._RISKY_LITERALS.get(literal_check)
            if not meta:
                literal_check = lhs.strip()
                meta = self._RISKY_LITERALS.get(literal_check)
                if meta:
                    lhs, rhs = rhs, lhs   # espejo: literal estaba a la izquierda

            if not meta:
                continue

            risk_level, scenario, explanation, recommendation = meta

            # Si el operador original era != (negado), el efecto es el contrario:
            # en vez de "dejar de capturar" algo al pasar a ===, EMPIEZA a capturar
            # algo que antes excluía. Usar explicación específica si existe.
            if original_op == '!=':
                neg_meta = self._NEGATED_RISKY_LITERALS.get(literal_check)
                if neg_meta:
                    risk_level, scenario, explanation, recommendation = neg_meta

            # Variable concreta para enriquecer la explicación
            var_name  = rhs if lhs in self._RISKY_LITERALS else lhs
            expr_display   = f"{lhs} {op} {rhs}"
            original_expr  = f"{lhs} {original_op} {rhs}"

            risks.append(FrontBackRiskFinding(
                scenario=scenario,
                risk_level=risk_level,
                expression=expr_display,
                explanation=(
                    f"ESLint corrigió `{original_expr}` → `{expr_display}` en `{var_name}`. "
                    + explanation
                ),
                recommendation=recommendation,
                context_line=a[:120] + ('…' if len(a) > 120 else ''),
            ))

        return risks

    def scan_existing_loose(
        self,
        line: str,
    ) -> List[FrontBackRiskFinding]:
        """
        Escanea una línea AÑADIDA que AÚN contiene == o != (no corregidos por ESLint).
        Solo reporta comparaciones donde el RHS/LHS es un literal de riesgo conocido.

        Niveles:
          • BAJO  → x == null / x == undefined  (patrón idiomático JS intencional
                     que captura ambos valores falsy de referencia; se documenta
                     para visibilidad sin alarmar al revisor).
          • MEDIO → x == true/false/1/0/'1'  (coerción implícita accidental;
                     el cambio de back-end puede romper silenciosamente la lógica).
        """
        risks: List[FrontBackRiskFinding] = []

        # _RE_LOOSE tiene lookbehind/lookahead que correctamente excluye === y !==
        # sin necesidad de enmascarar; lo mantenemos para compatibilidad.
        masked = line.replace('===', '\x01\x01\x01').replace('!==', '\x02\x02\x02')
        if not self._RE_LOOSE.search(masked):
            return risks

        # _extract_comparisons prioriza === sobre == en la alternación regex,
        # por lo que devuelve los operadores correctos directamente.
        for lhs, op, rhs in self._extract_comparisons(line):
            if op not in ('==', '!='):
                continue
            literal_check = rhs.strip()
            meta = self._RISKY_LITERALS.get(literal_check)
            if not meta:
                literal_check = lhs.strip()
                meta = self._RISKY_LITERALS.get(literal_check)
                if meta:
                    lhs, rhs = rhs, lhs
            if not meta:
                continue

            _, scenario, explanation, recommendation = meta
            expr_display = f"{lhs} {op} {rhs}"
            var_name = rhs if lhs in self._RISKY_LITERALS else lhs

            # == null/undefined: patrón idiomático JS (captura ambos) → BAJO
            # Otros literales: coerción implícita accidental → MEDIO
            r_lvl = 'BAJO' if rhs in ('null', 'undefined') else 'MEDIO'

            risks.append(FrontBackRiskFinding(
                scenario=f"[Sin corregir] {scenario}",
                risk_level=r_lvl,
                expression=expr_display,
                explanation=(
                    f"La línea añadida AÚN usa `{op}` (igualdad laxa) en "
                    f"`{expr_display}` para `{var_name}`. "
                    + explanation
                ),
                recommendation=recommendation,
                context_line=line.strip()[:120] + ('…' if len(line) > 120 else ''),
            ))

        return risks


# =============================================================================
# MOTOR DE CAMBIOS ESTRUCTURALES JS/TS: GUARDAS DE NULIDAD Y EVALUACIÓN LÓGICA
# =============================================================================

class JSStructuralChangeAnalyzer:
    """
    Detecta cambios estructurales en TypeScript/JavaScript que alteran la semántica
    de evaluación de nulidad, falsy/nullish, coerción implícita o aserciones de tipo.

    Patrones cubiertos (comparando línea eliminada vs línea añadida):

      1. Guarda && → optional chaining (?.)
         'x && x.prop'  →  'x?.prop'
         Diferencia clave: && evalúa como false para CUALQUIER valor falsy (0, false,
         '', NaN, null, undefined). ?. solo protege contra null y undefined; si x es
         0 o false, el código anterior cortocircuitaba y el nuevo accede a la propiedad.

      2. Logical OR (||) → nullish coalescing (??)
         'x || default'  →  'x ?? default'
         Con ||: el default se aplica si x es cualquier falsy (0, '', false, null, undefined).
         Con ??: el default solo se aplica si x es null o undefined.
         Si x puede ser 0 (cantidad cero) o '' (campo vacío), || los reemplazaba; ?? no.

      3. Non-null assertion añadida (!.)
         'obj.prop'  →  'obj!.prop'
         Elimina el error de compilación TypeScript pero no añade guarda en runtime.
         Si obj es null/undefined en ejecución → TypeError.

      4. Eliminación de guarda typeof
         'typeof x !== "undefined"'  →  'x !== undefined'
         typeof no lanza ReferenceError con variables no declaradas; la comparación
         directa sí, si x no está en scope.

      5. Ternario de guarda nula → nullish coalescing (??)
         'x !== null ? x : d'  →  'x ?? d'
         El ternario solo guardaba null; ?? cubre también undefined (generalmente mejor).
    """

    # ── Regexes de detección ─────────────────────────────────────────────────

    # 1. && guard → ?.
    # En eliminada: EXPR && EXPR.PROP  (misma base antes del punto)
    # En añadida:   EXPR?.PROP
    _RE_AND_GUARD = re.compile(
        r'(this\.(?:[\w$]+\.?)+|[\w$]+(?:\.[\w$]+)*)'
        r'\s*&&\s*'
        r'(this\.(?:[\w$]+\.?)+|[\w$]+(?:\.[\w$]+)*)\.(\w+)'
    )
    _RE_OPT_CHAIN_EXPR = re.compile(
        r'(this\.(?:[\w$]+\.?)+|[\w$]+(?:\.[\w$]+)*)\?\.([\w$]+)'
    )

    # 2. || → ??
    _RE_OR_OP   = re.compile(r'(?<!\|)\|\|(?!\|)')
    _RE_NULLISH = re.compile(r'\?\?(?!\?)')

    # 3. Non-null assertion !.
    _RE_NON_NULL = re.compile(r'(?<=\w)!\.')

    # 4. typeof → direct comparison
    _RE_TYPEOF       = re.compile(r"typeof\s+([\w$]+(?:\.[\w$]+)*)\s*!==?\s*['\"]undefined['\"]")
    _RE_DIRECT_UNDEF = re.compile(r"([\w$]+(?:\.[\w$]+)*)\s*!==?\s*undefined")

    # 5. Ternario nulo/undefined → ??
    _RE_TERNARY_NULL  = re.compile(
        r'([\w$]+(?:\.[\w$]+)*)\s*!==?\s*null\s*\?\s*\1\s*:\s*([\w$\'".()\[\] ]+)'
    )
    _RE_TERNARY_UNDEF = re.compile(
        r'([\w$]+(?:\.[\w$]+)*)\s*!==?\s*undefined\s*\?\s*\1\s*:\s*([\w$\'".()\[\] ]+)'
    )

    # ── API pública ──────────────────────────────────────────────────────────

    def analyze_pair(
        self, removed: str, added: str
    ) -> List["JSStructuralFinding"]:
        """
        Compara una línea eliminada con su par añadida y devuelve todos los
        hallazgos estructurales detectados. Sólo analiza pares reales (ambas
        no vacías y distintas).
        """
        findings: List[JSStructuralFinding] = []
        r = removed.strip()
        a = added.strip()
        if not r or not a or r == a:
            return findings

        findings.extend(self._detect_opt_chain(r, a))
        findings.extend(self._detect_nullish_coalescing(r, a))
        findings.extend(self._detect_non_null_assertion(r, a))
        findings.extend(self._detect_typeof_removal(r, a))
        findings.extend(self._detect_ternary_to_nullish(r, a))
        return findings

    # ── Detectores internos ──────────────────────────────────────────────────

    def _detect_opt_chain(self, r: str, a: str) -> List["JSStructuralFinding"]:
        results: List[JSStructuralFinding] = []
        for m_a in self._RE_OPT_CHAIN_EXPR.finditer(a):
            base = m_a.group(1)
            base_re = re.escape(base)
            # Buscar en la eliminada: base && base.PROP
            pat = re.compile(
                rf'{base_re}\s*&&\s*{base_re}\.([\w$\[]+)'
            )
            m_r = pat.search(r)
            if not m_r:
                continue
            prop     = m_a.group(2)
            before_f = f"{base} && {base}.{m_r.group(1)}"
            after_f  = m_a.group(0)
            if prop == 'length':
                extra = (
                    f"En este caso se accede a `.length`: para arrays, el único valor "
                    f"práctico que suprime el resultado con `&&` pero no con `?.` sería "
                    f"si `{base}` fuera `0`, `false` o `''`, lo cual no ocurre con arrays "
                    f"correctamente tipados. El riesgo es bajo si TypeScript tipifica `{base}` "
                    f"como array (nunca puede ser falsy distinto de null/undefined)."
                )
            else:
                extra = (
                    f"Si `{base}` puede tener un valor falsy distinto de null/undefined "
                    f"(por ejemplo, `0`, `false` o `''`) antes de este punto, la guarda "
                    f"`&&` anterior lo bloqueaba y `?.` no lo hará, permitiendo evaluar "
                    f"`.{prop}` sobre ese valor."
                )
            results.append(JSStructuralFinding(
                pattern="Guarda && reemplazada por optional chaining (?.)",
                severity="ATENCION",
                before=before_f,
                after=after_f,
                explanation=(
                    f"La guarda explícita `{base} && {base}.{m_r.group(1)}` bloqueaba "
                    f"la evaluación para CUALQUIER valor falsy de `{base}` "
                    f"(null, undefined, 0, false, ''). "
                    f"El optional chaining `{after_f}` solo protege contra null y undefined; "
                    f"si `{base}` es `0`, `false` o `''`, el código continúa evaluando "
                    f"`.{prop}`. " + extra
                ),
                recommendation=(
                    f"Verificar que `{base}` únicamente puede ser null o undefined cuando "
                    f"no tiene valor (nunca `0`, `false` ni `''`). Si TypeScript lo tipifica "
                    f"como `T | null | undefined` sin incluir `number | boolean | string`, "
                    f"el cambio es seguro. De lo contrario, considerar mantener la guarda "
                    f"explícita o usar `{base} != null && {base}.{prop}`."
                ),
            ))
        return results

    def _detect_nullish_coalescing(self, r: str, a: str) -> List["JSStructuralFinding"]:
        results: List[JSStructuralFinding] = []
        norm_r = re.sub(r'\s+', ' ', r)
        norm_a = re.sub(r'\s+', ' ', a)
        # Si reemplazar || por ?? en la línea eliminada produce la añadida → cambio limpio
        candidate = self._RE_OR_OP.sub('??', norm_r)
        if candidate != norm_a:
            return results
        if not (self._RE_NULLISH.search(a) and self._RE_OR_OP.search(r)):
            return results
        m_or = self._RE_OR_OP.search(r)
        lhs  = r[:m_or.start()].strip()
        rhs  = r[m_or.end():].strip()
        lhs_d = ('…' + lhs[-50:]) if len(lhs) > 50 else lhs
        rhs_d = (rhs[:50] + '…') if len(rhs) > 50 else rhs
        results.append(JSStructuralFinding(
            pattern="Logical OR (||) reemplazado por nullish coalescing (??)",
            severity="ATENCION",
            before=f"{lhs_d} || {rhs_d}",
            after=f"{lhs_d} ?? {rhs_d}",
            explanation=(
                f"Con `||`, el valor por defecto `{rhs_d}` se aplica cuando la expresión "
                f"izquierda es CUALQUIER valor falsy: null, undefined, 0, false o '' "
                f"(string vacío). Con `??`, el valor por defecto solo se aplica cuando "
                f"la izquierda es null o undefined. "
                f"Si la variable puede retornar `0` (cantidad cero), `false` (flag), "
                f"o `''` (campo vacío), antes eran tratados como 'sin dato' y ahora NO lo serán: "
                f"se usará su valor real en lugar del default."
            ),
            recommendation=(
                f"Verificar que los únicos estados 'sin dato' posibles para la variable "
                f"son null y undefined (nunca 0, false ni ''). "
                f"Si esos valores son válidos de negocio (ej: cantidad=0, activo=false, "
                f"descripción=''), el cambio a `??` es correcto (evita reemplazarlos). "
                f"Si 0, false o '' también deben disparar el default, mantener `||`."
            ),
        ))
        return results

    def _detect_non_null_assertion(self, r: str, a: str) -> List["JSStructuralFinding"]:
        results: List[JSStructuralFinding] = []
        r_count = len(self._RE_NON_NULL.findall(r))
        a_count = len(self._RE_NON_NULL.findall(a))
        if a_count <= r_count:
            return results
        m = self._RE_NON_NULL.search(a)
        if not m:
            return results
        start   = max(0, m.start() - 25)
        snippet = a[start: min(len(a), m.start() + 35)].strip()
        results.append(JSStructuralFinding(
            pattern="Non-null assertion operator (!.) añadido",
            severity="ATENCION",
            before=r[:90] + ('…' if len(r) > 90 else ''),
            after=a[:90]  + ('…' if len(a) > 90 else ''),
            explanation=(
                f"El operador `!.` (non-null assertion) le indica a TypeScript que "
                f"el valor nunca es null ni undefined en ese punto, eliminando el error "
                f"de compilación. SIN EMBARGO, no añade ninguna protección en runtime: "
                f"si la asunción es incorrecta, se produce un TypeError en producción. "
                f"Contexto detectado: `{snippet}`."
            ),
            recommendation=(
                "Confirmar que la variable está garantizadamente inicializada antes de "
                "este punto (por ejemplo, se asigna en el constructor, dentro de un *ngIf, "
                "o después de una guarda explícita). Si hay duda, preferir optional chaining "
                "`x?.prop` o una guarda `if (x !== null) { ... }` en lugar de `x!.prop`."
            ),
        ))
        return results

    def _detect_typeof_removal(self, r: str, a: str) -> List["JSStructuralFinding"]:
        results: List[JSStructuralFinding] = []
        m_r = self._RE_TYPEOF.search(r)
        m_a = self._RE_DIRECT_UNDEF.search(a)
        if not (m_r and m_a and m_r.group(1) == m_a.group(1)):
            return results
        varname = m_r.group(1)
        results.append(JSStructuralFinding(
            pattern="Guarda typeof eliminada — comparación directa con undefined",
            severity="INFORMATIVO",
            before=m_r.group(0),
            after=m_a.group(0),
            explanation=(
                f"`typeof {varname} !== 'undefined'` funciona incluso si `{varname}` no "
                f"está declarada en el scope (devuelve 'undefined' sin lanzar error). "
                f"La comparación directa `{varname} !== undefined` lanzará ReferenceError "
                f"en runtime si `{varname}` no está declarada como variable o propiedad."
            ),
            recommendation=(
                f"Si `{varname}` siempre está declarada (propiedad de clase, variable de módulo "
                f"garantizada, parámetro de función), el cambio es seguro y más legible. "
                f"Si es una variable global o de entorno cuya existencia no está garantizada "
                f"(ej: `window.myPlugin`), mantener `typeof {varname} !== 'undefined'`."
            ),
        ))
        return results

    def _detect_ternary_to_nullish(self, r: str, a: str) -> List["JSStructuralFinding"]:
        results: List[JSStructuralFinding] = []
        if not self._RE_NULLISH.search(a):
            return results
        for regex, guarded, other in (
            (self._RE_TERNARY_NULL,  'null',      'undefined'),
            (self._RE_TERNARY_UNDEF, 'undefined', 'null'),
        ):
            m_r = regex.search(r)
            if not m_r:
                continue
            varname     = m_r.group(1)
            default_val = m_r.group(2).strip()
            pat = re.compile(
                rf'{re.escape(varname)}\s*\?\?\s*{re.escape(default_val)}'
            )
            if not pat.search(a):
                continue
            results.append(JSStructuralFinding(
                pattern="Ternario de guarda nula reemplazado por nullish coalescing (??)",
                severity="INFORMATIVO",
                before=m_r.group(0),
                after=f"{varname} ?? {default_val}",
                explanation=(
                    f"El ternario `{m_r.group(0)}` solo guardaba {guarded}: si `{varname}` "
                    f"era {other}, se devolvía `{varname}` tal cual (no se aplicaba el default). "
                    f"`??` cubre ambos: null Y undefined. "
                    f"Esto es generalmente una mejora al alinear ambos valores de 'sin dato', "
                    f"pero si el ternario original era intencional (solo {guarded} debía usar "
                    f"el default y {other} debía propagarse), el comportamiento cambia."
                ),
                recommendation=(
                    f"Confirmar que tanto null como undefined deben resolverse con "
                    f"`{default_val}`. En la mayoría de APIs REST esto es correcto: un campo "
                    f"ausente (undefined) y un campo nulo (null) merecen el mismo valor por "
                    f"defecto. Si no, restaurar el ternario con guarda específica."
                ),
            ))
            break
        return results


class ESLintAngularAnalyzer:
    """
    Motor de análisis ESLint para proyectos Angular 16 strict mode.

    Archivos .html  → reglas 16, 17, 18  (template/recommended)
    Archivos .ts    → reglas  1-15       (@angular-eslint/recommended)

    Reglas NO activas (no se reportan): no-call-expression, no-any,
    no-conflicting-lifecycle, @typescript-eslint/**, accesibilidad.
    """

    # ── Patrones HTML ────────────────────────────────────────────────────────
    # Extrae contenido entre comillas de atributos Angular y de interpolaciones
    _RE_EXPR_ATTR   = re.compile(
        r'(?:[\[(*][\w.@$-]+\]?|\*\w+|#\w+)\s*=\s*"([^"]*)"'
    )
    _RE_EXPR_INTERP = re.compile(r'\{\{([^}]+)\}\}')

    # Regla 17: == o != que NO sean === ni !==
    # Lookbehind/lookahead para evitar capturar ===, !==, <=, >=
    _RE_EQEQ        = re.compile(r'(?<![=!<>])={2}(?!=)|(?<![!<>])!={1}(?!=)')
    _RE_STRICT      = re.compile(r'===|!==')

    # Regla 16: banana-in-box — patrón ([attr])=
    _RE_BANANA      = re.compile(r'\(\[[\w.]+\]\)\s*=')

    # Regla 18: no-negated-async — !(expr | async)
    _RE_NEG_ASYNC   = re.compile(r'!\s*\([^)]*\|\s*async\s*\)')

    # Regla 19: accessibility-table-scope — <th> debe tener scope="col"|"row"|"colgroup"|"rowgroup"
    # Incluida en template/recommended v16.3.1; el preset accessibility está comentado en .eslintrc.json
    # pero esta regla específica ya viene en recommended desde v15+.
    _RE_TH_SCOPE    = re.compile(r'<th\b[^>]*\bscope\s*=', re.I)
    _RE_TH_TAG      = re.compile(r'<th\b', re.I)

    # ── Patrones TS ──────────────────────────────────────────────────────────
    _RE_COMP_DECOR  = re.compile(r'@Component\s*\(')
    _RE_DIR_DECOR   = re.compile(r'@Directive\s*\(')
    _RE_PIPE_DECOR  = re.compile(r'@Pipe\s*\(')

    # Regla 4: método de ciclo de vida con cuerpo vacío
    _RE_EMPTY_HOOK  = re.compile(
        r'\b(' + '|'.join(LIFECYCLE_HOOKS) + r')\s*\([^)]*\)\s*:\s*\w*\s*\{\s*\}|\b('
        + '|'.join(LIFECYCLE_HOOKS) + r')\s*\([^)]*\)\s*\{\s*\}'
    )

    # Regla 5: host property en decorador
    _RE_HOST_PROP   = re.compile(r'\bhost\s*:\s*\{')

    # Regla 6: @Input con alias (cualquier argumento de cadena)
    _RE_INPUT_ALIAS  = re.compile(r"@Input\s*\(\s*['\"](\w+)['\"]\s*\)")
    # Regla 10: @Output con alias
    _RE_OUTPUT_ALIAS = re.compile(r"@Output\s*\(\s*['\"](\w+)['\"]\s*\)")

    # Reglas 7/11: inputs/outputs en metadata del decorador
    _RE_INPUTS_META  = re.compile(r'\binputs\s*:\s*\[')
    _RE_OUTPUTS_META = re.compile(r'\boutputs\s*:\s*\[')

    # Regla 8: nombre de la variable después de @Output()
    _RE_OUTPUT_DECL  = re.compile(
        r'@Output\s*\(\s*\)\s+(?:public\s+|private\s+|protected\s+|readonly\s+)?(\w+)'
    )

    # Regla 9: @Output cuyo nombre empieza por "on"
    _RE_OUTPUT_ON    = re.compile(
        r'@Output\s*\([^)]*\)\s+(?:public\s+|private\s+|protected\s+|readonly\s+)?on[A-Z]\w*'
    )

    # Regla 12: método de ciclo de vida declarado en la clase
    _RE_HOOK_METHOD  = re.compile(r'\b(' + '|'.join(LIFECYCLE_HOOKS) + r')\s*\(')
    _RE_IMPLEMENTS   = re.compile(r'\bimplements\b([^{]+)')

    # Regla 13: @Pipe sin PipeTransform
    _RE_PIPE_TRANSFORM = re.compile(r'\bimplements\b[^{]*\bPipeTransform\b')

    # Reglas 14/15: selector
    _RE_SELECTOR     = re.compile(r"\bselector\s*:\s*['\"]([^'\"]+)['\"]")

    # Sufijos obligatorios (reglas 1 y 3)
    _RE_CLASS_DEF    = re.compile(r'\bclass\s+(\w+)')

    # ── Helpers internos ─────────────────────────────────────────────────────

    def _has_loose_equality(self, text: str) -> bool:
        """Devuelve True si 'text' contiene == o != no estrictos."""
        masked = self._RE_STRICT.sub('\x00\x00\x00', text)
        return bool(self._RE_EQEQ.search(masked))

    def _extract_angular_expressions(self, line: str) -> List[str]:
        """Extrae el contenido de atributos Angular e interpolaciones."""
        exprs: List[str] = []
        for m in self._RE_EXPR_ATTR.finditer(line):
            exprs.append(m.group(1))
        for m in self._RE_EXPR_INTERP.finditer(line):
            exprs.append(m.group(1))
        return exprs

    def _line_has_eqeqeq(self, line: str) -> bool:
        exprs = self._extract_angular_expressions(line)
        return any(self._has_loose_equality(e) for e in exprs)

    def _best_removed_match(self, added_l: str, removed_lines: List[str]) -> str:
        """
        Encuentra la línea eliminada más similar a la añadida (ratio > 0.5).

        Estrategia doble:
          1. Character-set overlap ratio  → funciona bien para líneas cortas.
          2. Normalización de operadores  → detecta líneas largas que solo difieren
             en == vs ===  (correcciones ESLint eqeqeq en templates HTML largos).
        """
        if not removed_lines:
            return ""

        # Normaliza operadores de comparación para igualar semánticamente
        _norm_ops = lambda s: re.sub(r'===', '==', re.sub(r'!==', '!=', s))

        a_ns = re.sub(r'\s+', '', added_l)
        a_ns_norm = _norm_ops(a_ns)

        best, best_score = "", 0.0
        for r_l in removed_lines:
            r_ns = re.sub(r'\s+', '', r_l)

            # Estrategia 1: set-overlap (para líneas cortas)
            common = sum(1 for c in set(a_ns) if c in r_ns)
            total  = max(len(a_ns), len(r_ns), 1)
            score  = common / total

            # Estrategia 2: coincidencia exacta tras normalizar == vs ===
            # Cubre el caso: línea HTML larga con == → === (la única diferencia)
            r_ns_norm = _norm_ops(r_ns)
            if r_ns_norm == a_ns_norm and len(a_ns_norm) > 10:
                score = max(score, 0.95)
            elif len(a_ns) > 50:
                # Para líneas largas: comparar caracteres ordenados (multiset sim.)
                common2 = sum(min(a_ns.count(c), r_ns.count(c)) for c in set(a_ns))
                total2  = max(len(a_ns), len(r_ns), 1)
                score = max(score, common2 / total2)

            if score > best_score:
                best_score, best = score, r_l
        return best if best_score > 0.50 else ""

    # ── Análisis HTML ────────────────────────────────────────────────────────

    def analyze_html_line(
        self,
        added: str,
        removed: str,
        fb_engine: Optional["FrontBackTypeRiskEngine"] = None,
        fb_risks_out: Optional[List["FrontBackRiskFinding"]] = None,
    ) -> Optional[ESLintFinding]:
        a, r = added.strip(), removed.strip()

        # Regla 17 – eqeqeq
        a_has = self._line_has_eqeqeq(a)
        r_has = self._line_has_eqeqeq(r) if r else False
        if r_has and not a_has:
            # ── Motor Front-Back: analizar qué se comparaba y qué riesgo genera ──
            if fb_engine is not None and fb_risks_out is not None:
                fb_risks_out.extend(fb_engine.analyze_line_pair(r, a))
            return ESLintFinding(
                rule="@angular-eslint/template/eqeqeq", severity="ERROR",
                category="CORRECCION", line_added=a, line_removed=r,
                description="Corrección de igualdad estricta (== → === o != → !==) en expresión de plantilla.",
            )
        if a_has and not r_has:
            # Violación nueva: también capturar riesgo de coerción
            if fb_engine is not None and fb_risks_out is not None:
                fb_risks_out.extend(fb_engine.scan_existing_loose(a))
            return ESLintFinding(
                rule="@angular-eslint/template/eqeqeq", severity="ERROR",
                category="VIOLACION", line_added=a, line_removed=r,
                description="Uso de == o != en expresión de plantilla Angular. Debe usarse === o !==.",
            )

        # Regla 16 – banana-in-box
        a_ban = bool(self._RE_BANANA.search(a))
        r_ban = bool(self._RE_BANANA.search(r)) if r else False
        if a_ban and not r_ban:
            return ESLintFinding(
                rule="@angular-eslint/template/banana-in-box", severity="ERROR",
                category="VIOLACION", line_added=a, line_removed=r,
                description="Two-way binding incorrecto: se usa ([attr])= en lugar de [(attr)]=.",
            )
        if r_ban and not a_ban:
            return ESLintFinding(
                rule="@angular-eslint/template/banana-in-box", severity="ERROR",
                category="CORRECCION", line_added=a, line_removed=r,
                description="Corrección de two-way binding: sintaxis ([]) reemplazada por [()].",
            )

        # Regla 18 – no-negated-async
        a_neg = bool(self._RE_NEG_ASYNC.search(a))
        r_neg = bool(self._RE_NEG_ASYNC.search(r)) if r else False
        if a_neg and not r_neg:
            return ESLintFinding(
                rule="@angular-eslint/template/no-negated-async", severity="ERROR",
                category="VIOLACION", line_added=a, line_removed=r,
                description="Negación directa sobre pipe async ( !(obs$ | async) ). Usar variable con 'as'.",
            )
        if r_neg and not a_neg:
            return ESLintFinding(
                rule="@angular-eslint/template/no-negated-async", severity="ERROR",
                category="CORRECCION", line_added=a, line_removed=r,
                description="Corrección: se elimina el patrón !(... | async) del template.",
            )

        # Regla 19 – accessibility-table-scope
        # Activa: viene en template/recommended v16.3.1 (el preset accessibility
        # está comentado en .eslintrc.json pero esta regla ya pertenece a recommended).
        a_th = bool(self._RE_TH_TAG.search(a))
        r_th = bool(self._RE_TH_TAG.search(r)) if r else False
        a_sc = bool(self._RE_TH_SCOPE.search(a)) if a_th else False
        r_sc = bool(self._RE_TH_SCOPE.search(r)) if r_th else False

        if a_th and a_sc and r_th and not r_sc:
            return ESLintFinding(
                rule="@angular-eslint/template/accessibility-table-scope", severity="ERROR",
                category="CORRECCION", line_added=a, line_removed=r,
                description=(
                    "Corrección: se añade scope=\"col\"/\"row\" al elemento <th> "
                    "para cumplir accesibilidad de tablas HTML."
                ),
            )
        if a_th and not a_sc and not (r_th and r_sc):
            return ESLintFinding(
                rule="@angular-eslint/template/accessibility-table-scope", severity="ERROR",
                category="VIOLACION", line_added=a, line_removed=r,
                description="Elemento <th> sin atributo scope. Debe añadirse scope=\"col\" o scope=\"row\".",
            )

        return None

    # ── Análisis TypeScript ──────────────────────────────────────────────────

    def analyze_ts_line(
        self,
        added: str,
        removed: str,
        full_content: List[str],
    ) -> Optional[ESLintFinding]:
        a, r = added.strip(), removed.strip()
        full_text = "\n".join(full_content)

        # Regla 4 – no-empty-lifecycle-method
        if self._RE_EMPTY_HOOK.search(a) and not self._RE_EMPTY_HOOK.search(r):
            return ESLintFinding(
                rule="@angular-eslint/no-empty-lifecycle-method", severity="ERROR",
                category="VIOLACION", line_added=a, line_removed=r,
                description="Método de ciclo de vida declarado con cuerpo vacío {}.",
            )
        if self._RE_EMPTY_HOOK.search(r) and not self._RE_EMPTY_HOOK.search(a):
            return ESLintFinding(
                rule="@angular-eslint/no-empty-lifecycle-method", severity="ERROR",
                category="CORRECCION", line_added=a, line_removed=r,
                description="Corrección: se elimina o implementa el método de ciclo de vida vacío.",
            )

        # Regla 5 – no-host-metadata-property
        if self._RE_HOST_PROP.search(a) and not self._RE_HOST_PROP.search(r):
            return ESLintFinding(
                rule="@angular-eslint/no-host-metadata-property", severity="ERROR",
                category="VIOLACION", line_added=a, line_removed=r,
                description="Propiedad 'host:' en decorador. Usar @HostListener / @HostBinding.",
            )
        if self._RE_HOST_PROP.search(r) and not self._RE_HOST_PROP.search(a):
            return ESLintFinding(
                rule="@angular-eslint/no-host-metadata-property", severity="ERROR",
                category="CORRECCION", line_added=a, line_removed=r,
                description="Corrección: se elimina propiedad 'host:' del decorador.",
            )

        # Regla 6 – no-input-rename
        m_in_a = self._RE_INPUT_ALIAS.search(a)
        m_in_r = self._RE_INPUT_ALIAS.search(r) if r else None
        if m_in_a and not m_in_r:
            return ESLintFinding(
                rule="@angular-eslint/no-input-rename", severity="ERROR",
                category="VIOLACION", line_added=a, line_removed=r,
                description=f"@Input() con alias '{m_in_a.group(1)}'. El alias debe coincidir con el nombre de la propiedad.",
            )
        if m_in_r and not m_in_a:
            return ESLintFinding(
                rule="@angular-eslint/no-input-rename", severity="ERROR",
                category="CORRECCION", line_added=a, line_removed=r,
                description=f"Corrección: se elimina alias '{m_in_r.group(1)}' de @Input().",
            )

        # Regla 7 – no-inputs-metadata-property
        if self._RE_INPUTS_META.search(a) and not self._RE_INPUTS_META.search(r):
            return ESLintFinding(
                rule="@angular-eslint/no-inputs-metadata-property", severity="ERROR",
                category="VIOLACION", line_added=a, line_removed=r,
                description="Propiedad 'inputs:' en decorador. Usar @Input() en cada campo.",
            )
        if self._RE_INPUTS_META.search(r) and not self._RE_INPUTS_META.search(a):
            return ESLintFinding(
                rule="@angular-eslint/no-inputs-metadata-property", severity="ERROR",
                category="CORRECCION", line_added=a, line_removed=r,
                description="Corrección: se elimina propiedad 'inputs:' del decorador.",
            )

        # Regla 8 – no-output-native
        m_out_decl = self._RE_OUTPUT_DECL.search(a)
        if m_out_decl:
            out_name = m_out_decl.group(1).strip()
            if out_name.lower() in DOM_NATIVE_EVENTS:
                return ESLintFinding(
                    rule="@angular-eslint/no-output-native", severity="ERROR",
                    category="VIOLACION", line_added=a, line_removed=r,
                    description=f"@Output() '{out_name}' colisiona con evento DOM nativo.",
                )

        # Regla 9 – no-output-on-prefix
        if self._RE_OUTPUT_ON.search(a) and not self._RE_OUTPUT_ON.search(r):
            return ESLintFinding(
                rule="@angular-eslint/no-output-on-prefix", severity="ERROR",
                category="VIOLACION", line_added=a, line_removed=r,
                description="@Output() cuyo nombre empieza por 'on'. Renombrar sin el prefijo.",
            )
        if self._RE_OUTPUT_ON.search(r) and not self._RE_OUTPUT_ON.search(a):
            return ESLintFinding(
                rule="@angular-eslint/no-output-on-prefix", severity="ERROR",
                category="CORRECCION", line_added=a, line_removed=r,
                description="Corrección: se elimina prefijo 'on' del @Output().",
            )

        # Regla 10 – no-output-rename
        m_out_a = self._RE_OUTPUT_ALIAS.search(a)
        m_out_r = self._RE_OUTPUT_ALIAS.search(r) if r else None
        if m_out_a and not m_out_r:
            return ESLintFinding(
                rule="@angular-eslint/no-output-rename", severity="ERROR",
                category="VIOLACION", line_added=a, line_removed=r,
                description=f"@Output() con alias '{m_out_a.group(1)}'. El alias debe coincidir con el nombre de la propiedad.",
            )
        if m_out_r and not m_out_a:
            return ESLintFinding(
                rule="@angular-eslint/no-output-rename", severity="ERROR",
                category="CORRECCION", line_added=a, line_removed=r,
                description=f"Corrección: se elimina alias '{m_out_r.group(1)}' de @Output().",
            )

        # Regla 11 – no-outputs-metadata-property
        if self._RE_OUTPUTS_META.search(a) and not self._RE_OUTPUTS_META.search(r):
            return ESLintFinding(
                rule="@angular-eslint/no-outputs-metadata-property", severity="ERROR",
                category="VIOLACION", line_added=a, line_removed=r,
                description="Propiedad 'outputs:' en decorador. Usar @Output() en cada campo.",
            )
        if self._RE_OUTPUTS_META.search(r) and not self._RE_OUTPUTS_META.search(a):
            return ESLintFinding(
                rule="@angular-eslint/no-outputs-metadata-property", severity="ERROR",
                category="CORRECCION", line_added=a, line_removed=r,
                description="Corrección: se elimina propiedad 'outputs:' del decorador.",
            )

        # Regla 12 – use-lifecycle-interface
        m_hook = self._RE_HOOK_METHOD.search(a)
        if m_hook:
            hook_name = m_hook.group(1)
            if hook_name not in r:
                iface = LIFECYCLE_INTERFACES.get(hook_name, "")
                if iface:
                    impl_match = self._RE_IMPLEMENTS.search(full_text)
                    if not impl_match or iface not in impl_match.group(1):
                        return ESLintFinding(
                            rule="@angular-eslint/use-lifecycle-interface", severity="ERROR",
                            category="VIOLACION", line_added=a, line_removed=r,
                            description=f"Método '{hook_name}' sin declarar 'implements {iface}' en la clase.",
                        )

        # Regla 13 – use-pipe-transform-interface
        if self._RE_PIPE_DECOR.search(a):
            if not self._RE_PIPE_TRANSFORM.search(full_text):
                return ESLintFinding(
                    rule="@angular-eslint/use-pipe-transform-interface", severity="ERROR",
                    category="VIOLACION", line_added=a, line_removed=r,
                    description="Clase con @Pipe sin declarar 'implements PipeTransform'.",
                )

        # Reglas 1/3 – component-class-suffix / directive-class-suffix
        m_cls = self._RE_CLASS_DEF.search(a)
        if m_cls:
            cls_name = m_cls.group(1)
            is_comp = self._RE_COMP_DECOR.search(full_text)
            is_dir  = self._RE_DIR_DECOR.search(full_text)
            if is_comp and not cls_name.endswith('Component'):
                return ESLintFinding(
                    rule="@angular-eslint/component-class-suffix", severity="ERROR",
                    category="VIOLACION", line_added=a, line_removed=r,
                    description=f"La clase '{cls_name}' debe terminar en 'Component'.",
                )
            if is_dir and not cls_name.endswith('Directive'):
                return ESLintFinding(
                    rule="@angular-eslint/directive-class-suffix", severity="ERROR",
                    category="VIOLACION", line_added=a, line_removed=r,
                    description=f"La clase '{cls_name}' debe terminar en 'Directive'.",
                )

        # Reglas 14/15 – component-selector / directive-selector
        m_sel = self._RE_SELECTOR.search(a)
        if m_sel:
            sel = m_sel.group(1).strip()
            is_comp = bool(self._RE_COMP_DECOR.search(full_text))
            is_dir  = bool(self._RE_DIR_DECOR.search(full_text))
            if is_comp and not re.match(r'^app-[a-z][a-z0-9-]*$', sel):
                return ESLintFinding(
                    rule="@angular-eslint/component-selector", severity="ERROR",
                    category="VIOLACION", line_added=a, line_removed=r,
                    description=f"Selector de componente '{sel}' no cumple: prefijo 'app-', kebab-case, tipo elemento.",
                )
            if is_dir and not re.match(r'^app[A-Z][a-zA-Z0-9]*$', sel):
                return ESLintFinding(
                    rule="@angular-eslint/directive-selector", severity="ERROR",
                    category="VIOLACION", line_added=a, line_removed=r,
                    description=f"Selector de directiva '{sel}' no cumple: prefijo 'app', camelCase, tipo atributo.",
                )

        return None

    # ── Análisis de archivo completo ─────────────────────────────────────────

    def analyze_file(self, fc: FileChange) -> ESLintFileReport:
        """
        Analiza todas las líneas añadidas de un FileChange.
        Empareja cada línea + con la línea - más similar del mismo hunk
        para distinguir correcciones de violaciones nuevas.
        También ejecuta el motor FrontBackTypeRiskEngine para detectar
        riesgos de incompatibilidad de tipos front-back en comparaciones,
        y el JSStructuralChangeAnalyzer para detectar cambios estructurales
        (optional chaining, nullish coalescing, non-null assertion, etc.).
        """
        report   = ESLintFileReport()
        fb_engine = FrontBackTypeRiskEngine()
        structural_analyzer = JSStructuralChangeAnalyzer()
        ext = fc.ext.lower()
        is_html = ext in ('.html', '.component.html')
        is_ts   = ext in (
            '.ts', '.component.ts', '.service.ts', '.pipe.ts',
            '.directive.ts', '.guard.ts', '.interceptor.ts',
            '.module.ts', '.resolver.ts',
        )

        if not (is_html or is_ts):
            report.functional = [l for l in fc.added]
            return report

        removed_pool = list(fc.removed)
        used_findings: Set[Tuple[str, str, str]] = set()
        fb_risks_acc: List[FrontBackRiskFinding] = []  # acumulador de riesgos front-back
        structural_acc: List[JSStructuralFinding] = []  # acumulador de cambios estructurales

        for added_l in fc.added:
            a_stripped = added_l.strip()
            paired_r   = self._best_removed_match(added_l, removed_pool)

            # ── Motor estructural JS/TS: optional chaining, ??, !., typeof ──
            if paired_r:
                structural_acc.extend(structural_analyzer.analyze_pair(paired_r, added_l))

            finding: Optional[ESLintFinding] = None
            if is_html:
                finding = self.analyze_html_line(
                    added_l, paired_r,
                    fb_engine=fb_engine, fb_risks_out=fb_risks_acc,
                )
            elif is_ts:
                finding = self.analyze_ts_line(added_l, paired_r, fc.full_content)
                # Motor Front-Back para TypeScript también
                if finding and finding.rule == "@angular-eslint/template/eqeqeq":
                    if finding.category == 'CORRECCION':
                        fb_risks_acc.extend(fb_engine.analyze_line_pair(paired_r, added_l))
                    else:
                        fb_risks_acc.extend(fb_engine.scan_existing_loose(added_l))

            if finding:
                key = (finding.rule, finding.category, a_stripped[:120])
                if key not in used_findings:
                    used_findings.add(key)
                    if finding.category == 'CORRECCION':
                        report.corrections.append(finding)
                    else:
                        report.violations.append(finding)
            else:
                report.functional.append(a_stripped)
                # Escanear riesgos en líneas funcionales (no clasificadas como ESLint)
                fb_risks_acc.extend(fb_engine.scan_existing_loose(added_l))

        # ── Detectar si las líneas ELIMINADAS tenían patrones ESLint ─────────
        for removed_l in fc.removed:
            test_finding: Optional[ESLintFinding] = None
            if is_html:
                test_finding = self.analyze_html_line(removed_l, "")
            elif is_ts:
                test_finding = self.analyze_ts_line(removed_l, "", fc.full_content)
            if test_finding:
                report.removed_had_violations = True
                break

        # ── Consolidar riesgos front-back (deduplicar por expression) ─────────
        seen_expr: Set[str] = set()
        for risk in fb_risks_acc:
            if risk.expression not in seen_expr:
                seen_expr.add(risk.expression)
                report.front_back_risks.append(risk)

        # ── Consolidar hallazgos estructurales (deduplicar por before+after) ──
        seen_struct: Set[Tuple[str, str]] = set()
        for sf in structural_acc:
            key_s = (sf.before, sf.after)
            if key_s not in seen_struct:
                seen_struct.add(key_s)
                report.structural_findings.append(sf)

        return report


def analyze_technical_impact(fc: FileChange) -> str:
    """
    FIX BUG 5: Archivos de entorno analizan full_content para detectar
    patrones criticos que estan en lineas de contexto (sin + ni -).
    """
    if fc.is_binary:
        return "Actualizacion de archivo binario (imagen, fuente, recurso)"
    if fc.is_lockfile:
        n = len(fc.added)
        d = len(fc.removed)
        return f"Actualizacion de dependencias: +{n} entradas nuevas, -{d} eliminadas"

    search_lines = fc.added + fc.removed
    if not search_lines:
        search_lines = fc.full_content[:40]
    search_text = "\n".join(search_lines)

    found = []
    for pattern, description in IMPACT_SIGNALS:
        if pattern.search(search_text):
            found.append(description)

    if found:
        return " | ".join(found[:4])

    category = get_category(fc)
    if len(fc.added) == 0 and len(fc.removed) == 0:
        return f"Ajuste de propiedades o permisos en {category}"
    if fc.kind == 'added':
        return f"Implementacion inicial de {category}"
    if fc.kind == 'deleted':
        return f"Eliminacion completa de {category}"
    return f"Modificacion de logica interna en {category}"


def calculate_deploy_impact(fc: FileChange) -> str:
    """
    FIX BUG 4: Regex de AuthService mas preciso (sin falsos positivos en rutas).
    FIX BUG 5: Archivos de entorno analizan full_content para score correcto.
    """
    if fc.is_binary:
        return "Leve"
    if fc.is_lockfile:
        return "Leve"

    changed_lines = "\n".join(fc.added + fc.removed)
    all_lines = changed_lines if changed_lines.strip() else "\n".join(fc.full_content[:60])

    n_add = len(fc.added)
    n_del = len(fc.removed)
    score = 0

    # --- CRITICO ---
    if re.search(r'\bAuthService\b|\bAuthGuard\b|jsonwebtoken|\.verify\s*\(.*token', all_lines):
        score += 60
    if re.search(r'(?<!\w)password(?!\w)|\bbcrypt\b|\bsalt\b', all_lines, re.I):
        score += 50
    if re.search(r'migration|alembic|flyway|DROP TABLE|ALTER TABLE', all_lines, re.I):
        score += 70

    # --- ALTO: entornos con URLs reales ---
    if fc.is_environment_file and re.search(r'apiLoginUrl|apiBackend|apiUrl\b', all_lines, re.I):
        score += 45

    # --- BAJO/MEDIO: nuevos endpoints o rutas de recursos en env ---
    if fc.is_environment_file:
        env_added = "\n".join(fc.added)
        env_entry_pattern = re.compile(r"\b[a-zA-Z_]\w*\s*:\s*['\"][^'\"]+['\"]")
        endpoint_hint = re.compile(r'api|url|endpoint|catalogo|mantenimiento|parametro|service', re.I)
        env_entries = env_entry_pattern.findall(env_added)
        if env_entries:
            score += 10
            if any(endpoint_hint.search(e) for e in env_entries):
                score += 8

    # --- ALTO: contratos publicos ---
    if re.search(r'\bexport\s+(interface|type|class|enum)\b', all_lines):
        score += 40
    if re.search(r'@Input\s*\(\)|@Output\s*\(\)|EventEmitter', all_lines):
        score += 35
    if re.search(r'@NgModule\s*\(|providers\s*:\s*\[|declarations\s*:\s*\[', all_lines):
        score += 35
    if re.search(r'dispatch\s*\(|createAction|createReducer|createEffect', all_lines):
        score += 30

    # --- MEDIO ---
    if re.search(r'HttpClient|http\.(get|post|put|delete|patch)', all_lines, re.I):
        score += 30
    if re.search(r'catchError|throwError|try\s*{|except\s+', all_lines):
        score += 20
    if re.search(r'SELECT\b|INSERT\b|UPDATE\b|DELETE\b|\.save\s*\(', all_lines, re.I):
        score += 25
    if re.search(r'router\.navigate|canActivate|canLoad|authGuard|menuGuard', all_lines):
        score += 20
    if re.search(r'cron|scheduler|setTimeout|setInterval', all_lines):
        score += 20
    if re.search(r'WebSocket|socket\.io', all_lines, re.I):
        score += 25

    # --- BAJO: UI y estilos ---
    if fc.ext in ('.scss', '.css', '.component.scss'):
        score += 5
    if fc.ext in ('.html', '.component.html'):
        score += 10

    # --- LEVE: tests, docs ---
    if fc.ext == '.spec.ts':
        score = min(score + 3, 15)
    if fc.ext == '.md':
        return "Nula"

    # --- PHP / Laravel: scoring por contexto del archivo ---
    if fc.ext == '.php':
        filepath_lower = fc.filepath.lower().replace('\\', '/')
        score += 20  # base: cualquier archivo PHP en produccion tiene impacto

        # Controlador de API: modifica contratos de respuesta
        if re.search(r'controller', filepath_lower):
            score += 15
        # Middleware: intercepta todas las peticiones
        if re.search(r'middleware', filepath_lower):
            score += 35
        # Migracion: altera el esquema de base de datos
        if re.search(r'migration[s]?/', filepath_lower):
            score += 55
        # Rutas: define o modifica endpoints publicos
        if re.search(r'routes?/(api|web|channels|console)\.php', filepath_lower):
            score += 30
        # Configuracion de framework
        if re.search(r'config/', filepath_lower):
            score += 25
        # Seeder / Factory: populate de datos, riesgo bajo
        if re.search(r'seeder|factory', filepath_lower):
            score = min(score, 25)
        # Vista Blade: solo UI, impacto reducido
        if '.blade.' in filepath_lower:
            score = min(score, 20)

        # Bonus por signales en el codigo cambiado
        if re.search(r'->where\s*\(|->select\s*\(|->save\s*\(|->create\s*\(|->update\s*\(|->delete\s*\(', all_lines, re.I):
            score += 20  # operaciones ORM directas
        if re.search(r'return\s+response\s*\(|->json\s*\(|\$this->success\b|\$this->error\b', all_lines, re.I):
            score += 15  # modifica respuesta de la API
        if re.search(r'\?\?|\?->|\?\?=', all_lines):
            score += 10  # null guards (correccion de bug potencial)
        if re.search(r'\blog::|try\s*\{|catch\s*\(\$e\b|catch\s*\(Exception', all_lines, re.I):
            score += 10  # manejo de excepciones
        if re.search(r'auth\s*\(|can\s*\(|Gate::|\bpolicy\b', all_lines, re.I):
            score += 20  # logica de autorizacion
        if re.search(r'->belongsTo\s*\(|->hasMany\s*\(|->hasOne\s*\(|->belongsToMany\s*\(', all_lines):
            score += 12  # relaciones entre modelos
        if re.search(r'->paginate\s*\(|->get\s*\(|->first\s*\(|->with\s*\(', all_lines, re.I):
            score += 8   # consulta de coleccion

    # Bonus por magnitud
    total_lines = n_add + n_del
    if total_lines > 200:   score += 20
    elif total_lines > 100: score += 12
    elif total_lines > 50:  score += 6
    elif total_lines > 20:  score += 3
    elif total_lines > 5:   score += 1  # cambios quirurgicos siguen puntuando

    # Archivo nuevo: cap de impacto
    if fc.kind == 'added' and score < 40:
        score = min(score, 30)
    # Archivo eliminado: riesgo de dependencias rotas
    if fc.kind == 'deleted':
        score += 20

    # Solo linter fixes
    if fc.removed and not fc.added:
        if all(fc.detect_linter_fix(l) for l in fc.removed if l.strip()):
            return "Nula"

    # Cambios muy pequenos en environment no deben inflarse por contexto completo
    critical_change = bool(re.search(
        r'\bAuthService\b|\bAuthGuard\b|jsonwebtoken|\.verify\s*\(.*token|(?<!\w)password(?!\w)|\bbcrypt\b|\bsalt\b|migration|alembic|flyway|DROP TABLE|ALTER TABLE',
        changed_lines,
        re.I,
    ))
    total_lines = n_add + n_del
    if fc.is_environment_file and total_lines <= 5 and not critical_change:
        score = min(score, 28)

    if score == 0:        return "Nula"
    elif score <= 10:     return "Leve"
    elif score <= 28:     return "Baja"
    elif score <= 50:     return "Media"
    elif score <= 75:     return "Alta"
    else:                 return "Critica"


# =============================================================================
# RESUMEN DE LOCK FILE
# =============================================================================
def summarize_lockfile(fc: FileChange) -> Dict[str, List[str]]:
    """Extrae nombres de paquetes nuevos/eliminados de un lock file."""
    added_pkgs:   List[str] = []
    removed_pkgs: List[str] = []
    pkg_pattern = re.compile(r'"(@?[\w/@.-]+)":\s*\{|^(@?[\w/@.-]+)@[\d.^~]')

    for line in fc.added:
        m = pkg_pattern.search(line.strip())
        if m:
            name = (m.group(1) or m.group(2) or '').strip('"')
            if name and name not in added_pkgs:
                added_pkgs.append(name)

    for line in fc.removed:
        m = pkg_pattern.search(line.strip())
        if m:
            name = (m.group(1) or m.group(2) or '').strip('"')
            if name and name not in removed_pkgs:
                removed_pkgs.append(name)

    return {"added": added_pkgs[:20], "removed": removed_pkgs[:20]}


# =============================================================================
# RECOMENDACIONES
# =============================================================================
def analyze_recommendations(changes: List[FileChange]) -> List[str]:
    recs = []
    all_removed = "\n".join(l for f in changes for l in f.removed).lower()

    if 'console.' in all_removed or 'print(' in all_removed:
        recs.append(
            "Se limpiaron logs de depuracion. Verificar que no falten "
            "trazas criticas en produccion."
        )

    env_changes = [f for f in changes if f.is_environment_file]
    if env_changes:
        ambientes = [f.filename for f in env_changes]
        recs.append(
            f"Se modificaron {len(env_changes)} archivo(s) de entorno "
            f"({', '.join(ambientes)}). Confirmar que los endpoints sean "
            f"correctos en cada ambiente antes del despliegue."
        )

    if any(f.is_lockfile for f in changes):
        recs.append(
            "Se actualizo el archivo de dependencias. Ejecutar 'npm install' "
            "en el servidor de CI/CD para sincronizar node_modules."
        )

    if any(f.kind == 'deleted' for f in changes):
        recs.append(
            "Se eliminaron archivos. Ejecutar build completo y verificar "
            "que no queden imports huerfanos."
        )

    if any(f.ext in ('.html', '.scss', '.css', '.component.html', '.component.scss')
           for f in changes):
        recs.append(
            "Cambios en interfaz de usuario. Realizar QA visual en "
            "distintas resoluciones (movil y escritorio)."
        )

    all_code = "\n".join(l for f in changes for l in f.added + f.removed)
    if re.search(r'\bexport\s+(interface|type)\b', all_code):
        recs.append(
            "Se modificaron interfaces o tipos exportados. Verificar que "
            "todos los consumidores del tipo esten actualizados y que el "
            "build compile sin errores de tipo."
        )

    if any(f.ext == '.spec.ts' for f in changes):
        recs.append(
            "Se modificaron tests unitarios. Ejecutar suite completa de "
            "pruebas antes del merge."
        )

    if re.search(r'migration|ALTER TABLE|DROP', all_code, re.I):
        recs.append(
            "ALERTA: Detectadas posibles migraciones de base de datos. "
            "Revisar con DBA antes del despliegue."
        )

    if re.search(r'@NgModule|providers\s*:|declarations\s*:', all_code):
        recs.append(
            "Se modifico un modulo Angular. Verificar que las declaraciones "
            "y providers sean correctas y que no haya duplicados."
        )

    if any(calculate_deploy_impact(f) in ("Alta", "Critica") for f in changes):
        recs.append(
            "Existen cambios de impacto ALTO o CRITICO. Se recomienda "
            "revision por pares (code review) antes del merge a desarrollo."
        )

    if not recs:
        recs.append(
            "No se detectaron riesgos criticos. Revisar cobertura de pruebas "
            "unitarias para las nuevas estructuras anadidas."
        )

    return list(dict.fromkeys(recs))


# =============================================================================
# HELPERS DOCX
# =============================================================================
def _set_bg(cell, hex_color: str):
    tcPr = cell._tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def _set_borders(cell, color: str = "CCCCCC"):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        borders.append(el)
    tcPr.append(borders)

def _set_width(cell, cm: float):
    tcPr = cell._tc.get_or_add_tcPr()
    tcW  = OxmlElement("w:tcW")
    tcW.set(qn("w:w"), str(int(cm * 567)))
    tcW.set(qn("w:type"), "dxa")
    tcPr.append(tcW)

def _run(para, text: str, bold=False, italic=False,
         color: RGBColor = None, size: float = 10, font: str = "Arial"):
    r = para.add_run(text)
    r.bold        = bold
    r.italic      = italic
    r.font.size   = Pt(size)
    r.font.name   = font
    if color:
        r.font.color.rgb = color
    return r

def _header_row(table, cols: List[Tuple[str, float]]):
    row = table.rows[0]
    for i, (text, w) in enumerate(cols):
        cell = row.cells[i]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        _set_bg(cell, C_HDR_BG)
        _set_borders(cell, C_ACCENT)
        _set_width(cell, w)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(p, text, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF), size=11)

def _data_row(table, cells: List[Tuple[str, float, RGBColor, bool, str]]):
    row = table.add_row()
    for i, (text, w, tc, bold, bg) in enumerate(cells):
        cell = row.cells[i]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        _set_bg(cell, bg)
        _set_borders(cell)
        _set_width(cell, w)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for j, line in enumerate(text.split("\n")):
            if j > 0:
                p = cell.add_paragraph()
            _run(p, line, bold=bold, color=tc, size=10)

def _bullet(doc, text: str, symbol: str, color: RGBColor, indent: float = 1.0):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent       = Cm(indent)
    p.paragraph_format.first_line_indent = Cm(-0.5)
    p.paragraph_format.space_after       = Pt(2)
    _run(p, f"{symbol}  ", bold=True, color=color, size=9.5)
    _run(p, text, color=color, size=9.5)

def _section_title(doc, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(5)
    _run(p, text, bold=True, color=C_TITLE, size=12)
    pPr    = p._p.get_or_add_pPr()
    pBdr   = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), C_ACCENT)
    pBdr.append(bottom)
    pPr.append(pBdr)

def _divider(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "4")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), C_BORDER)
    pBdr.append(bot)
    pPr.append(pBdr)

def _h1(doc, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(7)
    _run(p, text, bold=True, color=C_TITLE, size=14)


# =============================================================================
# GENERADOR DOCX
# =============================================================================
class ReportGenerator:
    def __init__(self, changes: List[FileChange], branch_from: str, branch_to: str, output: str):
        self.changes         = changes
        self.branch_from     = branch_from
        self.branch_to       = branch_to
        self.output          = output
        self.doc             = Document()
        self._eslint_reports: Dict[str, ESLintFileReport] = {}
        self._upgrade_compatibility()
        self._page_setup()

    def _upgrade_compatibility(self):
        settings = self.doc.settings.element
        compat   = OxmlElement('w:compat')
        cs       = OxmlElement('w:compatSetting')
        cs.set(qn('w:name'), 'compatibilityMode')
        cs.set(qn('w:uri'),  'http://schemas.microsoft.com/office/word')
        cs.set(qn('w:val'),  '15')
        compat.append(cs)
        settings.append(compat)

    def _page_setup(self):
        sec = self.doc.sections[0]
        sec.page_width    = Inches(8.5)
        sec.page_height   = Inches(11)
        sec.left_margin   = Cm(2.0)
        sec.right_margin  = Cm(2.0)
        sec.top_margin    = Cm(2.0)
        sec.bottom_margin = Cm(2.0)
        style = self.doc.styles["Normal"]
        style.font.name = "Arial"
        style.font.size = Pt(10)

    def _title(self):
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(3)
        _run(p, "INFORME DE CAMBIOS DE CODIGO", bold=True, color=C_TITLE, size=20)

        p2 = self.doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_after = Pt(2)
        _run(p2, f"Rama: {self.branch_from}  hacia  {self.branch_to}", color=C_SUBTITLE, size=11)

        p3 = self.doc.add_paragraph()
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p3.paragraph_format.space_after = Pt(12)
        _run(p3, f"Fecha: {fecha_espanol()}", color=C_MUTED, size=10)

        p4 = self.doc.add_paragraph()
        p4.paragraph_format.space_after = Pt(12)
        pPr  = p4._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bot  = OxmlElement("w:bottom")
        bot.set(qn("w:val"),   "single")
        bot.set(qn("w:sz"),    "12")
        bot.set(qn("w:space"), "1")
        bot.set(qn("w:color"), C_ACCENT)
        pBdr.append(bot)
        pPr.append(pBdr)

    def _summary(self):
        _h1(self.doc, "1. Tabla Resumen de Cambios")
        total_add = sum(len(f.added) for f in self.changes)
        total_del = sum(len(f.removed) for f in self.changes)

        p = self.doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        _run(
            p,
            f"Archivos afectados: {len(self.changes)}   "
            f"Lineas anadidas: +{total_add}   "
            f"Lineas eliminadas: -{total_del}",
            color=C_MUTED, size=9, italic=True
        )

        COLS = [
            ("Archivo",           3.8),
            ("Categoria",         2.6),
            ("Tipo de Cambio",    2.2),
            ("Impacto Tecnico",   4.0),
            ("Impacto en Deploy", 1.6),
        ]
        tbl = self.doc.add_table(rows=1, cols=len(COLS))
        tbl.style = "Table Grid"
        tbl.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _header_row(tbl, COLS)

        for fc in self.changes:
            tc_kind, bg_kind = fc.kind_colors
            category         = get_category(fc)
            deploy_level     = calculate_deploy_impact(fc)
            tc_dep, bg_dep   = impact_colors(deploy_level)
            tech_impact      = analyze_technical_impact(fc)

            _data_row(tbl, [
                (fc.filename,   3.8, C_MOD_TEXT, True,  "EFF6FF"),
                (category,      2.6, C_BODY,     False, C_ROW_ALT),
                (fc.kind_label, 2.2, tc_kind,    True,  bg_kind),
                (tech_impact,   4.0, C_BODY,     False, C_WHITE),
                (deploy_level,  1.6, tc_dep,     True,  bg_dep),
            ])

        self.doc.add_paragraph()

    def _render_lockfile_detail(self, fc: FileChange):
        """Renderizado compacto para archivos de dependencias (lock files)."""
        summary = summarize_lockfile(fc)
        p_info  = self.doc.add_paragraph()
        _run(
            p_info,
            f"Actualizacion de dependencias del proyecto. "
            f"Cambios: +{len(fc.added)} / -{len(fc.removed)} entradas. "
            f"Archivo gestionado automaticamente por el package manager.",
            color=C_MUTED, size=9, italic=True
        )
        if summary["added"]:
            p_a = self.doc.add_paragraph()
            _run(p_a, "Paquetes con entradas nuevas:", bold=True, color=C_ADD_TEXT, size=9)
            for pkg in summary["added"]:
                _bullet(self.doc, pkg, "+", C_ADD_TEXT)
        if summary["removed"]:
            p_r = self.doc.add_paragraph()
            _run(p_r, "Paquetes con entradas eliminadas:", bold=True, color=C_DEL_TEXT, size=9)
            for pkg in summary["removed"]:
                _bullet(self.doc, pkg, "-", C_DEL_TEXT)
        if not summary["added"] and not summary["removed"]:
            _bullet(self.doc, "Actualizacion de metadatos internos (integridad, resoluciones).", "-", C_MUTED)

    def _render_structural_summary(self, fc: FileChange):
        """Renderizado de resumen estructural para archivos extensos o con interfaces."""
        p_lbl = self.doc.add_paragraph()
        _run(p_lbl, "Resumen Estructural del archivo:", bold=True, color=C_TITLE, size=9)

        p_exec = self.doc.add_paragraph()
        _run(p_exec, "Sintesis funcional del cambio:", bold=True, color=C_SUBTITLE, size=9)
        for item in fc.build_functional_summary():
            _bullet(self.doc, item, "-", C_BODY, indent=1.5)

        struct = fc.extract_structure()

        if struct["decorators"]:
            p_d = self.doc.add_paragraph()
            _run(p_d, "Decoradores detectados:", bold=True, color=C_SUBTITLE, size=9)
            for d in struct["decorators"]:
                _bullet(self.doc, d, "-", C_REF_TEXT, indent=1.5)

        if struct["imports"]:
            p_i = self.doc.add_paragraph()
            _run(p_i, "Dependencias / Imports:", bold=True, color=C_SUBTITLE, size=9)
            for imp in struct["imports"]:
                _bullet(self.doc, imp, "-", C_BODY, indent=1.5)

        if struct["entities"]:
            p_e = self.doc.add_paragraph()
            _run(p_e, "Estructuras Logicas (Clases / Funciones / Interfaces / Tipos):",
                 bold=True, color=C_SUBTITLE, size=9)
            for ent in struct["entities"]:
                _bullet(self.doc, ent, "-", C_BODY, indent=1.5)

        if struct["routes"]:
            p_r = self.doc.add_paragraph()
            _run(p_r, "Rutas registradas:", bold=True, color=C_SUBTITLE, size=9)
            for rt in struct["routes"]:
                _bullet(self.doc, rt, "-", C_MOD_TEXT, indent=1.5)

        if struct["ui_components"]:
            p_ui = self.doc.add_paragraph()
            _run(p_ui, "Componentes de interfaz detectados:", bold=True, color=C_SUBTITLE, size=9)
            for comp in struct["ui_components"][:12]:
                _bullet(self.doc, comp, "-", C_MOD_TEXT, indent=1.5)

        if struct["forms"]:
            p_f = self.doc.add_paragraph()
            _run(p_f, "Controles de formulario/accion:", bold=True, color=C_SUBTITLE, size=9)
            for ctrl in struct["forms"][:12]:
                _bullet(self.doc, ctrl, "-", C_BODY, indent=1.5)

        if struct["angular_bindings"]:
            p_b = self.doc.add_paragraph()
            _run(p_b, "Bindings y directivas Angular:", bold=True, color=C_SUBTITLE, size=9)
            for bind in struct["angular_bindings"][:12]:
                _bullet(self.doc, bind, "-", C_BODY, indent=1.5)

        if struct["events"]:
            p_ev = self.doc.add_paragraph()
            _run(p_ev, "Eventos de usuario detectados:", bold=True, color=C_SUBTITLE, size=9)
            for ev in struct["events"][:12]:
                _bullet(self.doc, ev, "-", C_REF_TEXT, indent=1.5)

        if struct["css_selectors"]:
            p_css = self.doc.add_paragraph()
            _run(p_css, "Selectores/Reglas de estilo detectados:", bold=True, color=C_SUBTITLE, size=9)
            for sel in struct["css_selectors"][:12]:
                _bullet(self.doc, sel, "-", C_BODY, indent=1.5)

        if not any(struct.values()):
            _bullet(
                self.doc,
                "Contenido estructurado sin entidades detectables (datos, configuracion, markup).",
                "-", C_BODY
            )

    def _render_line_detail(self, fc: FileChange):
        """Renderizado linea a linea para cambios acotados con clasificacion de razon."""
        if fc.added:
            p_add = self.doc.add_paragraph()
            _run(p_add, "Lineas anadidas:", bold=True, color=C_ADD_TEXT, size=9)
            for line_no, line in (fc.added_with_line or [(None, l) for l in fc.added]):
                if line_no is not None:
                    _bullet(self.doc, f"L{line_no}: {line}", "+", C_ADD_TEXT)
                else:
                    _bullet(self.doc, line, "+", C_ADD_TEXT)

        if fc.removed:
            p_rem = self.doc.add_paragraph()
            p_rem.paragraph_format.space_before = Pt(4)
            _run(p_rem, "Lineas eliminadas:", bold=True, color=C_DEL_TEXT, size=9)

            TAG_COLORS = {
                "Linter":        C_ADD_TEXT,
                "Limpieza":      C_MUTED,
                "Debug":         C_MUTED,
                "Doc":           C_MUTED,
                "Deuda tecnica": C_REF_TEXT,
                "Config":        C_MOD_TEXT,
                "Refactor":      C_REF_TEXT,
                "Test":          C_SUBTITLE,
            }

            removed_items = fc.removed_with_line or [(None, l) for l in fc.removed]
            for line_no, line in removed_items:
                reason = fc.classify_removed_line(line)
                line_display = f"L{line_no}: {line}" if line_no is not None else line
                if reason:
                    tag_match = re.match(r'\[([^\]]+)\]\s*(.*)', reason)
                    if tag_match:
                        tag_key  = tag_match.group(1)
                        tag_desc = tag_match.group(2)
                        tc_tag   = TAG_COLORS.get(tag_key, C_SUBTITLE)
                        p_line   = self.doc.add_paragraph()
                        p_line.paragraph_format.left_indent       = Cm(1.0)
                        p_line.paragraph_format.first_line_indent = Cm(-0.5)
                        p_line.paragraph_format.space_after       = Pt(2)
                        _run(p_line, "-  ", bold=True, color=C_DEL_TEXT, size=9)
                        _run(p_line, line_display, color=C_DEL_TEXT, size=9)
                        _run(p_line, f"  [{tag_key}: {tag_desc}]",
                             bold=True, color=tc_tag, size=8, italic=True)
                    else:
                        _bullet(self.doc, f"{line_display}  ({reason})", "-", C_DEL_TEXT)
                else:
                    _bullet(self.doc, line_display, "-", C_DEL_TEXT)

    # Extensiones que NO deben mostrar detalle de líneas incluso en resumen estructural
    _SKIP_LINE_DETAIL_EXTS: frozenset = frozenset([
        '.bak', '.sln', '.vbproj', '.csproj', '.vcxproj', '.fsproj',
        '.resx', '.xml', '.config', '.manifest',
    ])

    def _render_line_detail_limited(self, fc: "FileChange", max_added: int = 80, max_removed: int = 80):
        """
        Renderizado de líneas con límite para archivos grandes que además tienen resumen
        estructural. Garantiza que los nuevos métodos y cambios de lógica sean siempre
        visibles en el informe aunque el archivo supere el umbral de líneas.
        """
        added_items   = fc.added_with_line   or [(None, l) for l in fc.added]
        removed_items = fc.removed_with_line or [(None, l) for l in fc.removed]

        # ── Líneas añadidas ───────────────────────────────────────────────────
        if added_items:
            n_total   = len(added_items)
            truncated = n_total > max_added
            label     = (
                f"Lineas anadidas ({n_total} total"
                f" — se muestran las primeras {max_added}):"
                if truncated else "Lineas anadidas:"
            )
            p_add = self.doc.add_paragraph()
            p_add.paragraph_format.space_before = Pt(6)
            _run(p_add, label, bold=True, color=C_ADD_TEXT, size=9)

            for line_no, line in added_items[:max_added]:
                if line_no is not None:
                    _bullet(self.doc, f"L{line_no}: {line}", "+", C_ADD_TEXT)
                else:
                    _bullet(self.doc, line, "+", C_ADD_TEXT)

            if truncated:
                _bullet(
                    self.doc,
                    f"... y {n_total - max_added} linea(s) adicional(es) no mostradas.",
                    "…", C_MUTED,
                )

        # ── Líneas eliminadas ─────────────────────────────────────────────────
        if removed_items:
            n_total   = len(removed_items)
            truncated = n_total > max_removed
            label     = (
                f"Lineas eliminadas ({n_total} total"
                f" — se muestran las primeras {max_removed}):"
                if truncated else "Lineas eliminadas:"
            )
            p_rem = self.doc.add_paragraph()
            p_rem.paragraph_format.space_before = Pt(4)
            _run(p_rem, label, bold=True, color=C_DEL_TEXT, size=9)

            TAG_COLORS = {
                "Linter":        C_ADD_TEXT,
                "Limpieza":      C_MUTED,
                "Debug":         C_MUTED,
                "Doc":           C_MUTED,
                "Deuda tecnica": C_REF_TEXT,
                "Config":        C_MOD_TEXT,
                "Refactor":      C_REF_TEXT,
                "Test":          C_SUBTITLE,
            }

            for line_no, line in removed_items[:max_removed]:
                reason       = fc.classify_removed_line(line)
                line_display = f"L{line_no}: {line}" if line_no is not None else line
                if reason:
                    tag_match = re.match(r'\[([^\]]+)\]\s*(.*)', reason)
                    if tag_match:
                        tag_key  = tag_match.group(1)
                        tag_desc = tag_match.group(2)
                        tc_tag   = TAG_COLORS.get(tag_key, C_SUBTITLE)
                        p_line   = self.doc.add_paragraph()
                        p_line.paragraph_format.left_indent       = Cm(1.0)
                        p_line.paragraph_format.first_line_indent = Cm(-0.5)
                        p_line.paragraph_format.space_after       = Pt(2)
                        _run(p_line, "-  ", bold=True, color=C_DEL_TEXT, size=9)
                        _run(p_line, line_display, color=C_DEL_TEXT, size=9)
                        _run(p_line, f"  [{tag_key}: {tag_desc}]",
                             bold=True, color=tc_tag, size=8, italic=True)
                    else:
                        _bullet(self.doc, f"{line_display}  ({reason})", "-", C_DEL_TEXT)
                else:
                    _bullet(self.doc, line_display, "-", C_DEL_TEXT)

            if truncated:
                _bullet(
                    self.doc,
                    f"... y {n_total - max_removed} linea(s) adicional(es) no mostradas.",
                    "…", C_MUTED,
                )

    def _render_eslint_report(self, eslint_report: "ESLintFileReport", fc: "FileChange"):
        """
        Renderiza las tablas ESLint (correcciones, violaciones y cambios funcionales)
        de un archivo Angular (.ts o .html) dentro del informe docx.
        Incluye la tabla de riesgos de incompatibilidad de tipos Front-Back.
        """
        ext = fc.ext.lower()
        is_html = ext in ('.html', '.component.html')
        is_ts   = ext in (
            '.ts', '.component.ts', '.service.ts', '.pipe.ts',
            '.directive.ts', '.guard.ts', '.interceptor.ts',
            '.module.ts', '.resolver.ts',
        )
        if not (is_html or is_ts):
            return
        if not eslint_report.corrections and not eslint_report.violations:
            if not eslint_report.removed_had_violations:
                p_ok = self.doc.add_paragraph()
                p_ok.paragraph_format.space_before = Pt(4)
                _run(p_ok, "Análisis ESLint: ", bold=True, color=C_TITLE, size=9)
                _run(p_ok,
                     "No se detectaron correcciones ni violaciones ESLint en las líneas añadidas.",
                     color=C_MUTED, size=9, italic=True)
            # Aún puede haber riesgos front-back o estructurales → no retornar sin renderizar
            if not eslint_report.front_back_risks and not eslint_report.structural_findings:
                return

        # ── Correcciones ESLint ───────────────────────────────────────────────
        if eslint_report.corrections:
            p_corr = self.doc.add_paragraph()
            p_corr.paragraph_format.space_before = Pt(6)
            _run(p_corr, f"✔ Correcciones ESLint aplicadas ({len(eslint_report.corrections)})",
                 bold=True, color=C_ADD_TEXT, size=10)

            CORR_COLS = [
                ("#",             0.5),
                ("Antes (−)",     4.8),
                ("Después (+)",   4.8),
                ("Regla corregida", 3.4),
                ("Severidad",     1.3),
            ]
            tbl_c = self.doc.add_table(rows=1, cols=len(CORR_COLS))
            tbl_c.style = "Table Grid"
            _header_row(tbl_c, CORR_COLS)

            for idx, f in enumerate(eslint_report.corrections, 1):
                before = (f.line_removed[:90] + "…") if len(f.line_removed) > 90 else f.line_removed
                after  = (f.line_added[:90]   + "…") if len(f.line_added)   > 90 else f.line_added
                _data_row(tbl_c, [
                    (str(idx),    0.5,  C_MUTED,    False, C_WHITE),
                    (before,      4.8,  C_DEL_TEXT, False, C_DEL_BG),
                    (after,       4.8,  C_ADD_TEXT, False, C_ADD_BG),
                    (f.rule,      3.4,  C_MOD_TEXT, False, C_ROW_ALT),
                    (f.severity,  1.3,  C_ADD_TEXT, True,  C_ADD_BG),
                ])
            self.doc.add_paragraph()

        # ── Violaciones ESLint ────────────────────────────────────────────────
        if eslint_report.violations:
            p_viol = self.doc.add_paragraph()
            p_viol.paragraph_format.space_before = Pt(6)
            _run(p_viol, f"✘ Violaciones ESLint introducidas ({len(eslint_report.violations)})",
                 bold=True, color=C_DEL_TEXT, size=10)

            VIOL_COLS = [
                ("#",            0.5),
                ("Línea añadida (+)", 6.0),
                ("Regla violada",    5.5),
                ("Severidad",        1.3),
            ]
            tbl_v = self.doc.add_table(rows=1, cols=len(VIOL_COLS))
            tbl_v.style = "Table Grid"
            _header_row(tbl_v, VIOL_COLS)

            for idx, f in enumerate(eslint_report.violations, 1):
                line_disp = (f.line_added[:110] + "…") if len(f.line_added) > 110 else f.line_added
                _data_row(tbl_v, [
                    (str(idx),    0.5,  C_MUTED,    False, C_WHITE),
                    (line_disp,   6.0,  C_DEL_TEXT, False, C_DEL_BG),
                    (f.rule,      5.5,  C_MOD_TEXT, False, C_ROW_ALT),
                    (f.severity,  1.3,  C_DEL_TEXT, True,  C_DEL_BG),
                ])
            self.doc.add_paragraph()

        # ── Riesgos de Incompatibilidad de Tipos Front-Back ───────────────────
        if eslint_report.front_back_risks:
            # Colores por nivel de riesgo
            _risk_color = {
                'ALTO':  (C_IMPACT_ALTA,   BG_IMPACT_ALTA),
                'MEDIO': (C_IMPACT_MEDIA,  BG_IMPACT_MEDIA),
                'BAJO':  (C_IMPACT_BAJA,   BG_IMPACT_BAJA),
            }
            p_fb = self.doc.add_paragraph()
            p_fb.paragraph_format.space_before = Pt(8)
            _run(p_fb,
                 f"⚠ Riesgos de incompatibilidad de tipos Front ↔ Back ({len(eslint_report.front_back_risks)})",
                 bold=True, color=C_IMPACT_ALTA, size=10)

            p_desc = self.doc.add_paragraph()
            p_desc.paragraph_format.space_before = Pt(2)
            p_desc.paragraph_format.space_after  = Pt(4)
            _run(p_desc,
                 "Estos riesgos surgen porque == (igualdad laxa) realiza coerción de tipos: "
                 "1 == true → TRUE, pero 1 === true → FALSE. "
                 "Al corregir == → === (ESLint), la lógica cambia si el back-end "
                 "devuelve un tipo distinto al esperado por el front.",
                 color=C_SUBTITLE, size=8, italic=True)

            FB_COLS = [
                ("#",           0.4),
                ("Expresión",   3.0),
                ("Escenario",   2.5),
                ("Riesgo",      0.8),
                ("Explicación", 4.5),
                ("Recomendación", 3.5),
            ]
            tbl_fb = self.doc.add_table(rows=1, cols=len(FB_COLS))
            tbl_fb.style = "Table Grid"
            _header_row(tbl_fb, FB_COLS)

            for idx, risk in enumerate(eslint_report.front_back_risks, 1):
                tc, bg = _risk_color.get(risk.risk_level, (C_BODY, C_WHITE))
                _data_row(tbl_fb, [
                    (str(idx),            0.4,  C_MUTED,    False, C_WHITE),
                    (risk.expression,     3.0,  C_MOD_TEXT, True,  C_MOD_BG),
                    (risk.scenario,       2.5,  tc,         True,  bg),
                    (risk.risk_level,     0.8,  tc,         True,  bg),
                    (risk.explanation,    4.5,  C_BODY,     False, C_WHITE),
                    (risk.recommendation, 3.5,  C_SUBTITLE, False, C_ROW_ALT),
                ])
            self.doc.add_paragraph()

        # ── Cambios Estructurales JS/TS (optional chaining, ??, !., typeof) ──
        if eslint_report.structural_findings:
            _sev_color = {
                'ATENCION':    (C_IMPACT_MEDIA,  BG_IMPACT_MEDIA),
                'INFORMATIVO': (C_IMPACT_BAJA,   BG_IMPACT_BAJA),
            }
            p_st = self.doc.add_paragraph()
            p_st.paragraph_format.space_before = Pt(8)
            _run(p_st,
                 f"🔍 Cambios estructurales JS/TS detectados ({len(eslint_report.structural_findings)})",
                 bold=True, color=C_MOD_TEXT, size=10)

            p_st_desc = self.doc.add_paragraph()
            p_st_desc.paragraph_format.space_before = Pt(2)
            p_st_desc.paragraph_format.space_after  = Pt(4)
            _run(p_st_desc,
                 "Cambios de sintaxis que alteran la semántica de evaluación de nulidad o coerción: "
                 "optional chaining (?.), nullish coalescing (??), non-null assertion (!.) o "
                 "eliminación de guardas typeof. Se documenta la diferencia de comportamiento "
                 "para validar que el cambio no introduce regresiones.",
                 color=C_SUBTITLE, size=8, italic=True)

            ST_COLS = [
                ("#",               0.4),
                ("Patrón",          3.2),
                ("Antes (−)",       3.8),
                ("Después (+)",     3.8),
                ("Severidad",       0.8),
                ("Explicación",     4.5),
                ("Recomendación",   3.5),
            ]
            tbl_st = self.doc.add_table(rows=1, cols=len(ST_COLS))
            tbl_st.style = "Table Grid"
            _header_row(tbl_st, ST_COLS)

            for idx, sf in enumerate(eslint_report.structural_findings, 1):
                tc, bg = _sev_color.get(sf.severity, (C_BODY, C_WHITE))
                _data_row(tbl_st, [
                    (str(idx),        0.4,  C_MUTED,    False, C_WHITE),
                    (sf.pattern,      3.2,  C_MOD_TEXT, True,  C_MOD_BG),
                    (sf.before,       3.8,  C_DEL_TEXT, False, C_DEL_BG),
                    (sf.after,        3.8,  C_ADD_TEXT, False, C_ADD_BG),
                    (sf.severity,     0.8,  tc,         True,  bg),
                    (sf.explanation,  4.5,  C_BODY,     False, C_WHITE),
                    (sf.recommendation, 3.5, C_SUBTITLE, False, C_ROW_ALT),
                ])
            self.doc.add_paragraph()

    def _detail(self):
        semantic_engine = SemanticInsightEngine()
        relation_engine = ChangeRelationAnalyzer(self.changes)
        eslint_analyzer  = ESLintAngularAnalyzer()
        _h1(self.doc, "2. Detalle de Cambios por Archivo")

        for i, fc in enumerate(self.changes):
            _section_title(self.doc, fc.filename)

            p_path = self.doc.add_paragraph()
            p_path.paragraph_format.space_after = Pt(3)
            _run(p_path, fc.filepath, color=C_MUTED, size=8, italic=True)

            if fc.contexts:
                p_ctx = self.doc.add_paragraph()
                p_ctx.paragraph_format.space_after = Pt(3)
                _run(p_ctx, "Bloques / Funciones afectadas: ", bold=True, color=C_SUBTITLE, size=9)
                _run(p_ctx, ", ".join(fc.contexts), color=C_BODY, size=9, italic=True)

            deploy_level = calculate_deploy_impact(fc)
             # -----------------------------------------
            # ANALISIS SEMANTICO CONTEXTUAL
            # -----------------------------------------
            semantic_insights = semantic_engine.analyze(fc)
            if semantic_insights:
                p_sem = self.doc.add_paragraph()
                _run(p_sem, "Analisis Semantico Detectado:",
                     bold=True, color=C_TITLE, size=9)

                for insight in semantic_insights:
                    _bullet(self.doc, insight, ">", C_MOD_TEXT)

            relation_insights = relation_engine.analyze(fc)
            if relation_insights:
                p_rel = self.doc.add_paragraph()
                _run(p_rel, "Interrelaciones Detectadas:",
                     bold=True, color=C_TITLE, size=9)
                for insight in relation_insights:
                    _bullet(self.doc, insight, " > ", C_REF_TEXT)

            tc_dep, _    = impact_colors(deploy_level)
            p_stats = self.doc.add_paragraph()
            p_stats.paragraph_format.space_after = Pt(4)
            _run(p_stats, f"+{len(fc.added)} anadidas   -{len(fc.removed)} eliminadas   ",
                 color=C_MUTED, size=8)
            _run(p_stats, f"Impacto en deploy: {deploy_level}", bold=True, color=tc_dep, size=8)

            # ── Análisis ESLint (antes del renderizado de líneas) ──────────────
            eslint_report = eslint_analyzer.analyze_file(fc)
            self._eslint_reports[fc.filepath] = eslint_report

            # Elegir estrategia de renderizado
            if fc.is_binary:
                _bullet(self.doc, "Archivo binario. No se puede mostrar contenido de texto.",
                        "-", C_MUTED)
            elif fc.is_lockfile:
                self._render_lockfile_detail(fc)
            elif fc.needs_structural_summary:
                self._render_structural_summary(fc)
                # FIX: Para archivos MODIFICADOS (no nuevos, no extensiones sólo-estructura),
                # mostrar también las líneas cambiadas para que los nuevos métodos y cambios
                # de lógica queden explícitos en el informe.
                if fc.kind not in ('added',) and fc.ext not in self._SKIP_LINE_DETAIL_EXTS:
                    self._render_line_detail_limited(fc)
            else:
                self._render_line_detail(fc)

            # ── Reporte ESLint (correcciones / violaciones) ─────────────────────
            self._render_eslint_report(eslint_report, fc)

            if not fc.added and not fc.removed and not fc.is_binary:
                p_empty = self.doc.add_paragraph()
                _run(
                    p_empty,
                    "Sin modificaciones en el codigo fuente. Posible cambio de "
                    "permisos, archivo vacio o renombramiento.",
                    color=C_MUTED, size=9, italic=True
                )

            if i < len(self.changes) - 1:
                _divider(self.doc)

        self.doc.add_paragraph()

    def _impact_legend(self):
        _h1(self.doc, "3. Leyenda de Niveles de Impacto en Deploy")

        legend = [
            ("Nula",
             "Sin riesgo. Documentacion, solo linter/formatter o configuracion menor."),
            ("Leve",
             "Archivo de dependencias, nuevo archivo aislado o cambio estetico."),
            ("Baja",
             "Nuevo componente/servicio sin contratos externos. Logica interna acotada."),
            ("Media",
             "Cambio en logica de negocio, llamadas HTTP, rutas o estado compartido."),
            ("Alta",
             "Interfaces/tipos exportados, modulos Angular, guards o archivos de entorno."),
            ("Critica",
             "Autenticacion, JWT, passwords, migraciones de BD o contratos de API externa."),
        ]

        tbl = self.doc.add_table(rows=1, cols=2)
        tbl.style = "Table Grid"
        tbl.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _header_row(tbl, [("Nivel", 2.5), ("Descripcion del criterio", 11.7)])

        for level, desc in legend:
            tc_l, bg_l = impact_colors(level)
            _data_row(tbl, [
                (level, 2.5,  tc_l,   True,  bg_l),
                (desc,  11.7, C_BODY, False, C_WHITE),
            ])

        self.doc.add_paragraph()

    def _recommendations(self):
        _h1(self.doc, "4. Recomendaciones antes del Merge")
        recs = analyze_recommendations(self.changes)
        p = self.doc.add_paragraph()
        p.paragraph_format.space_after = Pt(7)
        _run(
            p,
            "Recomendaciones generadas a partir del analisis del codigo modificado:",
            color=C_BODY, size=10
        )
        for rec in recs:
            _bullet(self.doc, rec, " > ", C_MOD_TEXT)
        self.doc.add_paragraph()

    def _eslint_global_summary(self):
        """Sección 5: Resumen global del análisis ESLint + riesgos Front-Back."""
        if not self._eslint_reports:
            return

        all_corrections:  List[ESLintFinding]      = []
        all_violations:   List[ESLintFinding]      = []
        all_fb_risks:     List[FrontBackRiskFinding] = []
        for report in self._eslint_reports.values():
            all_corrections.extend(report.corrections)
            all_violations.extend(report.violations)
            all_fb_risks.extend(report.front_back_risks)

        # Contar frecuencia por regla
        from collections import Counter
        corr_counter = Counter(f.rule for f in all_corrections)
        viol_counter = Counter(f.rule for f in all_violations)
        fb_counter   = Counter(r.scenario.split(']')[-1].strip() for r in all_fb_risks)

        _h1(self.doc, "5. Resumen Global del Análisis ESLint")

        # Estadísticas generales
        ts_files   = sum(1 for fc in self.changes
                         if fc.ext.lower() in ('.ts', '.component.ts', '.service.ts',
                                               '.pipe.ts', '.directive.ts', '.guard.ts',
                                               '.interceptor.ts', '.module.ts', '.resolver.ts'))
        html_files = sum(1 for fc in self.changes
                         if fc.ext.lower() in ('.html', '.component.html'))

        p_stat = self.doc.add_paragraph()
        p_stat.paragraph_format.space_after = Pt(6)
        _run(p_stat,
             f"Archivos TypeScript analizados: {ts_files}   "
             f"Archivos HTML analizados: {html_files}   "
             f"Correcciones ESLint: {len(all_corrections)}   "
             f"Violaciones ESLint: {len(all_violations)}",
             color=C_MUTED, size=9, italic=True)

        # ── Tabla de correcciones por regla ───────────────────────────────────
        if corr_counter:
            p_c = self.doc.add_paragraph()
            _run(p_c, "Reglas más frecuentemente corregidas:", bold=True,
                 color=C_ADD_TEXT, size=10)
            COLS_C = [("#", 0.5), ("Regla ESLint", 8.5), ("Veces corregida", 2.2), ("Severidad", 1.5)]
            tbl_c = self.doc.add_table(rows=1, cols=len(COLS_C))
            tbl_c.style = "Table Grid"
            tbl_c.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _header_row(tbl_c, COLS_C)
            for rank, (rule, count) in enumerate(corr_counter.most_common(10), 1):
                _data_row(tbl_c, [
                    (str(rank), 0.5,  C_MUTED,    False, C_WHITE),
                    (rule,      8.5,  C_MOD_TEXT, False, C_ROW_ALT),
                    (str(count),2.2,  C_ADD_TEXT, True,  C_ADD_BG),
                    ("ERROR",   1.5,  C_ADD_TEXT, True,  C_ADD_BG),
                ])
            self.doc.add_paragraph()

        # ── Tabla de violaciones por regla ────────────────────────────────────
        if viol_counter:
            p_v = self.doc.add_paragraph()
            _run(p_v, "Reglas con violaciones introducidas:", bold=True,
                 color=C_DEL_TEXT, size=10)
            COLS_V = [("#", 0.5), ("Regla ESLint", 8.5), ("Ocurrencias", 2.2), ("Severidad", 1.5)]
            tbl_v = self.doc.add_table(rows=1, cols=len(COLS_V))
            tbl_v.style = "Table Grid"
            _header_row(tbl_v, COLS_V)
            for rank, (rule, count) in enumerate(viol_counter.most_common(10), 1):
                _data_row(tbl_v, [
                    (str(rank), 0.5,  C_MUTED,    False, C_WHITE),
                    (rule,      8.5,  C_MOD_TEXT, False, C_ROW_ALT),
                    (str(count),2.2,  C_DEL_TEXT, True,  C_DEL_BG),
                    ("ERROR",   1.5,  C_DEL_TEXT, True,  C_DEL_BG),
                ])
            self.doc.add_paragraph()

        # ── Tabla global de riesgos Front-Back ────────────────────────────────
        if all_fb_risks:
            p_fb_title = self.doc.add_paragraph()
            p_fb_title.paragraph_format.space_before = Pt(10)
            _run(p_fb_title,
                 f"Resumen de Riesgos de Tipos Front ↔ Back ({len(all_fb_risks)} detectados)",
                 bold=True, color=C_IMPACT_ALTA, size=11)

            p_fb_intro = self.doc.add_paragraph()
            p_fb_intro.paragraph_format.space_after = Pt(6)
            _run(p_fb_intro,
                 "El motor de análisis detectó comparaciones donde la corrección ESLint (== → ===) "
                 "puede alterar la lógica si el back-end retorna tipos distintos (1 en vez de true, "
                 "null en vez de undefined, string '1' en vez del número 1, etc.). "
                 "Revisar el contrato de cada campo con el back-end.",
                 color=C_SUBTITLE, size=8.5, italic=True)

            # Tabla consolidada por escenario
            if fb_counter:
                FB_SUM_COLS = [
                    ("#",             0.4),
                    ("Escenario de riesgo", 5.5),
                    ("Ocurrencias",   1.5),
                    ("Nivel",         1.0),
                    ("Qué verificar", 6.3),
                ]
                tbl_fb_sum = self.doc.add_table(rows=1, cols=len(FB_SUM_COLS))
                tbl_fb_sum.style = "Table Grid"
                tbl_fb_sum.alignment = WD_ALIGN_PARAGRAPH.CENTER
                _header_row(tbl_fb_sum, FB_SUM_COLS)

                # Mapa escenario → (nivel, recomendación)
                esc_meta: Dict[str, Tuple[str, str]] = {}
                for r in all_fb_risks:
                    key = r.scenario.split(']')[-1].strip()
                    if key not in esc_meta:
                        esc_meta[key] = (r.risk_level, r.recommendation[:180])

                _risk_color = {
                    'ALTO':  (C_IMPACT_ALTA,  BG_IMPACT_ALTA),
                    'MEDIO': (C_IMPACT_MEDIA, BG_IMPACT_MEDIA),
                    'BAJO':  (C_IMPACT_BAJA,  BG_IMPACT_BAJA),
                }

                for rank, (scenario, count) in enumerate(fb_counter.most_common(10), 1):
                    lvl, reco = esc_meta.get(scenario, ('MEDIO', ''))
                    tc, bg = _risk_color.get(lvl, (C_BODY, C_WHITE))
                    _data_row(tbl_fb_sum, [
                        (str(rank), 0.4,  C_MUTED,    False, C_WHITE),
                        (scenario,  5.5,  C_MOD_TEXT, False, C_ROW_ALT),
                        (str(count),1.5,  tc,         True,  bg),
                        (lvl,       1.0,  tc,         True,  bg),
                        (reco,      6.3,  C_SUBTITLE, False, C_WHITE),
                    ])
                self.doc.add_paragraph()

        if not all_corrections and not all_violations and not all_fb_risks:
            self.doc.add_paragraph()
            return

    def _footer(self):
        _divider(self.doc)
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(6)
        total_add = sum(len(f.added) for f in self.changes)
        total_del = sum(len(f.removed) for f in self.changes)
        _run(
            p,
            f"Archivos: {len(self.changes)}  |  "
            f"+{total_add} lineas  |  "
            f"-{total_del} lineas  |  "
            f"Archivo generado automaticamente - Ryu Gabo -",
            color=C_MUTED, size=9.5, italic=True
        )

    def generate(self) -> str:
        self._title()
        self._summary()
        self._detail()
        self._impact_legend()
        self._recommendations()
        self._eslint_global_summary()
        self._footer()
        self.doc.save(self.output)
        return self.output


# =============================================================================
# CLI Y EJECUCION
# =============================================================================
def find_input(arg: Optional[str]) -> Path:
    if arg:
        p = Path(arg)
        if not p.exists():
            print(f"[ERROR] No se encontro: {p}", file=sys.stderr)
            sys.exit(1)
        return p
    for candidate in (Path.cwd() / DEFAULT_INPUT, Path(__file__).parent / DEFAULT_INPUT):
        if candidate.exists():
            return candidate
    print(
        f"[ERROR] No se encontro '{DEFAULT_INPUT}'.\n"
        f"Genera el archivo con:\n"
        f"  git --no-pager diff --staged -U9999 > {DEFAULT_INPUT}",
        file=sys.stderr
    )
    sys.exit(1)


def main():
    ap = argparse.ArgumentParser(
        description="Convierte 'informe.txt' (git diff) en un informe .docx profesional."
    )
    ap.add_argument("--input",       "-i",  default=None,       help=f"Archivo diff (default: {DEFAULT_INPUT})")
    ap.add_argument("--output",      "-o",  default=None,       help=f"Archivo salida (default: {DEFAULT_OUTPUT})")
    ap.add_argument("--branch-from", "-bf", default="gabotest", help="Rama origen  (default: gabotest)")
    ap.add_argument("--branch-to",   "-bt", default="desarrollo",  help="Rama destino (default: desarrollo)")
    args = ap.parse_args()

    input_path = find_input(args.input)
    print(f"[INFO] Leyendo: {input_path}")

    text = clean(input_path.read_text(encoding="utf-8", errors="replace"))
    if not text.strip():
        print("[ERROR] El archivo esta vacio.", file=sys.stderr)
        sys.exit(1)

    changes = DiffParser().parse(text)
    if not changes:
        print("[AVISO] No se encontraron cambios legibles.", file=sys.stderr)
        sys.exit(1)

    output = args.output or str(input_path.parent / DEFAULT_OUTPUT)
    result = ReportGenerator(
        changes=changes,
        branch_from=args.branch_from,
        branch_to=args.branch_to,
        output=output
    ).generate()

    print(f"\n[OK] Informe generado: {result}")


if __name__ == "__main__":
    main()