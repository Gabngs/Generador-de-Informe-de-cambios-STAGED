#!/usr/bin/env python3
"""
git_diff_to_docx.py
────────────────────
Lee 'informe.txt' (en la misma carpeta) generado con:
    git --no-pager diff --staged > informe.txt

Analiza cualquier tipo de diff de forma genérica y produce
un informe profesional .docx con:
  • Tabla resumen por archivo
  • Detalle de líneas añadidas / eliminadas por archivo
  • Análisis automático de impacto y recomendaciones

Uso:
    python git_diff_to_docx.py
    python git_diff_to_docx.py --input otro.txt --output reporte.docx
    python git_diff_to_docx.py --branch-from feature/x --branch-to develop
"""

import sys
import re
import argparse
from datetime import datetime
from pathlib import Path
from dataclasses import dataclass, field
from typing import List, Tuple, Optional, Dict

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ─────────────────────────────────────────────────────────────────────────────
# PALETA DE COLORES
# ─────────────────────────────────────────────────────────────────────────────
C_TITLE     = RGBColor(0x1E, 0x3A, 0x5F)
C_SUBTITLE  = RGBColor(0x55, 0x55, 0x55)
C_BODY      = RGBColor(0x22, 0x22, 0x22)
C_MUTED     = RGBColor(0x88, 0x88, 0x88)
C_ADD_TEXT  = RGBColor(0x16, 0x65, 0x34)
C_ADD_BG    = "F0FDF4"
C_DEL_TEXT  = RGBColor(0x99, 0x1B, 0x1B)
C_DEL_BG    = "FEF2F2"
C_MOD_TEXT  = RGBColor(0x1E, 0x3A, 0x5F)
C_MOD_BG    = "EFF6FF"
C_REF_TEXT  = RGBColor(0x92, 0x40, 0x0E)
C_REF_BG    = "FFFBEB"
C_HDR_BG    = "1E3A5F"
C_ACCENT    = "2563EB"
C_BORDER    = "CCCCCC"
C_ROW_ALT   = "F8FAFC"
C_WHITE     = "FFFFFF"

DEFAULT_INPUT  = "informe.txt"
DEFAULT_OUTPUT = "informe_cambios.docx"

# Regex para limpiar ANSI
ANSI_RE = re.compile(r'\x1B(?:[@-Z\\-_]|\[[0-?]*[ -/]*[@-~])')


# ─────────────────────────────────────────────────────────────────────────────
# LIMPIEZA
# ─────────────────────────────────────────────────────────────────────────────
def clean(text: str) -> str:
    text = text.replace('\r\n', '\n').replace('\r', '\n')
    text = ANSI_RE.sub('', text)
    return text.replace('\x00', '')


# ─────────────────────────────────────────────────────────────────────────────
# MODELO: FileChange
# ─────────────────────────────────────────────────────────────────────────────
@dataclass
class FileChange:
    filepath: str
    added:    List[str] = field(default_factory=list)   # líneas +
    removed:  List[str] = field(default_factory=list)   # líneas -
    kind: str = "modified"   # added | deleted | renamed | modified

    # ── propiedades de presentación ──────────────────────────────────────────
    @property
    def filename(self) -> str:
        return Path(self.filepath).name

    @property
    def ext(self) -> str:
        name = self.filename.lower()
        # extensiones dobles conocidas
        for double in ('.component.ts', '.component.html', '.component.scss',
                       '.service.ts', '.spec.ts', '.module.ts', '.pipe.ts',
                       '.directive.ts', '.guard.ts', '.interceptor.ts'):
            if name.endswith(double):
                return double
        return Path(self.filepath).suffix.lower()

    @property
    def kind_label(self) -> str:
        if self.kind == 'added':
            return 'Adición'
        if self.kind == 'deleted':
            return 'Eliminación'
        if self.kind == 'renamed':
            return 'Renombrado'
        n_add = len(self.added)
        n_del = len(self.removed)
        if n_del == 0 and n_add > 0:
            return 'Adición'
        if n_add == 0 and n_del > 0:
            return 'Eliminación'
        if n_del > 0 and n_add == 0:
            return 'Eliminación'
        if n_del > n_add * 2:
            return 'Refactor'
        return 'Modificación'

    @property
    def kind_colors(self) -> Tuple[RGBColor, str]:
        lbl = self.kind_label
        if lbl == 'Adición':    return C_ADD_TEXT, C_ADD_BG
        if lbl == 'Eliminación': return C_DEL_TEXT, C_DEL_BG
        if lbl == 'Refactor':   return C_REF_TEXT, C_REF_BG
        return C_MOD_TEXT, C_MOD_BG   # Modificación / Renombrado


# ─────────────────────────────────────────────────────────────────────────────
# PARSER DE DIFF  (genérico, sin supuestos de contenido)
# ─────────────────────────────────────────────────────────────────────────────
class DiffParser:
    _FILE  = re.compile(r'^diff --git a/(.+?) b/(.+)$')
    _NEW   = re.compile(r'^new file mode')
    _DEL   = re.compile(r'^deleted file mode')
    _REN   = re.compile(r'^rename to (.+)$')
    _HUNK  = re.compile(r'^@@')
    _PLUS3 = re.compile(r'^\+\+\+')
    _MIN3  = re.compile(r'^---')

    def parse(self, text: str) -> List[FileChange]:
        text  = clean(text)
        files: List[FileChange] = []
        cur:   Optional[FileChange] = None

        for line in text.splitlines():
            m = self._FILE.match(line)
            if m:
                cur = FileChange(filepath=m.group(2).strip())
                files.append(cur)
                continue

            if cur is None:
                continue

            if self._NEW.match(line):
                cur.kind = 'added'
            elif self._DEL.match(line):
                cur.kind = 'deleted'
            elif self._REN.match(line):
                cur.kind = 'renamed'
                cur.filepath = self._REN.match(line).group(1).strip()
            elif self._HUNK.match(line) or self._PLUS3.match(line) or self._MIN3.match(line):
                continue
            elif line.startswith('+'):
                s = line[1:].strip()
                if s:
                    cur.added.append(s)
            elif line.startswith('-'):
                s = line[1:].strip()
                if s:
                    cur.removed.append(s)

        return [f for f in files if f.filepath]


# ─────────────────────────────────────────────────────────────────────────────
# ANALIZADOR GENÉRICO  — toda la inteligencia está aquí
# ─────────────────────────────────────────────────────────────────────────────

# Categorías de archivos por extensión/patrón
FILE_CATEGORIES: Dict[str, str] = {
    '.component.html':   'Template Angular',
    '.component.ts':     'Componente Angular',
    '.component.scss':   'Estilos de componente',
    '.service.ts':       'Servicio Angular',
    '.spec.ts':          'Test unitario',
    '.module.ts':        'Módulo Angular',
    '.pipe.ts':          'Pipe Angular',
    '.directive.ts':     'Directiva Angular',
    '.guard.ts':         'Guard Angular',
    '.interceptor.ts':   'Interceptor Angular',
    '.html':             'Template HTML',
    '.scss':             'Estilos SCSS',
    '.css':              'Estilos CSS',
    '.ts':               'TypeScript',
    '.js':               'JavaScript',
    '.py':               'Python',
    '.json':             'Configuración JSON',
    '.md':               'Documentación',
    '.yml':              'Configuración YAML',
    '.yaml':             'Configuración YAML',
    '.env':              'Variables de entorno',
    '.sql':              'SQL',
}

# Señales genéricas en líneas de código → descripción de impacto
# Cada entrada: (regex_en_linea, descripcion)  — se evalúan sobre líneas añadidas y eliminadas
IMPACT_SIGNALS: List[Tuple[re.Pattern, str]] = [
    # Observables / RxJS
    (re.compile(r'\.subscribe\s*\(\s*\{'),         'Migración a sintaxis de observer objeto (next/error/complete)'),
    (re.compile(r'catchError|throwError'),          'Manejo de errores con operadores RxJS'),
    (re.compile(r'takeUntilDestroyed|takeUntil'),   'Gestión del ciclo de vida de suscripciones'),
    # CSS / layout
    (re.compile(r'has-data'),                       'Clase condicional para estado con/sin datos'),
    (re.compile(r'scrollHeight.*flex|flex.*scrollHeight', re.I), 'Scroll flex en tabla PrimeNG'),
    (re.compile(r'calc\(100vh'),                    'Cálculo de altura relativa al viewport'),
    (re.compile(r'overflow\s*:\s*hidden'),          'Control de desbordamiento de contenedor'),
    # Variables de entorno / config
    (re.compile(r'apiUrl|API_URL|baseUrl'),          'Cambio de URL de API'),
    (re.compile(r'enableLogs|debug|DEBUG'),          'Cambio en configuración de logs/debug'),
    (re.compile(r'production\s*:\s*(true|false)'),   'Cambio de flag de producción'),
    # Angular
    (re.compile(r'@Injectable|providedIn'),          'Servicio inyectable'),
    (re.compile(r'HttpClient|HttpParams'),           'Comunicación HTTP'),
    (re.compile(r'Observable|Subject|BehaviorSubject'), 'Uso de observables RxJS'),
    (re.compile(r'router\.navigate'),               'Navegación programática'),
    (re.compile(r'@Input\(\)|@Output\(\)'),          'Comunicación entre componentes'),
    (re.compile(r'ngOnInit|ngOnDestroy'),            'Ciclo de vida del componente'),
    # Seguridad / auth
    (re.compile(r'AuthService|authService'),         'Lógica de autenticación'),
    (re.compile(r'token|Token|JWT'),                 'Gestión de tokens de acceso'),
    # Código obsoleto / limpieza
    (re.compile(r'console\.(log|warn|error)'),       'Llamadas a console (log/warn/error)'),
    (re.compile(r'debugMode|debug_mode'),            'Flag de modo debug'),
    (re.compile(r'TODO|FIXME|HACK'),                 'Comentarios pendientes (TODO/FIXME)'),
    # Exportación / descarga
    (re.compile(r'responseType.*blob|blob.*responseType', re.I), 'Descarga de archivos (Blob)'),
    (re.compile(r'exportExcel|exportCSV|export.*xlsx', re.I),    'Exportación a Excel/CSV'),
    # Tests
    (re.compile(r'describe\(|it\(|expect\('),        'Casos de prueba unitaria'),
    (re.compile(r'beforeEach|afterEach'),            'Setup/teardown de tests'),
]


def analyze_impact(fc: FileChange) -> str:
    """
    Genera una descripción de impacto 100% a partir del contenido real del diff.
    Sin supuestos hardcodeados — solo observa líneas añadidas y eliminadas.
    """
    all_added   = "\n".join(fc.added)
    all_removed = "\n".join(fc.removed)
    all_lines   = all_added + "\n" + all_removed

    found: List[str] = []
    seen:  set = set()

    for pattern, description in IMPACT_SIGNALS:
        if description not in seen and pattern.search(all_lines):
            found.append(description)
            seen.add(description)

    if found:
        return "; ".join(found)

    # Fallback genérico basado en tipo de archivo y volumen de cambios
    category = FILE_CATEGORIES.get(fc.ext, 'Archivo')
    n_add = len(fc.added)
    n_del = len(fc.removed)

    if fc.kind == 'added':
        return f"Nuevo archivo — {category} ({n_add} líneas)"
    if fc.kind == 'deleted':
        return f"Archivo eliminado — {category} ({n_del} líneas)"
    if fc.kind == 'renamed':
        return f"Renombrado — {category}"
    if n_del > n_add * 2:
        return f"Refactor — {category} (−{n_del} / +{n_add} líneas)"
    if n_add > 0 and n_del == 0:
        return f"Código añadido — {category} (+{n_add} líneas)"
    if n_del > 0 and n_add == 0:
        return f"Código eliminado — {category} (−{n_del} líneas)"
    return f"Modificación — {category} (+{n_add} / −{n_del} líneas)"


def analyze_recommendations(changes: List[FileChange]) -> List[str]:
    """
    Genera recomendaciones genéricas observando los patrones reales del diff.
    No asume nada — solo lee lo que está en los cambios.
    """
    recs: List[str] = []
    seen: set = set()

    def add(r: str):
        if r not in seen:
            recs.append(r)
            seen.add(r)

    all_added   = "\n".join(l for f in changes for l in f.added).lower()
    all_removed = "\n".join(l for f in changes for l in f.removed).lower()
    all_lines   = all_added + "\n" + all_removed

    exts = {f.ext for f in changes}
    kinds = {f.kind for f in changes}

    # ── Observaciones sobre los cambios ──────────────────────────────────────

    # Console.log eliminados
    if re.search(r'console\.(log|warn|error)', all_removed):
        add("Se eliminaron llamadas a console.log/warn/error. "
            "Verificar que no queden trazas de depuración en otros archivos del módulo.")

    # Suscripciones RxJS migradas
    if '.subscribe({' in all_added and '.subscribe(' in all_removed:
        add("Se migró al patrón de observer objeto en suscripciones RxJS. "
            "Aplicar el mismo patrón en todas las suscripciones del proyecto para consistencia.")

    # Manejo de errores añadido
    if 'error:' in all_added and ('error' not in all_removed or 'catch' not in all_removed):
        add("Se añadió manejo de errores (error callback). "
            "Confirmar que el mensaje de error sea informativo para el usuario final.")

    # URL de API modificada
    if 'apiurl' in all_lines or 'api_url' in all_lines or 'baseurl' in all_lines:
        add("Se detectó un cambio en la URL de API o configuración de entorno. "
            "Verificar que los entornos de staging y producción tengan sus propios valores "
            "y que no haya URLs hardcodeadas en otros componentes.")

    # Logs desactivados
    if 'enablelogs' in all_lines or 'debug' in all_lines:
        add("Se modificó la configuración de logs/debug. "
            "Confirmar que enableLogs: false esté presente en environment.prod.ts.")

    # Nuevos servicios
    new_services = [f for f in changes if f.kind == 'added' and '.service.ts' in f.ext]
    if new_services:
        names = ", ".join(f.filename for f in new_services)
        add(f"Nuevo(s) servicio(s) añadido(s): {names}. "
            "Asegurarse de registrarlos en el módulo correspondiente si no usan providedIn: 'root'.")

    # Archivos eliminados
    deleted = [f for f in changes if f.kind == 'deleted']
    if deleted:
        names = ", ".join(f.filename for f in deleted)
        add(f"Archivo(s) eliminado(s): {names}. "
            "Verificar que no existan importaciones huérfanas en otros archivos.")

    # Cambios en estilos
    style_changes = [f for f in changes if f.ext in ('.scss', '.css', '.component.scss')]
    if style_changes:
        add("Se modificaron archivos de estilos. "
            "Validar el layout en resoluciones mobile (320px), tablet (768px) y desktop (1440px).")

    # Cambios en templates
    template_changes = [f for f in changes if f.ext in ('.html', '.component.html')]
    if template_changes:
        add("Se modificaron templates HTML. "
            "Revisar que no haya expresiones de template con errores en producción "
            "(activar ng build --prod para verificar).")

    # Clase condicional has-data
    if 'has-data' in all_added:
        add("Se aplicó clase condicional basada en datos. "
            "Replicar el mismo patrón en todos los componentes con tablas que usen altura fija, "
            "para evitar zonas vacías cuando no hay registros.")

    # scrollHeight flex
    if 'scrollheight' in all_lines and 'flex' in all_lines:
        add('Se actualizó scrollHeight a "flex" en tabla PrimeNG. '
            "Confirmar que el contenedor padre tenga una altura definida, "
            "ya que scrollHeight='flex' requiere que el contenedor tenga height o flex para funcionar.")

    # Tests modificados
    test_changes = [f for f in changes if '.spec.ts' in f.ext]
    if test_changes:
        add("Se modificaron tests unitarios. Ejecutar ng test para verificar que todos pasen.")

    # Archivos de configuración
    config_changes = [f for f in changes if f.ext in ('.json', '.yml', '.yaml', '.env')]
    if config_changes:
        add("Se modificaron archivos de configuración. "
            "Confirmar que los cambios estén replicados en todos los entornos (dev, staging, prod) "
            "y que no contengan credenciales en texto plano.")

    # Si no se detectó nada específico
    if not recs:
        add("Revisar que los cambios sean compatibles con el resto de módulos de la aplicación.")
        add("Ejecutar las pruebas unitarias antes del merge para validar que no hay regresiones.")
        add("Verificar el comportamiento en distintos navegadores y resoluciones de pantalla.")

    # Siempre recomendado
    has_ui = any(f.ext in ('.html', '.component.html', '.scss', '.css', '.component.scss')
                 for f in changes)
    if has_ui:
        add("Realizar pruebas visuales en múltiples resoluciones antes del merge.")

    return recs


# ─────────────────────────────────────────────────────────────────────────────
# HELPERS PARA CONSTRUIR EL DOCX
# ─────────────────────────────────────────────────────────────────────────────

def _set_bg(cell, hex_color: str):
    tcPr = cell._tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def _set_borders(cell, color: str = "CCCCCC"):
    tcPr    = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:val"),   "single")
        el.set(qn("w:sz"),    "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        borders.append(el)
    tcPr.append(borders)


def _set_width(cell, cm: float):
    tcPr = cell._tc.get_or_add_tcPr()
    tcW  = OxmlElement("w:tcW")
    tcW.set(qn("w:w"),    str(int(cm * 567)))
    tcW.set(qn("w:type"), "dxa")
    tcPr.append(tcW)


def _run(para, text: str, bold=False, italic=False,
         color: RGBColor = None, size: float = 10, font: str = "Arial"):
    r = para.add_run(text)
    r.bold        = bold
    r.italic      = italic
    r.font.size   = Pt(size)
    r.font.name   = font
    if color:
        r.font.color.rgb = color
    return r


def _header_row(table, cols: List[Tuple[str, float]]):
    row = table.rows[0]
    for i, (text, w) in enumerate(cols):
        cell = row.cells[i]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        _set_bg(cell, C_HDR_BG)
        _set_borders(cell, C_ACCENT)
        _set_width(cell, w)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(p, text, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))


def _data_row(table, cells: List[Tuple[str, float, RGBColor, bool, str]]):
    """cells: (text, width_cm, text_color, bold, bg_hex)"""
    row = table.add_row()
    for i, (text, w, tc, bold, bg) in enumerate(cells):
        cell = row.cells[i]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        _set_bg(cell, bg)
        _set_borders(cell)
        _set_width(cell, w)
        lines = text.split("\n")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for j, line in enumerate(lines):
            if j > 0:
                p = cell.add_paragraph()
            _run(p, line, bold=bold, color=tc, size=9)


def _bullet(doc, text: str, symbol: str, color: RGBColor, indent: float = 1.0):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent       = Cm(indent)
    p.paragraph_format.first_line_indent = Cm(-0.5)
    p.paragraph_format.space_after       = Pt(3)
    _run(p, f"{symbol}  ", bold=True, color=color, size=10)
    _run(p, text, color=color, size=10)


def _section_title(doc, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(5)
    _run(p, text, bold=True, color=C_TITLE, size=12)
    pPr    = p._p.get_or_add_pPr()
    pBdr   = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), C_ACCENT)
    pBdr.append(bottom)
    pPr.append(pBdr)


def _divider(doc):
    p   = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "4")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), C_BORDER)
    pBdr.append(bot)
    pPr.append(pBdr)


def _h1(doc, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(7)
    _run(p, text, bold=True, color=C_TITLE, size=14)


# ─────────────────────────────────────────────────────────────────────────────
# GENERADOR
# ─────────────────────────────────────────────────────────────────────────────
class ReportGenerator:

    def __init__(self, changes: List[FileChange],
                 branch_from: str, branch_to: str, output: str):
        self.changes     = changes
        self.branch_from = branch_from
        self.branch_to   = branch_to
        self.output      = output
        self.doc         = Document()
        self._page_setup()

    def _page_setup(self):
        sec = self.doc.sections[0]
        sec.page_width    = Inches(8.5)
        sec.page_height   = Inches(11)
        sec.left_margin   = Cm(2.5)
        sec.right_margin  = Cm(2.5)
        sec.top_margin    = Cm(2.0)
        sec.bottom_margin = Cm(2.0)
        style           = self.doc.styles["Normal"]
        style.font.name = "Arial"
        style.font.size = Pt(10)

    # ── Portada ───────────────────────────────────────────────────────────────
    def _title(self):
        fecha = datetime.now().strftime("%d de %B de %Y")

        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(3)
        _run(p, "INFORME DE CAMBIOS DE CÓDIGO", bold=True, color=C_TITLE, size=20)

        p2 = self.doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_after = Pt(2)
        _run(p2, f"Rama: {self.branch_from}  →  {self.branch_to}",
             color=C_SUBTITLE, size=11)

        p3 = self.doc.add_paragraph()
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p3.paragraph_format.space_after = Pt(12)
        _run(p3, f"Fecha: {fecha}", color=C_MUTED, size=10)

        # línea decorativa
        p4 = self.doc.add_paragraph()
        p4.paragraph_format.space_after = Pt(12)
        pPr = p4._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bot  = OxmlElement("w:bottom")
        bot.set(qn("w:val"), "single"); bot.set(qn("w:sz"), "12")
        bot.set(qn("w:space"), "1");    bot.set(qn("w:color"), C_ACCENT)
        pBdr.append(bot); pPr.append(pBdr)

    # ── Sección 1: Resumen ────────────────────────────────────────────────────
    def _summary(self):
        _h1(self.doc, "1. Tabla Resumen de Cambios")

        # estadísticas globales antes de la tabla
        total_add = sum(len(f.added)   for f in self.changes)
        total_del = sum(len(f.removed) for f in self.changes)
        p = self.doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        _run(p, f"Archivos afectados: {len(self.changes)}   "
                f"Líneas añadidas: +{total_add}   "
                f"Líneas eliminadas: −{total_del}",
             color=C_MUTED, size=9, italic=True)

        COLS = [("Archivo", 4.8), ("Categoría", 3.0),
                ("Tipo de Cambio", 2.6), ("Impacto detectado", 4.6)]
        tbl = self.doc.add_table(rows=1, cols=len(COLS))
        tbl.style = "Table Grid"
        _header_row(tbl, COLS)

        for fc in self.changes:
            tc, bg = fc.kind_colors
            category = FILE_CATEGORIES.get(fc.ext, Path(fc.filepath).suffix.upper() or "Archivo")
            _data_row(tbl, [
                (fc.filename,        4.8, C_MOD_TEXT, True,  "EFF6FF"),
                (category,           3.0, C_BODY,     False, C_ROW_ALT),
                (fc.kind_label,      2.6, tc,         True,  bg),
                (analyze_impact(fc), 4.6, C_BODY,     False, C_WHITE),
            ])

        self.doc.add_paragraph()

    # ── Sección 2: Detalle ────────────────────────────────────────────────────
    def _detail(self):
        _h1(self.doc, "2. Detalle de Cambios por Archivo")

        for i, fc in enumerate(self.changes):
            _section_title(self.doc, fc.filename)

            # ruta completa
            p = self.doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            _run(p, fc.filepath, color=C_MUTED, size=8, italic=True)

            # estadísticas del archivo
            p2 = self.doc.add_paragraph()
            p2.paragraph_format.space_after = Pt(6)
            _run(p2, f"+{len(fc.added)} líneas añadidas   "
                     f"−{len(fc.removed)} líneas eliminadas",
                 color=C_MUTED, size=8)

            if fc.added:
                _run(self.doc.add_paragraph(),
                     "Líneas añadidas:", bold=True, color=C_ADD_TEXT, size=9)
                for line in fc.added:
                    _bullet(self.doc, line, "✔", C_ADD_TEXT)

            if fc.removed:
                p3 = self.doc.add_paragraph()
                p3.paragraph_format.space_before = Pt(4)
                _run(p3, "Líneas eliminadas:", bold=True, color=C_DEL_TEXT, size=9)
                for line in fc.removed:
                    _bullet(self.doc, line, "✖", C_DEL_TEXT)

            if not fc.added and not fc.removed:
                p4 = self.doc.add_paragraph()
                _run(p4,
                     "Sin cambios de contenido detectados "
                     "(cambio de permisos, renombrado o binario).",
                     color=C_MUTED, size=9, italic=True)

            if i < len(self.changes) - 1:
                _divider(self.doc)

        self.doc.add_paragraph()

    # ── Sección 3: Recomendaciones ────────────────────────────────────────────
    def _recommendations(self):
        _h1(self.doc, "3. Recomendaciones antes del Merge")

        recs = analyze_recommendations(self.changes)

        p = self.doc.add_paragraph()
        p.paragraph_format.space_after = Pt(7)
        _run(p,
             "Las siguientes recomendaciones se generaron automáticamente a partir "
             "del análisis del contenido de los cambios:",
             color=C_BODY, size=10)

        for rec in recs:
            _bullet(self.doc, rec, "➤", C_MOD_TEXT)

        self.doc.add_paragraph()

    # ── Pie ───────────────────────────────────────────────────────────────────
    def _footer(self):
        _divider(self.doc)
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(6)
        total_add = sum(len(f.added)   for f in self.changes)
        total_del = sum(len(f.removed) for f in self.changes)
        _run(p,
             f"Archivos: {len(self.changes)}  ·  "
             f"+{total_add} líneas  ·  −{total_del} líneas  ·  "
             "Informe Generado automaticamente con Python - Ryu Gabo -",
             color=C_MUTED, size=8, italic=True)

    # ── Entrada ───────────────────────────────────────────────────────────────
    def generate(self) -> str:
        self._title()
        self._summary()
        self._detail()
        self._recommendations()
        self._footer()
        self.doc.save(self.output)
        return self.output


# ─────────────────────────────────────────────────────────────────────────────
# BÚSQUEDA DEL ARCHIVO DE ENTRADA
# ─────────────────────────────────────────────────────────────────────────────
def find_input(arg: Optional[str]) -> Path:
    if arg:
        p = Path(arg)
        if not p.exists():
            print(f"[ERROR] No se encontró: {p}", file=sys.stderr)
            sys.exit(1)
        return p

    for candidate in (Path.cwd() / DEFAULT_INPUT,
                      Path(__file__).parent / DEFAULT_INPUT):
        if candidate.exists():
            return candidate

    print(
        f"[ERROR] No se encontró '{DEFAULT_INPUT}'.\n\n"
        f"  Genera el archivo con:\n"
        f"    git --no-pager diff --staged > {DEFAULT_INPUT}\n\n"
        f"  O indícalo con:\n"
        f"    python git_diff_to_docx.py --input <archivo.txt>",
        file=sys.stderr
    )
    sys.exit(1)


# ─────────────────────────────────────────────────────────────────────────────
# CLI
# ─────────────────────────────────────────────────────────────────────────────
def main():
    ap = argparse.ArgumentParser(
        description=(
            "Convierte 'informe.txt' (git --no-pager diff --staged) "
            "en un informe .docx profesional."
        ),
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog=__doc__
    )
    ap.add_argument("--input",       "-i",  default=None,
                    help=f"Archivo de diff (default: '{DEFAULT_INPUT}' en carpeta actual)")
    ap.add_argument("--output",      "-o",  default=None,
                    help=f"Archivo de salida (default: '{DEFAULT_OUTPUT}' junto al .txt)")
    ap.add_argument("--branch-from", "-bf", default="feature/cambios",
                    help="Rama origen  (default: feature/cambios)")
    ap.add_argument("--branch-to",   "-bt", default="develop",
                    help="Rama destino (default: develop)")
    args = ap.parse_args()

    # Localizar
    input_path = find_input(args.input)
    print(f"[INFO] Leyendo: {input_path}")

    raw  = input_path.read_text(encoding="utf-8", errors="replace")
    text = clean(raw)

    if not text.strip():
        print("[ERROR] El archivo está vacío.", file=sys.stderr)
        sys.exit(1)

    # Parsear
    changes = DiffParser().parse(text)
    if not changes:
        print(
            "[AVISO] No se encontraron cambios.\n"
            "        Comprueba que el archivo fue generado con:\n"
            "          git --no-pager diff --staged > informe.txt",
            file=sys.stderr
        )
        sys.exit(1)

    print(f"[INFO] {len(changes)} archivo(s) detectado(s):")
    for fc in changes:
        print(f"       [{fc.kind_label:12}] {fc.filepath}"
              f"  +{len(fc.added)} / -{len(fc.removed)}")

    # Generar
    output = args.output or str(input_path.parent / DEFAULT_OUTPUT)
    result = ReportGenerator(
        changes=changes,
        branch_from=args.branch_from,
        branch_to=args.branch_to,
        output=output
    ).generate()

    print(f"\n[OK] Informe generado: {result}")


if __name__ == "__main__":
    main()